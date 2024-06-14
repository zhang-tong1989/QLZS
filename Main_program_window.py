from PyQt5.QtMultimedia import *# 播放视频库
from PyQt5.QtMultimediaWidgets import QVideoWidget# 播放视频库
from PyQt5.QtCore import QUrl# 播放视频库
from PyQt5 import QtCore, QtGui, QtWidgets,uic
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import QHeaderView,QTableWidgetItem
from PyQt5.QtDesigner import *
from PyQt5.QtCore import QThread, pyqtSignal
from PyQt5.QtCore import QThread, pyqtSignal,QDateTime
import PyQt5.sip
import PyQt5.QAxContainer
from PyQt5.QtGui import QColor, QBrush
from PyQt5.QtGui import QKeyEvent
from PyQt5.QtWidgets import QApplication, QWidget, QMainWindow, QPushButton, QVBoxLayout,QMessageBox,QMenu,QUndoStack
from PyQt5.QtWidgets import QDialog
import subprocess
import sys,openpyxl,os,time,datetime,re,copy,pyperclip
from pathlib import Path
import json
from openpyxl.styles import Font
from openpyxl.styles import Alignment
from openpyxl.styles.borders import Border, Side
import multiprocessing
from multiprocessing import Process, Queue
import base64
import requests
from docx import Document
import threading# 线程模块
from pdf2docx import Converter
from PyPDF2 import PdfReader, PdfWriter
import pyttsx4# 文字转语音
from io import BytesIO
# from pydub import AudioSegment
# from pydub.playback import play
# import comtypes
# import comtypes.client
from aip import AipSpeech
class Thread_write_old(QThread):# 多线程写入Excel
    qmut = QMutex()
    runSing=pyqtSignal(int)
    stopSing=pyqtSignal()
    def __init__(self):
        super(Thread_write_old, self).__init__()
    def run(self):
        try:
            self.qmut.lock()  # 加锁
            win.write_old_Excel()
            self.qmut.unlock()  # 解锁
            self.stopSing.emit()
        except:
            pass
class Work(QObject):# 匹配编清单多线程
    qmut = QMutex()
    startSing = pyqtSignal()
    runSing=pyqtSignal(int)
    stopSing=pyqtSignal()
    def __init__(self):
        super().__init__()
        self.is_running = True
    def run_time_save(self):# 保存
        self.qmut.lock()  # 加锁
        print('进入循环')
        # time.sleep(10)
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_intelligent_qdk(self):# 多线程匹配编清单
        self.qmut.lock()  # 加锁
        win.intelligent_qdk()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_hand_qdk(self):# 多线程手动编清单
        self.qmut.lock()  # 加锁
        win.hand_qdk()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_import_dek(self):# 多线程加载定额库
        self.qmut.lock()  # 加锁
        print(11)
        win.import_dek('加载定额库')
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_hand_quota(self):# 多线程手动套定额
        self.qmut.lock()  # 加锁
        win.hand_quota()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_copy_name(self):# 按名称匹配
        self.qmut.lock()  # 加锁
        win.copy_name()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_name_sp(self):# 按名称规格匹配
        self.qmut.lock()  # 加锁
        win.name_sp()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_check_name(self):# 按系统名称规格匹配
        self.qmut.lock()  # 加锁
        win.check_sps_name()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_all_qdk(self):# 复用标准清单
        self.qmut.lock()  # 加锁
        win.all_qdk()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_write_new_Excel(self):# 复用标准清单
        self.qmut.lock()  # 加锁
        win.write_new_Excel()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_merge_unit(self):# 重合并单位工程
        self.qmut.lock()  # 加锁
        win.merge_unit()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_show_name(self):  # 多线程显示名称
        self.qmut.lock()  # 加锁
        win.show_name_row()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_system_Clicked(self):  # 多线点击系统写入名称和清单名称
        self.qmut.lock()  # 加锁
        win.system_Clicked()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_row_height(self):  #自适应行高
        self.qmut.lock()  # 加锁
        win.row_height()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_tool_show_row(self):  # 点击按钮显示隐藏的行
        self.qmut.lock()  # 加锁
        win.tool_show_row()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_show_rows(self):  # 点击整个工程多线点击系统写入名称和清单名称
        self.qmut.lock()  # 加锁
        win.show_rows()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_tool_paste(self):  # 粘贴
        self.qmut.lock()  # 加锁
        win.tool_paste()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_write_sys(self):  # 重读系统
        self.qmut.lock()  # 加锁
        win.write_sys()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_argv(self):  # 多线程打开文件
        self.qmut.lock()  # 加锁
        win.argv()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_delelte_text(self):  # 删除文本内容
        self.qmut.lock()  # 加锁
        win.delelte_text()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_del_space(self):  # 清除空格
        self.qmut.lock()  # 加锁
        win.del_space()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
class items_un_re(QObject):  #删除单位工程撤销回撤多线程类
    qmut = QMutex()
    startSing = pyqtSignal()
    stopSing = pyqtSignal()
    def __init__(self,items_old,items_mew,text_old_dict,text_new_dict):
        super().__init__()
        self.items_old = items_old
        self.items_mew = items_mew
        self.text_old_dict=text_old_dict
        self.text_new_dict=text_new_dict
    def run_undo(self):  # 撤销
        self.qmut.lock()  # 加锁,
        undo=items_tableCommand(self.items_old,self.items_mew,self.text_old_dict,self.text_new_dict)
        undo.do()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
    def run_redo(self):  # 回撤
        self.qmut.lock()  # 加锁,
        redo=items_tableCommand(self.items_old,self.items_mew,self.text_old_dict,self.text_new_dict)
        redo.red()
        self.qmut.unlock()  # 解锁
        self.stopSing.emit()
class items_tableCommand(QUndoCommand):# 删除单位工程撤销类
    def __init__(self,items_old,items_mew,text_old_dict,text_new_dict):
        super(items_tableCommand, self).__init__()
        self.items_old = items_old
        self.items_mew = items_mew
        self.text_old_dict=text_old_dict
        self.text_new_dict=text_new_dict
        # print(self.items_old,self.items_mew,self.text_old_dict,self.text_new_dict)
    # 回撤
    def redo(self):
        self.startThread_run_redo()
    def startThread_run_redo(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork =items_un_re(self.items_old,self.items_mew,self.text_old_dict,self.text_new_dict)  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_redo)
        self.Mywork.stopSing.connect(self.stopThread_redo)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
    def stopThread_redo(self):
        self.thread.quit()  # 退出
        self.thread.wait()# 回收资源
    def red(self):
        if self.items_old!=self.items_mew:
            win.table_do = '不执行'
            win.window.treeWidget_Items.topLevelItem(0).takeChildren()
            win.window.stackedWidget.setCurrentIndex(0)  # 设置显示页面
            for son in self.items_mew:
                win.son = QTreeWidgetItem(win.window.treeWidget_Items.topLevelItem(0))
                # self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                win.son.setText(0, son[0])
                win.son.setText(1, son[1])
                win.son.setCheckState(0, Qt.Unchecked)
                win.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
                win.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                win.son.setSizeHint(0, QSize(0, 30))
                font = QFont()
                font.setPointSize(10)  # 设置字体大小为10像素
                font.setFamily("宋体")
                win.son.setFont(0, font)
            for self.table,self.text_new_list in self.text_new_dict.items():
                if self.text_new_list!=[]:
                    self.table.clearContents()
                    self.table.setRowCount(len(self.text_new_list))  # 设置行数
                    for self.Single_rows, value in enumerate(self.text_new_list):
                        if '《定额》' in value:
                            for self.Single_colum ,text_new in enumerate(value):
                                self.item_new=QTableWidgetItem(str(text_new))
                                # self.item_new.setFlags(Qt.ItemIsEnabled)内容不能编辑
                                self.item_new.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                brush.setStyle(QtCore.Qt.SolidPattern)
                                self.item_new.setBackground(brush)  # 背景颜色
                                brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                brush.setStyle(QtCore.Qt.SolidPattern)
                                self.item_new.setForeground(brush)# 字体颜色QtCore.Qt.NoBrush
                                self.table.setItem(self.Single_rows, self.Single_colum, self.item_new)
                        else:
                            for self.Single_colum ,text_new in enumerate(value):
                                self.item_new=QTableWidgetItem(str(text_new))
                                if self.Single_colum == 7:
                                    self.item_new.setCheckState(Qt.Unchecked)
                                self.table.setItem(self.Single_rows, self.Single_colum,self.item_new)
                    self.table.viewport().update()  # 刷新tab内容
            win.table_do = '执行'
        # 回撤
        if self.items_old==self.items_mew:
            win.table_do = '不执行'
            for self.table,self.text_new_list in self.text_new_dict.items():
                self.table.clearContents()
                self.table.setRowCount(len(self.text_new_list))  # 设置行数
                for self.Single_rows,value in enumerate(self.text_new_list):
                    if '《定额》' in value:
                        for self.Single_colum ,text_new in enumerate(value):
                            self.item_new=QTableWidgetItem(str(text_new))
                            # self.item_new.setFlags(Qt.ItemIsEnabled)# 内容不能编辑
                            self.item_new.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                            brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                            brush.setStyle(QtCore.Qt.SolidPattern)
                            self.item_new.setBackground(brush)  # 背景颜色
                            brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                            brush.setStyle(QtCore.Qt.SolidPattern)
                            self.item_new.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                            self.table.setItem(self.Single_rows, self.Single_colum, self.item_new)
                    else:
                        for self.Single_colum, text_new in enumerate(value):
                            self.item_new = QTableWidgetItem(str(text_new))
                            if self.Single_colum == 7:
                                self.item_new.setCheckState(Qt.Unchecked)
                            self.table.setItem(self.Single_rows, self.Single_colum, self.item_new)

                self.table.viewport().update()  # 刷新tab内容
            win.table_do = '执行'
    def undo(self):
        self.startThread_run_undo()
    def startThread_run_undo(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork =items_un_re(self.items_old,self.items_mew,self.text_old_dict,self.text_new_dict)  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_undo)
        self.Mywork.stopSing.connect(self.stopThread_undo)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
    def stopThread_undo(self):
        self.thread.quit()  # 退出
        self.thread.wait()# 回收资源
    def do(self):
        if self.items_old!=self.items_mew:
            win.table_do = '不执行'
            win.window.treeWidget_Items.topLevelItem(0).takeChildren()
            win.window.stackedWidget.setCurrentIndex(0)  # 设置显示页面
            for son in self.items_old:
                win.son = QTreeWidgetItem(win.window.treeWidget_Items.topLevelItem(0))
                # self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                win.son.setText(0, son[0])
                win.son.setText(1, son[1])
                win.son.setCheckState(0, Qt.Unchecked)
                win.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
                win.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                win.son.setSizeHint(0, QSize(0, 30))
                font = QFont()
                font.setPointSize(10)  # 设置字体大小为10像素
                font.setFamily("宋体")
                win.son.setFont(0, font)
            for self.table, self.text_old_list in self.text_old_dict.items():
                if self.text_old_list != []:
                    self.table.clearContents()
                    self.table.viewport().update()  # 刷新tab内容
                    self.table.setRowCount(len(self.text_old_list))  # 设置行数
                    for self.Single_rows, value in enumerate(self.text_old_list):
                        if '《定额》' in value:
                            for self.Single_colum, text_old in enumerate(value):
                                self.item_old = QTableWidgetItem(str(text_old))
                                self.item_old.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                brush.setStyle(QtCore.Qt.SolidPattern)
                                self.item_old.setBackground(brush)  # 背景颜色
                                brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                brush.setStyle(QtCore.Qt.SolidPattern)
                                self.item_old.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                                self.table.setItem(self.Single_rows, self.Single_colum, self.item_old)
                        else:
                            for self.Single_colum, text_old in enumerate(value):
                                self.item_old = QTableWidgetItem(str(text_old))
                                if self.Single_colum == 7:
                                    if text_old != '《定额》':
                                        self.item_old.setCheckState(Qt.Unchecked)
                                self.table.setItem(self.Single_rows, self.Single_colum, self.item_old)
                    self.table.viewport().update()  # 刷新tab内容
            win.table_do = '执行'
        if self.items_old==self.items_mew:
            print('撤销执行')
            win.table_do = '不执行'
            for self.table,self.text_old_list in self.text_old_dict.items():
                self.table.clearContents()
                self.table.setRowCount(len(self.text_old_list))  # 设置行数
                for self.Single_rows,value in enumerate(self.text_old_list):
                    if '《定额》' in value:
                        for self.Single_colum ,text_old in enumerate(value):
                            self.item_old=QTableWidgetItem(str(text_old))
                            self.item_old.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                            brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                            brush.setStyle(QtCore.Qt.SolidPattern)
                            self.item_old.setBackground(brush)  # 背景颜色
                            brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                            brush.setStyle(QtCore.Qt.SolidPattern)
                            self.item_old.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                            self.table.setItem(self.Single_rows, self.Single_colum, self.item_old)
                    else:
                        for self.Single_colum ,text_old in enumerate(value):
                            self.item_old=QTableWidgetItem(str(text_old))
                            if self.Single_colum == 7:
                                self.item_old.setCheckState(Qt.Unchecked)
                            self.table.setItem(self.Single_rows, self.Single_colum,self.item_old)
                self.table.viewport().update()  # 刷新tab内容
            win.table_do = '执行'
class Windows():#主窗口类
    def __init__(self):
        BASE_DIR = os.path.dirname(__file__)
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        print(sys.argv[0])
        file_path = os.path.join(BASE_DIR, "Main_Window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('清量计价助手V2.0')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.85)
        self.width = int(self.screenwidth * 0.85)
        self.window.resize(self.width, self.height)
        self.window.closeEvent = self.closeEvent# 实现关闭窗口弹出保存对话框
        # 拆分
        self.window.splitter.setStretchFactor(0, 35)
        self.window.splitter.setStretchFactor(1, 65)
        # self.window.splitter.setStyleSheet("QSplitter::handle { background-color: green }")
        progreesbar = os.path.join(BASE_DIR, "progreesbar.ui")
        self.pwindow = uic.loadUi(progreesbar)
        self.pwindow.setWindowModality(Qt.ApplicationModal)  # 阻塞主窗口不能点击WindowModal
        self.pwindow.setWindowFlags(Qt.WindowStaysOnTopHint)  # 窗口始终在前面
        self.pwindow.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.WindowCloseButtonHint)  # 禁止最大化
        self.pwindow.setWindowFlags(Qt.FramelessWindowHint)# 设置窗口无边框
        # 设置窗口背景透明
        self.pwindow.setAttribute(Qt.WA_TranslucentBackground, True)
        # 设置透明度，0为全透明
        self.pwindow.setWindowOpacity(1)
        self.window.treeWidget_Items.setColumnWidth(0, 600)#0列列宽
        self.window.treeWidget_Items.setColumnWidth(1, 1)
        self.window.toolButton_save_items.clicked.connect(self.json_save)#保存工程
        self.window.toolButton_save_items.setShortcut('ctrl+S')
        self.window.pushButton_new_item.clicked.connect(self.new_Item) # 新建单项工程
        self.window.pushButton_del.clicked.connect(self.new_del)  # 删除工程
        self.window.pushButton_Excel.clicked.connect(self.Excel_values)#导入工程量
        self.window.toolButton_quantity.clicked.connect(self.import_files)  # 导入非标工程量
        self.window.pushButton_qdk.clicked.connect(self.Excel_qdk)  # 加载清单库
        self.window.toolButton_import_dek.clicked.connect(self.startThread_run_import_dek)  # 加载定额库
        self.window.pushButton_intelligentqdk.clicked.connect(self.startThread_run_intelligent_qdk)#匹配编清单,多线程运行
        self.window.pushButton_all_qdk.clicked.connect(self.startThread_run_all_qdk)  # 复用标准清单
        self.window.toolButton_hand_quota.clicked.connect(self.startThread_run_hand_quota)  # 手动套定额
        self.window.toolButton_specifications.clicked.connect(self.specifications)  # 解析规格
        self.window.pushButton_save_old.clicked.connect(self.startThread_write_old_Excel)#写入原Excel,多线程运行
        self.window.pushButton_save_old.setEnabled(True)  # 禁用按钮
        self.window.pushButton_Export_Excel.clicked.connect(self.write_new_Excel)#导出到Excel
        self.window.pushButton_Export_Excel.setEnabled(True)  # 禁用按钮
        self.window.toolButton_standard_save.clicked.connect(self.write_zero_Excel)  # 导出到Excel
        self.window.toolButton_standard_save.setEnabled(True)  # 禁用按钮
        self.window.toolButton_copy_excel.clicked.connect(self.copy_excel)  # 复制到Excel
        self.window.toolButton_copy_excel.setEnabled(True)  # 禁用按钮
        # self.window.pushButton_Export_Excel.setEnabled(True)# 按钮不禁用
        self.window.pushButton_create_table.clicked.connect(self.create_table)#创建表格
        self.window.pushButton_clear_table.clicked.connect(self.clear_table)  # 创建表格
        self.window.toolButton_delelte_text.clicked.connect(self.startThread_run_delelte_text)  # 删除文本
        self.window.toolButton_delelte_text.setShortcut('delete')
        self.window.toolButton_del_quota.clicked.connect(self.del_quota)  # 删除文本
        self.window.toolButton_del_space.clicked.connect(self.startThread_run_del_space)  # 清除空格
        self.window.pushButton_insert_rows.clicked.connect(self.insert_rows)#插入行
        self.window.pushButton_del_rows.clicked.connect(self.del_rows) #删除行
        self.window.toolButton_row_height.clicked.connect(self.startThread_run_row_height)  # 自适应行高
        self.cwd = os.getcwd()#获取当前文件所在位置
        self.window.pushButton_open_qdk.clicked.connect(self.open_qdk_window)
        self.window.pushButton_copy_names.clicked.connect(self.copy_names) #复用匹配清单名称
        self.window.toolButton_sum.clicked.connect(self.sum_quantity)  # 工程量求和
        self.window.toolButton_part_sum.clicked.connect(self.part_sum)  # 局部求和
        self.window.toolButton_pdf.clicked.connect(self.pdf_tool)  # PDF工具
        self.window.toolButton_text_speck.clicked.connect(self.speck_tool)  # 语音工具
        self.window.toolButton_text_speck.setEnabled(False)  # 禁用按钮
        self.window.toolButton_all_checks.clicked.connect(self.all_checks)  # 检查
        self.window.pushButton_merge_unit.clicked.connect(self.startThread_run_merge_unit)#合并单位工程
        self.window.pushButton_show_name.clicked.connect(self.show_name)#名称筛选
        self.window.toolButton_filter_sp.clicked.connect(self.filter_sp)  # 规格筛选
        # self.window.pushButton_copyn.setCheckable(True)#按钮默认被选中
        self.window.pushButton_select_allname.clicked.connect(self.select_allname)#全选名称
        self.window.pushButton_Cancel_all.clicked.connect(self.unselect_allname)#取消全选
        self.window.pushButton_find_replace.clicked.connect(self.find_replace)#查找替换
        self.window.pushButton_find_replace.setShortcut('ctrl+F')
        self.window.toolButton_shear_text.clicked.connect(self.shear_text)#剪切文本
        self.window.toolButton_shear_text.setShortcut('ctrl+X')
        self.window.toolButton_copy.clicked.connect(self.tool_copy)#复制文本
        self.window.toolButton_copy.setShortcut('ctrl+C')
        self.window.toolButton_paste.clicked.connect(self.startThread_run_tool_paste)#粘贴文本
        self.window.toolButton_paste.setShortcut('ctrl+V')
        # 剪切板
        self.clipboard = QApplication.clipboard()
        self.clipboard.dataChanged.connect(self.Clipboard)# 剪切板内容发生变化连接函数
        self.window.toolButton_copy_row.clicked.connect(self.copy_row)  # 复制行
        self.window.toolButton_shear_row.clicked.connect(self.shear_row)  # 剪切行
        self.window.toolButton_paste_row.clicked.connect(self.paste_row)  # 粘贴行
        self.window.toolButton_hide_rows.clicked.connect(self.choise_rows)  # 隐藏行
        menu_row= QMenu()
        image1 = os.path.join(BASE_DIR, 'image', "隐藏行.png")
        hide_row = QAction(QIcon(image1),'隐藏行',menu_row)
        menu_row.addAction(hide_row)
        hide_row.setData("隐藏行")
        hide_row.triggered.connect(self.hide_rows)
        image2 = os.path.join(BASE_DIR, 'image', "显示行.png")
        show_row = QAction(QIcon(image2),'显示行',menu_row)
        menu_row.addAction(show_row)
        show_row.setData("显示行")
        show_row.triggered.connect(self.startThread_run_tool_show_row)
        menu_row.addSeparator()  # 分隔符
        self.window.toolButton_hide_rows.setMenu(menu_row)
        menu_row.triggered.connect(self.menu_rows_clicked)
        self.window.toolButton_hide_column.clicked.connect(self.choise_column)  #隐藏列
        menu_colum = QMenu()
        image1 = os.path.join(BASE_DIR, 'image', "隐藏列.png")
        hide_column = QAction(QIcon(image1),'隐藏列',menu_colum)
        menu_colum.addAction(hide_column)
        hide_column.setData("隐藏列")
        hide_column.triggered.connect(self.hide_column)
        image2 = os.path.join(BASE_DIR, 'image', "显示列.png")
        show_column = QAction(QIcon(image2),'显示列',menu_colum)
        menu_colum.addAction(show_column)
        show_column.setData("显示列")
        show_column.triggered.connect(self.show_column)
        menu_colum.addSeparator()  # 分隔符
        self.window.toolButton_hide_column.setMenu(menu_colum)
        menu_colum.triggered.connect(self.menu_clicked)
        self.window.lineEdit.setPlaceholderText('输入名称，点击Enter')  # 提示文本
        self.window.lineEdit.returnPressed.connect(self.up_check)  # 绑定enter键
        self.window.lineEdit_specification.setPlaceholderText('输入规格，点击Enter')
        self.window.lineEdit_specification.returnPressed.connect(self.specification_check)  # 绑定enter
        self.window.pushButton_match_name.clicked.connect(self.match_name)  # 提取清单匹配名称
        self.window.listWidget_sys.itemDoubleClicked.connect(self.D_click) # 双击
        self.window.pushButton_unit_sys.clicked.connect(self.unit_sys)  #合并系统
        self.window.pushButton_rewrite_name.clicked.connect(self.write_sys)
        self.window.stackedWidget.setCurrentIndex(0)
        self.window.label_url_01.setText("<A href='https://v.qq.com/x/page/t3552ncvcmp.html'>打开网页学习</a>")
        self.window.label_url_01.setOpenExternalLinks(True)
        self.window.label_url_02.setText("<A href='https://v.qq.com/x/page/w3552y59051.html'>打开网页学习</a>")
        self.window.label_url_02.setOpenExternalLinks(True)
        self.window.label_url_03.setText("<A href='https://v.qq.com/x/page/i35527cn3nh.html'>打开网页学习</a>")
        self.window.label_url_03.setOpenExternalLinks(True)
        self.window.label_url_04.setText("<A href='https://v.qq.com/x/page/i3553anudcr.html'>打开网页学习</a>")
        self.window.label_url_04.setOpenExternalLinks(True)
        self.window.label_url_05.setText("<A href='https://v.qq.com/x/page/w3553hgzh7q.html'>打开网页学习</a>")
        self.window.label_url_05.setOpenExternalLinks(True)
        self.window.label_url_06.setText("<A href='https://v.qq.com/x/page/d3553j7f3pr.html'>打开网页学习</a>")
        self.window.label_url_06.setOpenExternalLinks(True)
        self.window.label_url_07.setText("<A href='https://v.qq.com/x/page/f3553pcnj4b.html'>打开网页学习</a>")
        self.window.label_url_07.setOpenExternalLinks(True)
        self.window.label_url_08.setText("<A href='https://item.taobao.com/item.htm?ft=t&id=785883260964&skuId=5364859856981&spm=a21dvs.23580594.0.0.4fee645en5fSGB'>打开网页购买软件</a>")
        self.window.label_url_08.setOpenExternalLinks(True)
        self.window.label_url_09.setText("<A href='https://pan.baidu.com/s/1GM_9XWMG_bvD8zi7I-ewng?pwd=r6bd'>下载软件和教程</a>")
        self.window.label_url_09.setOpenExternalLinks(True)
        self.font = Font(
            name="宋体",  # 字体linkActivated
            size=9,  # 字体大小
            color="000000",  # 字体颜色，用16进制rgb表示
            bold=False,  # 是否加粗，True/False
            italic=False,  # 是否斜体，True/False
            strike=None,  # 是否使用删除线，True/False
            underline=None,)  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
        self.alig = Alignment()
        self.alig.wrap_text = True# 自动换行
        self.alig.vertical = 'center'# 垂直方向居中对齐
        # 设置边框样式
        # l_side =openpyxl.styles.borders.Side(style='dashDot', color=None)
        # r_side = openpyxl.styles.borders.Side(style='dashDotDot', color=None)
        # t_side = openpyxl.styles.borders.Side(style='dashed', color=None)
        # b_side = openpyxl.styles.borders.Side(style='mediumDashDot', color=None)
        # style  == 'dashDot'左四划线, 'dashDotDot'左三划线, 'dashed'左六划线, 'mediumDashDot'右三划线, 'double'右七双划线, 'slantDashDot',右二划线 'thin'左六实线, 'hair',左一划线 'dotted',左二划线
        # 'thick',右六粗线 'mediumDashed',右四点划线 'mediumDashDotDot'右一点划线, 'medium'右五中粗实线
        l_side =Side(style='thin', color=None)
        r_side = Side(style='thin', color=None)
        t_side = Side(style='thin', color=None)
        b_side =Side(style='thin', color=None)
        self.border = Border(left=l_side, right=r_side, top=t_side, bottom=b_side, vertical=l_side)
        # 设置标签
        font = QFont()
        font.setFamily("宋体")
        font.setPointSize(9)
        # 树窗口
        self.window.treeWidget_Items.expandAll()
        self.window.treeWidget_Items.itemClicked.connect(self.onClicked)
        self.window.treeWidget_Items.hideColumn(1)# 隐藏第一列
        self.window.treeWidget_Items.hideColumn(2)  # 隐藏第二列
        self.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
        self.window.treeWidget_Items.itemChanged.connect(self.tree_check)
        self.window.treeWidget_Items.setStyleSheet("QHeaderView::section{background:rgb(196, 223, 255);}")# 列表头颜色
        self.window.treeWidget_system.itemChanged.connect(self.check_unit_sys)
        self.window.treeWidget_system.itemClicked.connect(self.startThread_run_system_Clickede)
        self.window.treeWidget_system.hideColumn(1)  # 隐藏第一列
        self.window.treeWidget_system.setStyleSheet("QHeaderView::section{background:rgb(196, 223, 255);}")
        self.window.treeWidget_name.itemChanged.connect(self.check_names)
        self.window.treeWidget_specification.itemChanged.connect(self.check_specification)
        self.window.treeWidget_name.setStyleSheet("QHeaderView::section{background:rgb(196, 223, 255);}")
        # self.window.treeWidget_Items.setSortingEnabled(True)排序
        # 树窗口添加右键
        self.window.treeWidget_Items.setContextMenuPolicy(Qt.CustomContextMenu)
        self.window.treeWidget_Items.customContextMenuRequested.connect(self.show_right)  # 绑定菜单
        self.undoStack_del = QUndoStack() #删除单位工程撤销回撤
        self.undoStack_del.setUndoLimit(100)
        undoAction_del = self.undoStack_del.createUndoAction(self.undoStack_del, "Undo")
        undoAction_del.setShortcut("Ctrl+Z")
        redoAction_del = self.undoStack_del.createRedoAction(self.undoStack_del, "Redo")
        redoAction_del.setShortcut("Ctrl+y")
        self.window.addAction(undoAction_del)
        self.window.addAction(redoAction_del)
        self.window.toolButton_image_text.clicked.connect(self.image_text)  # 图片文字识别
        self.tablelist = [self.window.tableWidget_0,self.window.tb_1, self.window.tb_2, self.window.tb_3, self.window.tb_4, self.window.tb_5, self.window.tb_6, self.window.tb_7
                          , self.window.tb_8, self.window.tb_9, self.window.tb_10, self.window.tb_11, self.window.tb_12, self.window.tb_13
            , self.window.tb_14, self.window.tb_15, self.window.tb_16, self.window.tb_17, self.window.tb_18, self.window.tb_19, self.window.tb_20
            , self.window.tb_21, self.window.tb_22, self.window.tb_23, self.window.tb_24, self.window.tb_25, self.window.tb_26, self.window.tb_27
            , self.window.tb_28, self.window.tb_29, self.window.tb_30, self.window.tb_31, self.window.tb_32, self.window.tb_33, self.window.tb_34
            , self.window.tb_35, self.window.tb_36, self.window.tb_37, self.window.tb_38, self.window.tb_39, self.window.tb_40, self.window.tb_41
            , self.window.tb_42, self.window.tb_43, self.window.tb_44, self.window.tb_45, self.window.tb_46, self.window.tb_47, self.window.tb_48, self.window.tb_49
            , self.window.tb_50, self.window.tb_51, self.window.tb_52, self.window.tb_53, self.window.tb_54,self.window.tb_55, self.window.tb_56, self.window.tb_57
            , self.window.tb_58, self.window.tb_59, self.window.tb_60, self.window.tb_61, self.window.tb_62,self.window.tb_63, self.window.tb_64
            , self.window.tb_65, self.window.tb_66, self.window.tb_67, self.window.tb_68, self.window.tb_69,self.window.tb_70, self.window.tb_71
            , self.window.tb_72, self.window.tb_73, self.window.tb_74, self.window.tb_75, self.window.tb_76, self.window.tb_77, self.window.tb_78
            , self.window.tb_79, self.window.tb_80, self.window.tb_81, self.window.tb_82, self.window.tb_83,self.window.tb_84, self.window.tb_85
            , self.window.tb_86, self.window.tb_87, self.window.tb_88, self.window.tb_89, self.window.tb_90, self.window.tb_91, self.window.tb_92, self.window.tb_93
            , self.window.tb_94, self.window.tb_95, self.window.tb_96, self.window.tb_97, self.window.tb_98,self.window.tb_99, self.window.tb_100,
                          ]
        for i,tab in enumerate(self.tablelist):
            tab.setColumnCount(12)
            # tab.horizontalHeader()
            tab.setHorizontalHeaderLabels(
                ['清单编码', '清单名称', '项目特征', '清单单位', '匹配清单名称', '匹配定额名称', '专业/系统',
                 '名称','规格', '单位', '工程量', '备注'])
            tab.setFocusPolicy(Qt.StrongFocus)
            tab.setStyleSheet("gridline-color: rgb(257, 1, 0)")
            tab.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)  # 设置列宽，列宽可调
            tab.horizontalHeader().resizeSection(0, 120)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(1, 120)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(2, 200)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(3, 120)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(4, 150)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(5, 150)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(6, 150)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(7, 150)  # 调整第11列的大小为500像素
            tab.horizontalHeader().resizeSection(8, 150)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(9,100)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(10, 100)  # 调整第2列的大小为500像素
            tab.horizontalHeader().resizeSection(11,100)  # 调整第2列的大小为500像素
            tableHeader =tab.horizontalHeader()
            tableHeader.setStretchLastSection(False)
            verticalHeader =tab.verticalHeader()
            verticalHeader.setStretchLastSection(False)
            # tab.itemClicked.connect(self.Click_Select)  # 单击获取单元格中的内容
            tab.itemDoubleClicked.connect(self.row_Interactive)  # 双击行高可调
            tab.itemDoubleClicked.connect(self.open_Text_window)  # 双击加载文本修改窗口
            # tab.cellChanged.connect(self.undo_redo)  # 改变单元格内容发出信号，连接槽撤销回撤。
            # 表格添加右键
            tab.setContextMenuPolicy(Qt.CustomContextMenu)
            tab.customContextMenuRequested.connect(self.table_show_right)  # 绑定菜单
            tab.viewport().update()# 刷新tab内容
        self.timer = QTimer(self.window)
        self.timer.start(1000)
        now_time=QDateTime.currentDateTime()#获取实时时间
        formatted_time=now_time.toString('yyyy-MM-dd hh:mm:ss')
        if formatted_time>= '2024-07-30 11:43:03':
            self.window.toolButton_standard_save.setEnabled(False)
            self.window.pushButton_save_old.setEnabled(False)
            self.window.pushButton_Export_Excel.setEnabled(False)
            self.window.toolButton_copy_excel.setEnabled(False)
        if sys.argv.__len__() < 2:
            if self.json_save_list==[]:
                self.timer.singleShot(180000, self.startThread_run_time_save)
        if sys.argv.__len__() >= 2:
            self. startThread_run_argv()
            self.json_save_list.append(sys.argv[1])
            if os.path.exists(sys.argv[1]):
                filename =os.path.join(os.path.abspath('.'),sys.argv[1])
                filest = time.localtime(os.stat(filename).st_ctime)#文件创建时间
                filest_time=time.strftime("%Y-%m-%d %H:%M:%S", filest)
                filemt = time.localtime(os.stat(filename).st_mtime)#文件修改时间
                filemt_time=time.strftime("%Y-%m-%d %H:%M:%S", filemt)
                if filest_time>=filemt_time:
                    self.window.toolButton_standard_save.setEnabled(False)
                    self.window.pushButton_save_old.setEnabled(False)
                    self.window.pushButton_Export_Excel.setEnabled(False)
                    self.window.toolButton_copy_excel.setEnabled(False)
                if filest_time>='2024-07-30 11:43:03':
                    self.window.toolButton_standard_save.setEnabled(False)
                    self.window.pushButton_save_old.setEnabled(False)
                    self.window.pushButton_Export_Excel.setEnabled(False)
                    self.window.toolButton_copy_excel.setEnabled(False)
                if filemt_time>='2024-07-30 11:43:03':
                    self.window.toolButton_standard_save.setEnabled(False)
                    self.window.pushButton_save_old.setEnabled(False)
                    self.window.pushButton_Export_Excel.setEnabled(False)
                    self.window.toolButton_copy_excel.setEnabled(False)
    def menu_rows_clicked(self,action):
        BASE_DIR = os.path.dirname(__file__)
        if action.data()=='隐藏行':
            image1 = os.path.join(BASE_DIR, 'image', "隐藏行.png")
            self.window.toolButton_hide_rows.setIcon(QIcon(image1))
            self.window.toolButton_hide_rows.setText(action.data())
        if action.data()=='显示行':
            image2 = os.path.join(BASE_DIR, 'image', "显示行.png")
            self.window.toolButton_hide_rows.setIcon(QIcon(image2))
            self.window.toolButton_hide_rows.setText(action.data())
    def choise_rows(self):
        but_text=self.window.toolButton_hide_rows.text()
        if but_text=='隐藏行':
            self.hide_rows()
        if but_text=='显示行':
            self.startThread_run_tool_show_row()
    def menu_clicked(self,action):
        BASE_DIR = os.path.dirname(__file__)
        if action.data()=='隐藏列':
            image1 = os.path.join(BASE_DIR, 'image', "隐藏列.png")
            self.window.toolButton_hide_column.setIcon(QIcon(image1))
            self.window.toolButton_hide_column.setText(action.data())
        if action.data()=='显示列':
            image2 = os.path.join(BASE_DIR, 'image', "显示列.png")
            self.window.toolButton_hide_column.setIcon(QIcon(image2))
            self.window.toolButton_hide_column.setText(action.data())
    def choise_column(self):
        but_text=self.window.toolButton_hide_column.text()
        if but_text=='隐藏列':
            self.hide_column()
            print(but_text)
        if but_text=='显示列':
            self.show_column()
            print(but_text)
    old_undo_dict = {}
    tab_change='提取'
    def Click_Select(self):#单击获取表格内任意值，并且或撤销值
        pass
    table_do='执行'
    new_undo_dict={}
    def undo_redo(self):#文本内容变化后转让撤销类中
        pass
    def startThread_run_argv(self):# 多线程打开文件
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_argv)
        self.Mywork.stopSing.connect(self.stopThread_argv)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
    def stopThread_argv(self):  #
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
        self.Mywork.deleteLater()
        self.thread.deleteLater()
    def argv(self):# 打开工程文件
        if os.path.exists(sys.argv[1]):
            self.tab_change = '不提取'
            jsonPath =os.path.join(os.path.abspath('.'),sys.argv[1])
            with open(jsonPath, 'r') as f:
                self.data_json = json.load(f)
                self.python_data = json.loads(self.data_json)  # json转python
                for value in self.python_data:
                    for key, value in value[0].items():
                        print(str(key).split('$')[1], value)
                        if str(key).split('$')[0] == '建设项目':
                            self.tableWidget = self.tablelist[0]
                            self.window.treeWidget_Items.topLevelItem(0).setText(2, str(key).split('$')[1])
                            for row, va in enumerate(value):
                                self.tableWidget.setRowCount(row + 1)  # 设置行数
                                if '《定额》' in va:
                                    for column, text in enumerate(va):
                                        self.tableWidget.setItem(row, column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        self.item1.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                        brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setBackground(brush)  # 背景颜色
                                        brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                                        self.tableWidget.setItem(row, column, self.item1)
                                else:
                                    for column, text in enumerate(va):
                                        self.tableWidget.setItem(row, column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        if column == 7:
                                            self.item1.setCheckState(Qt.Unchecked)
                                        self.tableWidget.setItem(row, column, self.item1)
                        else:
                            self.son = QTreeWidgetItem(self.window.treeWidget_Items.topLevelItem(0))
                            self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                            self.son.setText(0, str(key).split('$')[0])
                            self.son.setText(1, str(key).split('$')[1])
                            self.son.setCheckState(0, Qt.Unchecked)
                            self.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
                            self.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                            self.son.setSizeHint(0, QSize(0, 30))
                            font = QFont()
                            font.setPointSize(10)  # 设置字体大小为10像素
                            font.setFamily("宋体")
                            self.son.setFont(0, font)
                            self.tableWidget = self.tablelist[int(str(key).split('$')[1])]
                            for row, va in enumerate(value):
                                self.tableWidget.setRowCount(row + 1)  # 设置行数
                                if '《定额》' in va:
                                    for column ,text in enumerate(va):
                                        self.tableWidget.setItem(row,column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        self.item1.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                        brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setBackground(brush)  # 背景颜色
                                        brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                                        self.tableWidget.setItem(row, column, self.item1)
                                else:
                                    for column, text in enumerate(va):
                                        self.tableWidget.setItem(row, column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        if column == 7:
                                            self.item1.setCheckState(Qt.Unchecked)
                                        self.tableWidget.setItem(row, column, self.item1)
            self.tab_change = '提取'
    json_save_list=[]
    def json_save(self): # 保存工程
        if self.json_save_list!=[]:
            save_json = []
            self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:
                n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
                for i in range(0, n):
                    ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                    name = ite.text(0)
                    num_max = ite.text(2)
                    self.table0_json = []
                    save_json.append(self.table0_json)
                    table0_alldict = {}
                    table0_allvalues = []
                    for self.Single_rows in range(0, self.window.tableWidget_0.rowCount()):
                        table0_values = []
                        table0_allvalues.append(table0_values)
                        for self.Single_colum in range(0, self.window.tableWidget_0.columnCount()):
                            self.text0 = self.window.tableWidget_0.item(self.Single_rows, self.Single_colum).text()
                            table0_values.append(self.text0)
                    table0_alldict[name + '$' + num_max] = table0_allvalues
                    self.table0_json.append(table0_alldict)
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    for j in range(0, count):
                        self.table_json = []
                        save_json.append(self.table_json)
                        table_alldict = {}
                        table_allvalues = []
                        string = ite.child(j)  # 子节点的文字信息
                        self.item_zero = string.text(0)  # 获取0列内容
                        self.item_zer1 = string.text(1)  # 获取1列内容
                        self.item = self.item_zero, self.item_zer1
                        self.tableWidget = self.tablelist[int(self.item_zer1)]
                        self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                        self.table_column = int(self.tableWidget.columnCount())
                        for self.Single_rows in range(0, self.tableWidget_allrows):
                            table_values = []
                            table_allvalues.append(table_values)
                            for self.Single_colum in range(0, self.table_column):
                                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                table_values.append(self.text)
                        table_alldict[self.item_zero + '$' + self.item_zer1] = table_allvalues
                        self.table_json.append(table_alldict)
                        # print(self.table_json[0])
                self.data_json = json.dumps(save_json, ensure_ascii=False)  # python转json文。
                # BASE_DIR = os.path.dirname(__file__)
                # self.files_save = BASE_DIR + '/' + "json_save" + '/' + self.item_zero + '$' + self.item_zer1 + '.json'
                with open(self.json_save_list[-1], "w") as f:
                    json.dump(self.data_json, f)  # 禁止ascii转换，这样就可以打印中文。
            QMessageBox.information(self.window, '温馨提示', '工程保存成功')
        if self.json_save_list==[]:
            self.fileName_choose, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", self.cwd,'*.ZJB2.0')
            if self.fileName_choose == "":
                print("\n取消选择")
                return
            if self.fileName_choose!="":
                self.json_save_list.append(self.fileName_choose)
                save_json=[]
                # self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
                # if int(self.Item_row) != -1:
                n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
                for i in range(0, n):
                    ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                    name=ite.text(0)
                    num_max=ite.text(2)
                    self.table0_json = []
                    save_json.append(self.table0_json)
                    table0_alldict = {}
                    table0_allvalues = []
                    for self.Single_rows in range(0, self.window.tableWidget_0.rowCount()):
                        table0_values = []
                        table0_allvalues.append(table0_values)
                        for self.Single_colum in range(0,self.window.tableWidget_0.columnCount()):
                            self.text0 = self.window.tableWidget_0.item(self.Single_rows, self.Single_colum).text()
                            table0_values.append(self.text0)
                    table0_alldict[name + '$' + num_max] = table0_allvalues
                    self.table0_json.append(table0_alldict)
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    for j in range(0, count):
                        self.table_json = []
                        save_json.append(self.table_json)
                        table_alldict = {}
                        table_allvalues = []
                        string = ite.child(j)  # 子节点的文字信息
                        self.item_zero = string.text(0)  # 获取0列内容
                        self.item_zer1 = string.text(1)  # 获取1列内容
                        self.item = self.item_zero, self.item_zer1
                        self.tableWidget = self.tablelist[int(self.item_zer1)]
                        self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                        self.table_column = int(self.tableWidget.columnCount())
                        for self.Single_rows in range(0, self.tableWidget_allrows):
                            table_values = []
                            table_allvalues.append(table_values)
                            for self.Single_colum in range(0, self.table_column):
                                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                table_values.append(self.text)
                        table_alldict[self.item_zero + '$' + self.item_zer1] = table_allvalues
                        self.table_json.append(table_alldict)
                        # print(self.table_json[0])
                self.data_json = json.dumps(save_json, ensure_ascii=False)  # python转json文。
                with open(self.json_save_list[-1], "w") as f:
                    json.dump(self.data_json, f)  # 禁止ascii转换，这样就可以打印中文。
                QMessageBox.information(self.window, '温馨提示', '工程保存成功')
    def open(self):#打开工程
        jsonPath, _ = QFileDialog.getOpenFileName(self.window, '选择文件', ' ','文件类型(*.ZJB2.0)')
        if jsonPath:
            self.tab_change = '不提取'
            with open(jsonPath, 'r') as f:
                self.data_json = json.load(f)
                self.python_data = json.loads(self.data_json)  # json转python
                print(self.python_data)
                for value in self.python_data:
                    for key, value in value[0].items():
                        # print(str(key).split('$')[1], value)
                        if str(key).split('$')[0]=='建设项目':
                            self.tableWidget = self.tablelist[0]
                            self.window.treeWidget_Items.topLevelItem(0).setText(2, str(key).split('$')[1])
                            for row,va in enumerate(value):
                                self.tableWidget.setRowCount(row + 1)  # 设置行数
                                if '《定额》' in va:
                                    for column ,text in enumerate(va):
                                        self.tableWidget.setItem(row,column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        self.item1.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                        brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setBackground(brush)  # 背景颜色
                                        brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                                        self.tableWidget.setItem(row, column, self.item1)
                                else:
                                    for column, text in enumerate(va):
                                        self.item1 = QTableWidgetItem(str(text))
                                        self.tableWidget.setItem(row, column, QTableWidgetItem(''))
                                        if column == 7:
                                            self.item1.setCheckState(Qt.Unchecked)
                                        self.tableWidget.setItem(row,column, self.item1)

                        else:
                            self.son = QTreeWidgetItem(self.window.treeWidget_Items.topLevelItem(0))
                            self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                            self.son.setText(0, str(key).split('$')[0])
                            self.son.setText(1, str(key).split('$')[1])
                            self.son.setCheckState(0, Qt.Unchecked)
                            self.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
                            self.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                            self.son.setSizeHint(0, QSize(0, 30))
                            font = QFont()
                            font.setPointSize(10)  # 设置字体大小为10像素
                            font.setFamily("宋体")
                            self.son.setFont(0, font)
                            self.tableWidget = self.tablelist[int(str(key).split('$')[1])]
                            for row,va in enumerate(value):
                                self.tableWidget.setRowCount(row + 1)  # 设置行数
                                if '《定额》' in va:
                                    for column ,text in enumerate(va):
                                        self.tableWidget.setItem(row,column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        self.item1.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                        brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setBackground(brush)  # 背景颜色
                                        brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        self.item1.setForeground(brush)  # 字体颜色QtCore.Qt.NoBrush
                                        self.tableWidget.setItem(row, column, self.item1)
                                else:
                                    for column ,text in enumerate(va):
                                        self.tableWidget.setItem(row,column, QTableWidgetItem(''))
                                        self.item1 = QTableWidgetItem(str(text))
                                        if column == 7:
                                            self.item1.setCheckState(Qt.Unchecked)
                                        self.tableWidget.setItem(row,column, self.item1)
            self.tab_change = '提取'
    def show_right(self,pos):#绑定右键
        treeWidget_Menu = QMenu()
        BASE_DIR = os.path.dirname(__file__)
        image_path1 = os.path.join(BASE_DIR, 'image', "新建工程.png")
        new_items = QAction(QIcon(image_path1),'新建单位工程')
        treeWidget_Menu.addAction(new_items)
        new_items.triggered.connect(self.new_Item)

        image_path2 = os.path.join(BASE_DIR, 'image', "删除单位工程.png")
        openAct = QAction(QIcon(image_path2),'删除单位工程')
        treeWidget_Menu.addAction(openAct)
        openAct.triggered.connect(self.new_del)
        screenPos = self.window.treeWidget_Items.mapToGlobal(pos)#转换坐标系
        treeWidget_Menu.exec_(screenPos)

    # 建立工程
    def new_Item(self):
        self.window.stackedWidget.setCurrentIndex(0)
        self.tableWidget = self.tablelist[0]
        # self.tableWidget.setSelectionBehavior(QAbstractItemView.SelectRows)  # 选择整行
        if int(self.window.treeWidget_Items.topLevelItem(0).text(2))<=99:
            print(self.window.treeWidget_Items.topLevelItem(0).text(2))
            self.son = QTreeWidgetItem(self.window.treeWidget_Items.topLevelItem(0))
            self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled |  QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            self.son.setText(0, '新建工程')
            self.tre_2=self.window.treeWidget_Items.topLevelItem(0).text(2)
            self.tre_num=int(self.tre_2)+1
            self.window.treeWidget_Items.topLevelItem(0).setText(2,str(self.tre_num))
            self.son.setText(1, str(self.tre_num))
            self.son.setCheckState(0, Qt.Unchecked)
            self.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
            self.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
            self.son.setSizeHint(0, QSize(0, 30))
            self.window.stackedWidget.setCurrentIndex(self.tre_num)
            font = QFont()
            font.setPointSize(10)  # 设置字体大小为10像素
            font.setFamily("宋体")
            self.son.setFont(0, font)

    tabldict={}
    choose = '选择'
    def tree_check(self, item, cloumn):#树窗口选中状态
        if item.checkState(cloumn) == Qt.Checked:
            for check_num in range(0,len(self.tablelist)):
                if int(item.text(1))==check_num and int(item.text(1))!=0:#确保总工程不会提取表格内容
                    self.tabldict[(str(item.text(1)),str(item.text(0)))] = self.tablelist[check_num]
            # print(self.tabldict)
            # print(self.tabldict[item.text(1)])  # 用键来查找值
        if item.checkState(cloumn) == Qt.Unchecked:
            if (str(item.text(1)),str(item.text(0))) in self.tabldict:
                self.tabldict.pop((str(item.text(1)),str(item.text(0))))

    def new_del(self):# 删除单位工程
        self.tab_change = '不提取'
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            # 撤销
            self.old_undo_dict = {}
            self.items_old=[]
            n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
            for i in range(0, n):
                ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                count = ite.childCount()  # 获取当前根节点的子节点数量
                for j in range(0, count):
                    string = ite.child(j)  # 子节点的文字信息
                    self.item_zero = string.text(0)  # 获取0列内容
                    self.item_zer1 = string.text(1)  # 获取1列内容
                    self.item = self.item_zero, self.item_zer1
                    self.items_old.append(self.item)
                    self.tableWidget = self.tablelist[int(self.item_zer1)]
                    self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                    self.table_column = int(self.tableWidget.columnCount())
                    table_allvalues = []
                    self.old_undo_dict[self.tableWidget]=table_allvalues
                    for self.Single_rows in range(0, self.tableWidget_allrows):
                        table_values = []
                        table_allvalues.append(table_values)
                        for self.Single_colum in range(0, self.table_column):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            table_values.append(self.text)
            # 删除单位工程
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            if int(self.item_zero)!=0:
                self.tablelist[int(self.item_zero)].clearContents()  # 可以清除表格所有的内容
                self.tablelist[int(self.item_zero)].setRowCount(0)

                self.item = self.window.treeWidget_Items.currentItem()
                remove_item = str(self.item.text(1)),str(self.item.text(0))
                if remove_item in self.tabldict:
                    self.tabldict.pop(remove_item)
                root = self.window.treeWidget_Items.invisibleRootItem()
                for it in self.window.treeWidget_Items.selectedItems():
                    print(it)
                    (it.parent() or root).removeChild(it)
                self.window.stackedWidget.setCurrentIndex(0)
                self.write_sys()
                self.match_name()
            # 回撤
            self.new_undo_dict = {}
            self.items_new = []
            n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
            for i in range(0, n):
                ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                count = ite.childCount()  # 获取当前根节点的子节点数量
                for j in range(0, count):
                    string = ite.child(j)  # 子节点的文字信息
                    self.item_zero = string.text(0)  # 获取0列内容
                    self.item_zer1 = string.text(1)  # 获取1列内容
                    self.item = self.item_zero, self.item_zer1
                    self.items_new.append(self.item)
                    self.tableWidget = self.tablelist[int(self.item_zer1)]
                    self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                    self.table_column = int(self.tableWidget.columnCount())
                    table_allvalues = []
                    self.new_undo_dict[self.tableWidget] = table_allvalues
                    for self.Single_rows in range(0, self.tableWidget_allrows):
                        table_values = []
                        table_allvalues.append(table_values)
                        for self.Single_colum in range(0, self.table_column):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            table_values.append(self.text)
            if self.items_old!= self.items_new:
                command = items_tableCommand(self.items_old,self.items_new,self.old_undo_dict, self.new_undo_dict)
                self.window.treeWidget_Items.topLevelItem(0).takeChildren()
                self.tableWidget.viewport().update()  # 刷新tab内容
                self.undoStack_del.push(command)
            self.tab_change = '提取'
    def onClicked(self):# 点击树窗口获取table
        self.tab_change = '不提取'
        self.check_names_dict={}
        self.check_specification_dict={}
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:# 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            self.item = self.window.treeWidget_Items.currentItem()
            self.item_0=self.item.text(0)#获取0列内容
            self.item_1 = self.item.text(1)#获取1列内容
            self.window.treeWidget_Items.headerItem().setText(0,self.item_0)#表头写入内容
            for t in range(0,len(self.tablelist)):#点击单位工程显示page和table
                if int(self.item_1)==(t):
                    self.window.stackedWidget.setCurrentIndex(t)
                    self.tableWidget = self.tablelist[t]
                    # self.tableWidget.setSelectionBehavior(QAbstractItemView.SelectRows)# 选择整行
            print(self.item_0,self.item_1)
            if self.item_0=='建设项目':
                self.window.pushButton_unit_sys.setEnabled(True)  # 禁用按钮
            else:
                self.window.pushButton_unit_sys.setEnabled(False)
            self.write_sys()
            self.match_name()
            self.tab_change = '提取'
    def table_show_right(self, pos):#右键
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:# 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            table_Menu = QMenu()
            BASE_DIR = os.path.dirname(__file__)
            image1 = os.path.join(BASE_DIR, 'image', "插入行.png")
            item1 = QAction(QIcon(image1),'插入行')
            table_Menu.addAction(item1)
            item1.triggered.connect(self.insert_rows)

            image2 = os.path.join(BASE_DIR, 'image', "删除行.png")
            item2 = QAction(QIcon(image2),'删除行')
            table_Menu.addAction(item2)
            item2.triggered.connect(self.del_rows)

            image3 = os.path.join(BASE_DIR, 'image', "剪切.png")
            item3 = QAction(QIcon(image3),'剪切行')
            table_Menu.addAction(item3)
            item3.triggered.connect(self.shear_row)

            image4 = os.path.join(BASE_DIR, 'image', "复制行.png")
            item4 = QAction(QIcon(image4),'复制行')
            table_Menu.addAction(item4)
            item4.triggered.connect(self.copy_row)

            image5 = os.path.join(BASE_DIR, 'image', "粘贴行.png")
            item5 = QAction(QIcon(image5),'粘贴行')
            table_Menu.addAction(item5)
            item5.triggered.connect(self.paste_row)

            image13 = os.path.join(BASE_DIR, 'image', "剪切文本.png")
            item13 = QAction(QIcon(image13),'剪切文本')
            table_Menu.addAction(item13)
            item13.triggered.connect(self.shear_text)

            image6 = os.path.join(BASE_DIR, 'image', "复制.png")
            item6 = QAction(QIcon(image6),'复制文本')
            table_Menu.addAction(item6)
            item6.triggered.connect(self.tool_copy)

            image7 = os.path.join(BASE_DIR, 'image', "粘贴.png")
            item7 = QAction(QIcon(image7),'粘贴文本')
            table_Menu.addAction(item7)
            item7.triggered.connect(self.startThread_run_tool_paste)

            image8 = os.path.join(BASE_DIR, 'image', "删除内容.png")
            item8 = QAction(QIcon(image8),'删除文本')
            table_Menu.addAction(item8)
            item8.triggered.connect(self.startThread_run_delelte_text)

            image9 = os.path.join(BASE_DIR, 'image', "清除空格.png")
            item9 = QAction(QIcon(image9),'清除空格')
            table_Menu.addAction(item9)
            item9.triggered.connect(self.startThread_run_del_space)

            image10 = os.path.join(BASE_DIR, 'image', "全选名称.png")
            item10 = QAction(QIcon(image10),'勾选名称')
            table_Menu.addAction(item10)
            item10.triggered.connect(self.select_allname)

            image11 = os.path.join(BASE_DIR, 'image', "取消全选.png")
            item11 = QAction(QIcon(image11),'取消勾选')
            table_Menu.addAction(item11)
            item11.triggered.connect(self.unselect_allname)

            image12 = os.path.join(BASE_DIR, 'image', "删除定额.png")
            item12 = QAction(QIcon(image12),'删除定额')
            table_Menu.addAction(item12)
            item12.triggered.connect(self.del_quota)

            screenPos = self.tableWidget.mapToGlobal(pos)  # 转换坐标系
            table_Menu.exec_(screenPos)
    # 多线程运行函数
    def startThread_run_intelligent_qdk(self):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.all_filesqdk != []:
                self.thread = QThread()  # 实例化一个线程
                self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
                self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
                self.thread.started.connect(self.Mywork.run_intelligent_qdk)
                self.Mywork.runSing.connect(self.run_progressBar)
                self.Mywork.stopSing.connect(self.stopThread_intelligent_qdk)  # 停止信号连接到stopThread方法
                self.pwindow.show()
                self.pwindow.progressBar.setValue(0)
                self.thread.start()  # 开始线程的运行
            else:
                QMessageBox.information(self.window, '提示', '请加载清单库，再匹配编清单')
                return
    def stopThread_intelligent_qdk(self):# 匹配编清单
        if self.intell == '继续执行':
            # self.tableWidget.verticalHeader().resizeSection(0, 100)  # 调整每一行的大小为100像素
            self.pwindow.close()
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
            self.Mywork.deleteLater()
            self.thread.deleteLater()
        if self.intell == '不执行':
            self.pwindow.close()
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
            self.Mywork.deleteLater()
            self.thread.deleteLater()
            QMessageBox.information(self.window, '提示', '错误，请检查！')
    def startThread_run_hand_qdk(self):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1 :
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_hand_qdk)
            self.Mywork.stopSing.connect(self.stopThread_hand_qdk)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
    def stopThread_hand_qdk(self):
        if self.row_num=='继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
            # self.tableWidget.verticalHeader().setSectionResizeMode(self.row, QHeaderView.Interactive)
        elif self.row_num=='不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
    def startThread_write_old_Excel(self):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.old_filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.xlsx)')
            if self.old_filePaths != []:
                self.table_rows = self.tableWidget.rowCount()
                for self.old_file in self.old_filePaths:
                    hide_f = '(.*)/(.*)'
                    self.hide = re.compile(hide_f, re.S)
                    self.t_hied = self.hide.findall(self.old_file)
                    hidefilenames = self.t_hied[0][0] + '/' + '~$' + self.t_hied[0][1]
                    if os.path.exists(hidefilenames):
                        QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')
                        return
                self.worker =Thread_write_old()
                self.worker.runSing.connect(self.run_progressBar)  # 连接worker中的信号resultReady到主控窗口的handle_result去更新进度条
                self.worker.stopSing.connect(self.stopThread_write_old_Excel)
                self.pwindow.progressBar.setValue(0)
                self.pwindow.show()
                self.worker.start()
    def stopThread_write_old_Excel(self):
        if self.result=='文件打开了':
            self.pwindow.close()
            self.worker.quit()  # 退出
            self.worker.wait()# 回收资源
            self.worker.deleteLater()
            QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')
        if self.result=='执行完毕':
            self.pwindow.close()
            self.worker.quit()  # 退出
            self.worker.wait()# 回收资源
            self.worker.deleteLater()
            QMessageBox.information(self.window, '温馨提示', '数据写入完成，请检查。')
    def run_progressBar(self,msg):# 进度条传入数值
        self.pwindow.progressBar.setValue(msg)
    def startThread_run_copy_name(self):# 复制名称
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_copy_name)
            self.Mywork.stopSing.connect(self.stopThread_copy_name)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
    def stopThread_copy_name(self):
        if self.cn=='继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
        elif self.cn=='不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
    def startThread_run_name_sp(self):# 应用名称规格
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero) != 0:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_name_sp)
            self.Mywork.stopSing.connect(self.stopThread_name_sp)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
    def stopThread_name_sp(self):
        if self.ns=='继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
        elif self.ns=='不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
    def startThread_run_check_name(self):# 应用系统名称规格
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero) != 0:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_check_name)
            self.Mywork.stopSing.connect(self.stopThread_run_check_name)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
    def stopThread_run_check_name(self):
        if self.ns=='继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
        elif self.ns=='不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
    def startThread_run_all_qdk(self):# 复用标准清单
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_all_qdk)
            self.Mywork.runSing.connect(self.run_progressBar)
            self.Mywork.stopSing.connect(self.stopThread_run_all_qdk)  # 停止信号连接到stopThread方法
            self.pwindow.show()
            self.pwindow.progressBar.setValue(0)
            self.thread.start()  # 开始线程的运行
    def stopThread_run_all_qdk(self):
        if self.ns=='继续执行':
            self.pwindow.close()
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
            self.Mywork.deleteLater()
            self.thread.deleteLater()
            QMessageBox.information(self.window, '温馨提示', '复用清单完成，请检查。')
        elif self.ns=='不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
            self.Mywork.deleteLater()
            self.thread.deleteLater()
    def copy_excel(self):# 复制到excel
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            row = self.tableWidget.currentRow()  # 获取单元格行数
            column = self.tableWidget.currentColumn()
            self.text_dict = {}
            if str(row) != '-1':
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    contents = self.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                    # print('第',r.row(),'行','第',r.column(),'列',contents)
                    if r.row() not in self.text_dict:
                        self.text_dict[r.row()] = ['{}\t'.format(contents)]
                    elif r.row() in self.text_dict:
                        self.text_dict[r.row()].append('{}\t'.format(contents))
                if self.text_dict!={}:
                    all_text = ''
                    for values in self.text_dict.values():
                        values[-1]=str(values[-1]).replace('\t', '')#values是列表修最后一个元素
                        copy_text = ''
                        for value in values:
                            if len(str(value).split('\n'))>=2:
                                copy_text+='"{}"\t'.format(str(value).split('\t')[0])
                            if len(str(value).split('\n')) == 1:
                                copy_text+='{}'.format(str(value))
                        all_text +='{}\n'.format(str(copy_text))
                    self.clipboard.setText(all_text)
    def startThread_run_write_new_Excel(self):# 保存到Excel
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero) != 0:
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_write_new_Excel)
            self.Mywork.stopSing.connect(self.stopThread_run_write_new_Excel)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
    def stopThread_run_write_new_Excel(self):
        if self.result=='文件打开了':
            QMessageBox.information(self.window, '温馨提示','无法保存到Excel，请先关闭或删除同名的Excel文件')
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
        elif self.result=='执行完毕':
            QMessageBox.information(self.window, '温馨提示', '保存到Excel完成，请检查。')
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
        elif self.result=='取消':
            self.thread.quit()  # 退出
            self.thread.wait()# 回收资源
    def import_files(self):#加载非标工程量
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.file_window = file_window()
            self.file_window.window.show()
    def create_table(self):#创建表格
        self.tab_change = '不提取'
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[self.tableWidget] = self.Click_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            # 创建表格
            table_column = self.tableWidget.columnCount()
            # self.tableWidget.setRowCount(2000)  # 设置行数
            # rows=self.tableWidget.rowCount()
            for i in range(0,1000):
                self.tableWidget.setRowCount(i+1)  # 设置行数
                for j in range(0,table_column):
                    item= QTableWidgetItem('')
                    self.tableWidget.setItem(int(i), j, item)
                    if j ==7:
                        item.setCheckState(Qt.Unchecked)
                        self.tableWidget.setItem(int(i),7, item)
            # 回撤
            if self.table_do == '执行':
                self.new_undo_dict = {}
                new_text_list = []
                self.new_undo_dict[self.tableWidget] = new_text_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    new_text_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.Click_list != new_text_list:
                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                    self.undoStack_del.push(command)
            self.tab_change = '提取'
    def clear_table(self):# 清空表格
        self.tab_change = '不提取'
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            # 撤销
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[self.tableWidget] = self.Click_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if int(self.item_zero)>=0:
                self.tableWidget.clearContents()#可以清除表格所有的内容
                self.tableWidget.setRowCount(0)
                self.write_sys()
                self.match_name()
            # 回撤
            if self.table_do == '执行':
                self.new_undo_dict = {}
                new_text_list = []
                self.new_undo_dict[self.tableWidget] = new_text_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    new_text_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.Click_list != new_text_list:
                    print(new_text_list)
                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                    self.undoStack_del.push(command)
            self.tab_change = '提取'
    def insert_rows(self):#插入行
        self.tab_change = '不提取'
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            # 撤销
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[self.tableWidget] = self.Click_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            # 插入行
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if self.row!=-1:
                self.tableWidget.insertRow(self.row+1)
                for j in range(0,self.tableWidget.columnCount()):
                    item= QTableWidgetItem('')
                    self.tableWidget.setItem(self.row+1, j, item)
                    if j == 7:
                        item.setCheckState(Qt.Unchecked)
                        self.tableWidget.setItem(self.row+1, 7, item)
                # 回撤
                if self.table_do == '执行':
                    self.new_undo_dict={}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        print(new_text_list)
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'
    def startThread_run_row_height(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_row_height)
        self.Mywork.stopSing.connect(self.stopThread_row_height)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
    def stopThread_row_height(self):
        # QApplication.processEvents()  # 刷新软件
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
    def row_height(self):# 自适应行高
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            self.rows = self.tableWidget.rowCount()  # 获取单元格行数
            if self.rows!=0:
                for i in range(0,self.rows):
                    self.tableWidget.verticalHeader().setSectionResizeMode(i,QHeaderView.ResizeToContents) # 行高根据内容调整，但是行高不可调
    def del_quota(self):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tab_change = '不提取'
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row) != '-1':
                self.itemzero = self.window.treeWidget_Items.currentItem()
                self.item_zero = self.itemzero.text(1)  # 获取1列内容
                # if int(self.item_zero)!=0
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                self.rows = []
                self.column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    if r.row() not in self.rows:
                        self.rows.append(r.row())
                    self.column.append(r.column())
                self.rows.sort(reverse=False)
                num = 0
                for i in range(0, len(self.rows)):
                    if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                            and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                        if self.tableWidget.item(mytable[num].row(), 5).text() == '《定额》':
                            self.tableWidget.removeRow(mytable[num].row())
                        num += self.tableWidget.columnCount()
                self.write_sys()
                self.match_name()

                if self.table_do == '执行':
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'
    def del_rows(self):#删除行
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tab_change = '不提取'
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row) != '-1':
                self.itemzero = self.window.treeWidget_Items.currentItem()
                self.item_zero = self.itemzero.text(1)  # 获取1列内容
                # if int(self.item_zero)!=0
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                self.rows = []
                self.column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    if r.row() not in self.rows:
                        self.rows.append(r.row())
                    self.column.append(r.column())
                self.rows.sort(reverse=False)
                num = 0
                for i in range(0, len(self.rows)):
                    if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                            and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                        self.tableWidget.removeRow(mytable[num].row())
                        num+=self.tableWidget.columnCount()
                self.write_sys()
                self.match_name()

                if self.table_do == '执行':
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'
    def startThread_run_delelte_text(self):# 删除内容
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1 :
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row) != '-1':
                self.thread = QThread()  # 实例化一个线程
                self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
                self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
                self.thread.started.connect(self.Mywork.run_delelte_text)
                self.Mywork.stopSing.connect(self.stopThread_delelte_text)  # 停止信号连接到stopThread方法
                self.thread.start()  # 开始线程的运行
    def stopThread_delelte_text(self):
        self.thread.quit()  # 退出
        # self.thread.wait()  # 回收资源
    def delelte_text(self):# 删除文本内容
        self.tab_change = '不提取'
            # 撤销
        self.old_undo_dict = {}
        self.Click_list = []
        self.old_undo_dict[self.tableWidget] = self.Click_list
        for self.Single_rows in range(0, self.tableWidget.rowCount()):
            t_list = []
            self.Click_list.append(t_list)
            for self.Single_colum in range(0, self.tableWidget.columnCount()):
                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                t_list.append(self.text)
        # 删除文本
        mytable = self.tableWidget.selectedItems()
        for r in mytable:#第三种方法获取值
            self.item1 = QTableWidgetItem('')
            if r.column() == 7:
                self.item1.setCheckState(Qt.Unchecked)
            if self.tableWidget.item(r.row(), r.column()).text()!='《定额》':
                self.tableWidget.setItem(r.row(), r.column(), self.item1)
        self.tableWidget.viewport().update()  # 刷新tab内容
        # 撤回
        if self.table_do == '执行':
            self.new_undo_dict = {}
            new_text_list = []
            self.new_undo_dict[self.tableWidget] = new_text_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                new_text_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if self.Click_list != new_text_list:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                self.undoStack_del.push(command)
        self.tab_change = '提取'
    def startThread_run_del_space(self):# 清除空格
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1 :
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            print(self.row)
            if str(self.row) != '-1':
                self.thread = QThread()  # 实例化一个线程
                self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
                self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
                self.thread.started.connect(self.Mywork.run_del_space)
                self.Mywork.stopSing.connect(self.stopThread_del_space)  # 停止信号连接到stopThread方法
                self.thread.start()  # 开始线程的运行
    def stopThread_del_space(self):
        self.thread.quit()  # 退出
        # self.thread.wait()  # 回收资源
    def del_space(self):# 清除空格
        self.tab_change = '不提取'
        self.old_undo_dict = {}
        self.Click_list = []
        self.old_undo_dict[self.tableWidget] = self.Click_list
        for self.Single_rows in range(0, self.tableWidget.rowCount()):
            t_list = []
            self.Click_list.append(t_list)
            for self.Single_colum in range(0, self.tableWidget.columnCount()):
                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                t_list.append(self.text)
        #清除空格
        mytable = self.tableWidget.selectedItems()
        for r in mytable:#第三种方法获取值
            contents =self.tableWidget.item(r.row(), r.column()).text()#获取单元格内容
            contents_replace=str(contents).replace(' ','')
            self.item1= QTableWidgetItem(contents_replace)
            # print(r.row(),r.column(),contents_replace)
            if r.column() == 7:
                self.item1.setCheckState(Qt.Unchecked)
            self.tableWidget.setItem(r.row(), r.column(), self.item1)
        self.tableWidget.viewport().update()  # 刷新tab内容
        self.write_sys()
        self.match_name()
        if self.table_do == '执行':
            self.new_undo_dict = {}
            new_text_list = []
            self.new_undo_dict[self.tableWidget] = new_text_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                new_text_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if self.Click_list != new_text_list:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                self.undoStack_del.push(command)
        self.tab_change = '提取'
    all_files = []
    def Excel_values(self):#导入工程量
        self.tab_change = '不提取'
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1 :
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[self.tableWidget] = self.Click_list
            for self.Single_rows in range(0, self.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, self.tableWidget.columnCount()):
                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)>=0:
            self.filePaths, _ =QFileDialog.getOpenFileNames(self.window,'选择文件',' ','文件类型(*.xlsx)')
            if self.filePaths!=[]:
                messageBox = QMessageBox()
                messageBox.resize(300, 600)
                messageBox.setWindowTitle('选择框')
                messageBox.setText('"提示", "请选择导入表格的形式！"')
                messageBox.setStandardButtons(QMessageBox.Yes | QMessageBox.Ok | QMessageBox.Close)
                buttonYes = messageBox.button(QMessageBox.Yes)
                buttonYes.setText("追加导入")
                buttonOk = messageBox.button(QMessageBox.Ok)
                buttonOk.setText("清空导入")
                buttonC = messageBox.button(QMessageBox.Close)
                buttonC.setText('取消')
                messageBox.exec_()
                if messageBox.clickedButton() ==buttonYes:
                    self.all_files.append(self.filePaths)
                    self.tableWidget_row = int(self.tableWidget.rowCount())#获取总行数
                    self.num = 0
                    for self.files in self.filePaths:
                        self.wb = openpyxl.load_workbook(self.files, read_only=False, data_only=True, keep_links=False)
                        for self.ws in self.wb.worksheets:
                            self.rows = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                            max_num=len(self.rows)
                            for wb_ws_row in range(0, max_num):
                                data=self.rows[wb_ws_row]
                            #     Data_sources = self.files+',' + self.ws.title + '第{}行'.format(wb_ws_row+1)
                                self.num += 1
                                self.tableWidget.setRowCount(int(self.tableWidget_row)+self.num)#设置行数
                                for colum in range(0,self.tableWidget.columnCount()):
                                    self.tableWidget.setItem(int(self.num)+int(self.tableWidget_row)-1, colum, QTableWidgetItem(''))
                                for j in range(0, len(data)):
                                    self.item1 = QTableWidgetItem(str(data[j]).replace('None', ''))
                                    if j== 7:
                                        self.item1.setCheckState(Qt.Unchecked)
                                    self.tableWidget.setItem(int(self.num)+int(self.tableWidget_row)-1, int(j), self.item1)
                    self.tab_change = '提取'
                    self.write_sys()
                    self.match_name()
                if messageBox.clickedButton() == buttonOk:
                    self.all_files=[]
                    self.all_files.append(self.filePaths)
                    self.tableWidget.clearContents()#可以清除表格所有的内容
                    self.num = 0
                    for self.files in self.filePaths:
                        self.wb = openpyxl.load_workbook(self.files, read_only=False, data_only=True, keep_links=False)
                        t = time.localtime()  # 获取当前本地时间
                        strtime = time.strftime("%Y年-%m月-%d日-%H时-%M分-%S秒", t)
                        content = '.*/(.*).*.xlsx'
                        save_file = re.compile(content, re.S)
                        self.file = save_file.findall(self.files)
                        # self.wb.save(files_address+'\{}{}.xlsx'.format(self.file[0],strtime))
                        for self.ws in self.wb.worksheets:
                            self.rows = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                            max_num=len(self.rows)
                            for wb_ws_row in range(0, max_num):
                                data=self.rows[wb_ws_row]
                                self.num += 1
                                self.tableWidget.setRowCount(self.num)#设置行数
                                for colum in range(0,self.tableWidget.columnCount()):
                                    self.tableWidget.setItem(int(self.num)-1, colum, QTableWidgetItem(''))
                                for j in range(0, len(data)):
                                    self.item1 = QTableWidgetItem(str(data[j]).replace('None', ''))
                                    if j == 7:
                                        self.item1.setCheckState(Qt.Unchecked)
                                    self.tableWidget.setItem(int(self.num)-1, int(j), self.item1)
                    self.write_sys()
                    self.match_name()
                    self.tab_change = '提取'
                # 撤回
                if self.table_do == '执行':
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)


                # self.tableWidget.setStyleSheet("gridline-color: rgb(257, 1, 0)")
                # self.tableWidget.horizontalHeader().setStyleSheet(
                #     "color: rgb(0, 83, 128);border:1px solid rgb(210, 210, 210);")
                # self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive) #设置列宽，列宽可调
                # self.tableWidget.horizontalHeader().resizeSection(2, 400)  # 调整第2列的大小为500像素
                # self.tableWidget.horizontalHeader().resizeSection(7, 200)  # 调整第2列的大小为500像素
                # self.tableWidget.horizontalHeader().resizeSection(8, 300)  # 调整第11列的大小为500像素
                # tableHeader = self.tableWidget.horizontalHeader()
                # tableHeader.setStretchLastSection(False)
                # self.tableWidget.setColumnWidth(0, 220)  # 手动调整列宽
                # it = self.tableWidget.item(30, 0)
                # self.tableWidget.scrollToItem(it)#滚轮定位
    def row_Interactive(self):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.rows = self.tableWidget.rowCount()  # 获取单元格行数
            if self.rows != 0:
                for i in range(0, self.rows):
                    self.tableWidget.verticalHeader().setSectionResizeMode(i, QHeaderView.Interactive)
    all_filesqdk=[]
    files_qd =[]
    def Excel_qdk(self,result):#读取清单文件地址
        self.rows1 = []
        if result != '重新加载清单库':
            BASE_DIR = os.path.dirname(__file__)
            self.files_qdk = BASE_DIR + '/' + "清单数据库" + '/' + "清单库" + '/'
            self.file_qdk=QFileDialog.getExistingDirectory(self.window, "选择文件夹",self.files_qdk)
            if '/' in self.file_qdk:
                self.all_filesqdk = []
                self.files_qd = []
                self.files_qd.append(self.file_qdk)
                for dirpath, dirnames, filenames in os.walk(self.file_qdk):
                    for filename in filenames:
                        files = os.path.join(dirpath, filename)
                        if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                            self.all_filesqdk.append(str(files))
                            self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                            for self.ws in self.wb.worksheets:
                                self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,values_only=True))
                                for value in self.values:
                                    self.rows1.append(value)
                QMessageBox.about(self.window, '清单库', '清单库加载完成,可以编制清单！')
        if result == '重新加载清单库':
            if self.files_qd!=[]:
                # print(self.files_qd)
                for dirpath, dirnames, filenames in os.walk(self.files_qd[-1]):
                    for filename in filenames:
                        files = os.path.join(dirpath, filename)
                        if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                            self.all_filesqdk.append(str(files))
                            self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                            for self.ws in self.wb.worksheets:
                                self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,
                                                                     values_only=True))
                                for value in self.values:
                                    self.rows1.append(value)
        if self.files_qd!=[]:
            # QToolTip.setFont(QFont('SansSerif', 10))
            self.window.pushButton_qdk.setToolTip(str(self.files_qd[0]))
    all_filesdek=[]
    file_de =[]
    def startThread_run_import_dek(self):#加载定额库
        BASE_DIR = os.path.dirname(__file__)
        self.files_dek = BASE_DIR + '/' + "定额库" + '/'
        self.file_dek = QFileDialog.getExistingDirectory(self.window, "选择文件夹",self.files_dek)
        if '/' in self.file_dek:
            self.all_filesdek = []
            self.file_de = []
            self.file_de.append(self.file_dek)
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_import_dek)
            self.Mywork.stopSing.connect(self.stopThread_import_dek)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
        else:
            return
    def stopThread_import_dek(self):
        if self.quota_go=='继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
            QMessageBox.about(self.window, '定额库', '定额库加载完成,可以套定额了！')
        if self.quota_go == '不继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源

    def import_dek(self,quota):#加载定额库
        self.quota_values = []
        if quota == '加载定额库':
            for dirpath, dirnames, filenames in os.walk(self.file_dek):
                for filename in filenames:
                    files = os.path.join(dirpath, filename)
                    if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                        self.all_filesdek.append(str(files))
                        self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                        for self.ws in self.wb.worksheets:
                            self.values = list(
                                self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,
                                                  values_only=True))
                            for value in self.values:
                                self.quota_values.append(value)
            self.quota_go='继续执行'
        else:
            self.quota_go='不继续执行'

        if quota == '重新加载定额库':
            if self.file_de!=[]:
                # print(self.files_qd)
                for dirpath, dirnames, filenames in os.walk(self.file_de[-1]):
                    for filename in filenames:
                        files = os.path.join(dirpath, filename)
                        if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                            self.all_filesqdk.append(str(files))
                            self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                            for self.ws in self.wb.worksheets:
                                self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,
                                                                     values_only=True))
                                for value in self.values:
                                    self.quota_values.append(value)
        if self.file_de!=[]:
            self.window.toolButton_import_dek.setToolTip(str(self.file_de[0]))
    def startThread_run_hand_quota(self):#多线程手动套定额
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.all_filesdek != []:
                self.thread = QThread()  # 实例化一个线程
                self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
                self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
                self.thread.started.connect(self.Mywork.run_hand_quota)
                self.Mywork.stopSing.connect(self.stopThread_hand_quota)  # 停止信号连接到stopThread方法
                self.thread.start()  # 开始线程的运行
            else:
                QMessageBox.information(self.window, '提示', '请加载定额库，再匹配套定额')
                return

    def stopThread_hand_quota(self):
        if self.row_num == '继续执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
        elif self.row_num == '不执行':
            self.thread.quit()  # 退出
            self.thread.wait()  # 回收资源
    def hand_quota(self):#手动套定额
        if self.tableWidget.currentRow() != -1:
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row) != '-1':
                self.tab_change = '不提取'
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                # 选中行插入定额
                mytable = self.tableWidget.selectedItems()
                for r in mytable:#第三种方法获取值
                    if r.column()==5:
                        quantity_text = self.tableWidget.item(r.row(), 10).text()
                        sys_text = self.tableWidget.item(r.row(), 6).text()
                        name_text = self.tableWidget.item(r.row(), 7).text()
                        specification_text = self.tableWidget.item(r.row(), 8).text()
                        self.quota_text=self.tableWidget.item(r.row(), 5).text()#获取单元格内容
                        if self.quota_text!='《定额》':
                            self.quota_sp = str(self.quota_text).split('&')
                            n = 0
                            for num,text in enumerate(self.quota_sp):
                                for v in self.quota_values:
                                    if text in str(v[0]) and text!='':
                                        n+=1
                                        self.tableWidget.insertRow(r.row()+num+1)
                                        for column in range(0, self.tableWidget.columnCount()):
                                            items = QTableWidgetItem('')
                                            items.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                            brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                            brush.setStyle(QtCore.Qt.SolidPattern)
                                            items.setBackground(brush)  # 背景颜色
                                            self.tableWidget.setItem(r.row()+num+1, column, items)
                                        for j, data in enumerate(v):
                                            print(j, data)
                                            self.item1 = QTableWidgetItem(str(data).replace('None', ''))
                                            self.item1.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                            brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                            brush.setStyle(QtCore.Qt.SolidPattern)
                                            self.item1.setBackground(brush)  # 背景颜色
                                            if j == 2:
                                                self.tableWidget.setItem(r.row()+num+1, 0, self.item1)
                                            if j == 3:
                                                self.tableWidget.setItem(r.row()+num+1, 1, self.item1)
                                            if j == 4:
                                                self.tableWidget.setItem(r.row()+num+1, 3, self.item1)

                                        item5 = QTableWidgetItem('《定额》')
                                        item5.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                        brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                        brush.setStyle(QtCore.Qt.SolidPattern)
                                        item5.setBackground(brush)  # 背景颜色
                                        self.tableWidget.setItem(r.row()+num+1, 5, item5)
                                        self.tableWidget.setItem(r.row() + num + 1, 6,QTableWidgetItem(str(sys_text)))
                                        self.tableWidget.setItem(r.row() + num + 1, 7, QTableWidgetItem(str(name_text)))
                                        self.tableWidget.setItem(r.row() + num + 1, 8, QTableWidgetItem(str(specification_text)))
                                        self.tableWidget.setItem(r.row()+num+1, 10, QTableWidgetItem(str(quantity_text)))
                                        self.tableWidget.verticalHeader().resizeSection(r.row()+num+1, 80)
                                        self.tableWidget.verticalHeader().setSectionResizeMode(r.row()+num+1,QHeaderView.ResizeToContents)
                                        self.tableWidget.viewport().update()
                                        if n==1:
                                            break
                if self.table_do == '执行':
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'

                self.row_num='继续执行'
        else:
            self.row_num = '不执行'

    def specifications(self):#解析材料规格窗口
        self.specifications=Specifications_window()
        self.specifications.window.show()
    def hand_qdk(self):#手动编清单
        if self.all_filesqdk != [] and self.tableWidget.currentRow()!=-1:
            self.tab_change = '不提取'
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row)!='-1':
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)

                self.t=self.tableWidget.item(self.row, 4).text()#获取单元格内容
                self.fre=self.tableWidget.item(self.row, 7).text()#获取单元格内容
                self.eight = str(self.tableWidget.item(self.row, 8).text()).split('&')

                if len(self.eight)==1:
                    self.sre = str(self.tableWidget.item(self.row, 8).text()).split('&')[0]
                    for v in self.rows1:
                        if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                and 'NP3' not in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                            for j in range(0,len(v[1:5])):
                                self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                self.item1 = QTableWidgetItem(self.redata)
                                self.tableWidget.setItem(int(self.row), int(j), self.item1)
                if len(self.eight)==2:
                    self.sre = str(self.tableWidget.item(self.row, 8).text()).split('&')[0]
                    self.four = str(self.tableWidget.item(self.row, 8).text()).split('&')[-1]
                    for v in self.rows1:
                        if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                and  'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                            for j in range(0,len(v[1:5])):
                                self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four)
                                self.item1 = QTableWidgetItem(self.redata)
                                self.tableWidget.setItem(int(self.row), int(j), self.item1)
                if len(self.eight)==3:
                    self.sre = str(self.tableWidget.item(self.row, 8).text()).split('&')[0]
                    self.four = str(self.tableWidget.item(self.row, 8).text()).split('&')[1]
                    self.five = str(self.tableWidget.item(self.row, 8).text()).split('&')[2]
                    for v in self.rows1:
                        if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                and  'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' in str(v[3]) and 'NP5' not in str(v[3]):
                            for j in range(0,len(v[1:5])):
                                self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4',self.five)
                                self.item1 = QTableWidgetItem(self.redata)
                                self.tableWidget.setItem(int(self.row), int(j), self.item1)
                if len(self.eight)==4:
                    self.sre = str(self.tableWidget.item(self.row, 8).text()).split('&')[0]
                    self.four = str(self.tableWidget.item(self.row, 8).text()).split('&')[1]
                    self.five = str(self.tableWidget.item(self.row, 8).text()).split('&')[2]
                    self.six = str(self.tableWidget.item(self.row, 8).text()).split('&')[3]
                    for v in self.rows1:
                        if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                and 'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' in str(v[3]) and 'NP5' in str(v[3]):
                            for j in range(0,len(v[1:5])):
                                self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4',self.five).replace('NP5',self.six)
                                self.item1 = QTableWidgetItem(self.redata)
                                self.tableWidget.setItem(int(self.row), int(j), self.item1)
                # elif len(self.eight)==1:
                #     for v in self.rows1:
                #         if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                #                 and 'NP1' in str(v[3]) and 'NP2' not in str(v[3]) and 'NP3' not in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                #             for j in range(0,len(v[1:5])):
                #                 self.redata = str(v[1:5][j]).replace('NP1', str(self.fre))
                #                 self.item1 = QTableWidgetItem(self.redata)
                #                 self.tableWidget.setItem(int(self.row), int(j), self.item1)

                self.tableWidget.verticalHeader().resizeSection(self.row, 150)  # 调整第2列的大小为500像素
                self.tableWidget.verticalHeader().setSectionResizeMode(self.row, QHeaderView.ResizeToContents)
                self.tableWidget.update()

                if self.table_do == '执行':
                    self.new_undo_dict={}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        print(new_text_list)
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'

                self.row_num='继续执行'
        else:
            self.row_num = '不执行'

    def intelligent_qdk(self):# 匹配编清单
        # 撤销
        self.old_undo_dict = {}
        self.Click_list = []
        self.old_undo_dict[self.tableWidget] = self.Click_list
        for self.Single_rows in range(0, self.tableWidget.rowCount()):
            t_list = []
            self.Click_list.append(t_list)
            for self.Single_colum in range(0, self.tableWidget.columnCount()):
                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                t_list.append(self.text)
        # 匹配编清单
        try:
            row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(row) != '-1':
                self.tab_change = '不提取'
                self.rows = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if r.row() not in self.rows and self.tableWidget.item(r.row(), 5).text()!='《定额》':
                        self.rows.append(r.row())
                self.pwindow.progressBar.setRange(0, len(self.rows))
                for r in mytable:  # 第三种方法获取值
                    if r.column() == 4:
                        self.quota_text = self.tableWidget.item(r.row(), 5).text()  # 获取单元格内容
                        if self.quota_text != '《定额》':
                            i=r.row()
                            self.t = self.tableWidget.item(i, 4).text()  # 获取单元格内容
                            self.fre = self.tableWidget.item(i, 7).text()  # 获取单元格内容
                            self.eight = str(self.tableWidget.item(i, 8).text()).split('&')
                            if len(self.eight) == 1:
                                self.sre = str(self.tableWidget.item(i, 8).text()).split('&')[0]  # 获取单元格内容
                                for v in self.rows1:
                                    if str(self.t) in str(v[0]) and str(self.t) != '' and str(self.t) != ' ' and str(self.t) != '   ' and str(self.t) != '    '\
                                           and 'NP3' not in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                                        for j in range(0, len(v[1:5])):
                                            self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2',self.sre)
                                            self.tableWidget.setItem(int(i), int(j), QTableWidgetItem(str(self.redata)))
                                            self.tableWidget.verticalHeader().setSectionResizeMode(int(i),QHeaderView.ResizeToContents)

                            if len(self.eight)==2:
                                self.sre = str(self.tableWidget.item(i, 8).text()).split('&')[0]# 获取单元格内容
                                self.four=str(self.tableWidget.item(i, 8).text()).split('&')[-1]
                                for v in self.rows1:
                                    if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                            and 'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                                        for j in range(0, len(v[1:5])):
                                            self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four)
                                            self.tableWidget.setItem(int(i), int(j), QTableWidgetItem(str(self.redata)))
                                            self.tableWidget.verticalHeader().setSectionResizeMode(int(i),QHeaderView.ResizeToContents)
                            if len(self.eight) == 3:
                                self.sre = str(self.tableWidget.item(i, 8).text()).split('&')[0]# 获取单元格内容
                                self.four=str(self.tableWidget.item(i, 8).text()).split('&')[1]
                                self.five = str(self.tableWidget.item(i, 8).text()).split('&')[2]
                                for v in self.rows1:
                                    if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                            and 'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' in str(v[3]) and 'NP5' not in str(v[3]):
                                        for j in range(0, len(v[1:5])):
                                            self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4',self.five)
                                            self.tableWidget.setItem(int(i), int(j), QTableWidgetItem(str(self.redata)))
                                            self.tableWidget.verticalHeader().setSectionResizeMode(int(i),QHeaderView.ResizeToContents)
                            if len(self.eight) == 4:
                                self.sre = str(self.tableWidget.item(i, 8).text()).split('&')[0]# 获取单元格内容
                                self.four=str(self.tableWidget.item(i, 8).text()).split('&')[1]
                                self.five = str(self.tableWidget.item(i, 8).text()).split('&')[2]
                                self.six = str(self.tableWidget.item(i, 8).text()).split('&')[3]
                                for v in self.rows1:
                                    if str(self.t) in str(v[0]) and str(self.t)!='' and str(self.t)!=' ' and str(self.t)!='   ' and str(self.t)!='    '\
                                            and 'NP2' in str(v[3]) and 'NP3' in str(v[3]) and 'NP4' in str(v[3]) and 'NP5' in str(v[3]):
                                        for j in range(0, len(v[1:5])):
                                            self.redata = str(v[1:5][j]).replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4',self.five).replace('NP5',self.six)
                                            self.tableWidget.setItem(int(i), int(j), QTableWidgetItem(str(self.redata)))
                                            self.tableWidget.verticalHeader().setSectionResizeMode(int(i),QHeaderView.ResizeToContents)
                            # elif len(self.eight) == 1:
                            #     for v in self.rows1:
                            #         if str(self.t) in str(v[0]) and str(self.t) != '' and str(self.t) != ' ' and str(self.t) != '   ' and str(self.t) != '    '\
                            #                 and 'NP1' in str(v[3]) and 'NP2' not in str(v[3]) and 'NP3' not in str(v[3]) and 'NP4' not in str(v[3]) and 'NP5' not in str(v[3]):
                            #             for j in range(0, len(v[1:5])):
                            #                 self.redata = str(v[1:5][j]).replace('NP1', str(self.fre))
                            #                 self.tableWidget.setItem(int(i), int(j), QTableWidgetItem(str(self.redata)))
                            #                 self.tableWidget.verticalHeader().setSectionResizeMode(int(i),QHeaderView.ResizeToContents)
                            self.Mywork.runSing.emit(i)
                self.tableWidget.viewport().update()
        except Exception as e:  # Exception捕获错误的类型，e保存具体错误内容
            print('出现异常', e)
            self.intell = '不执行'
        else:
            self.intell = '继续执行'
            if self.table_do == '执行':
                self.new_undo_dict = {}
                new_text_list = []
                self.new_undo_dict[self.tableWidget] = new_text_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    new_text_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.Click_list != new_text_list:
                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                    self.undoStack_del.push(command)
                self.tab_change = '提取'


    def write_old_Excel(self):#保存到原Excel表格
        self.count=0
        self.table_rows = self.tableWidget.rowCount()
        for self.old_file in self.old_filePaths:
            self.wb = openpyxl.load_workbook(self.old_file, read_only=False, data_only=True, keep_links=False)
            for self.ws in self.wb.worksheets:
                self.rows = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                print(self.rows)
                self.max_num = len(self.rows)
                self.pwindow.progressBar.setRange(0, int(self.max_num))
                for wb_ws_row in range(0, self.max_num):
                    for table_row in range(0, self.table_rows):
                        table_values0 = self.tableWidget.item(int(table_row), 0).text()
                        table_values1 = self.tableWidget.item(int(table_row), 1).text()
                        table_values2 = self.tableWidget.item(int(table_row), 2).text()
                        table_values3 = self.tableWidget.item(int(table_row), 3).text()
                        table_values4 = self.tableWidget.item(int(table_row), 4).text()
                        table_values5 = self.tableWidget.item(int(table_row), 5).text()
                        table_values6 = self.tableWidget.item(int(table_row), 6).text()
                        table_values7 = self.tableWidget.item(int(table_row), 7).text()
                        table_values8 = self.tableWidget.item(int(table_row), 8).text()
                        table_values9 = self.tableWidget.item(int(table_row), 9).text()
                        all_valus= table_values6,table_values7,table_values8,table_values9
                        if len(self.rows[wb_ws_row])>=10:
                            if str(self.rows[wb_ws_row][6]).replace('None','')==table_values6\
                                    and str(self.rows[wb_ws_row][7]).replace('None','')==table_values7\
                                    and str(self.rows[wb_ws_row][8]).replace('None','')==table_values8\
                                    and str(self.rows[wb_ws_row][9]).replace('None','')==table_values9:
                                self.count+=1
                                # label_value='第{}条清单匹配完成'.format(str(self.count))
                                self.ws.cell(wb_ws_row + 1,1, table_values0)
                                self.ws.cell(wb_ws_row + 1,2, table_values1)
                                self.ws.cell(wb_ws_row + 1,3, table_values2)
                                self.ws.cell(wb_ws_row + 1,4, table_values3)
                                self.ws.cell(wb_ws_row + 1,5, table_values4)
                                self.ws.cell(wb_ws_row + 1,6, table_values5)
                                self.ws.cell(wb_ws_row + 1,7, table_values6)
                                self.ws.cell(wb_ws_row + 1,8, table_values7)
                                self.ws.cell(wb_ws_row + 1,9, table_values8)
                                self.ws.cell(wb_ws_row + 1,10, table_values9)
                                self.worker.runSing.emit(int(wb_ws_row))
            hide_f = '(.*)/(.*)'
            self.hide = re.compile(hide_f, re.S)
            self.t_hied = self.hide.findall(self.old_file)
            hidefilenames = self.t_hied[0][0] + '/' + '~$' + self.t_hied[0][1]
            if os.path.exists(hidefilenames):
                self.result = '文件打开了'#防止在写入时打开文件，软件崩溃。
            else:
                self.wb.save(self.old_file)
                self.result = '执行完毕'


    def write_new_Excel(self):
        items_dict={}
        ite = self.window.treeWidget_Items.topLevelItem(0)  # 循环获取根节点
        count = ite.childCount()  # 获取当前根节点的子节点数量
        for j in range(0, count):
            string = ite.child(j)  # 子节点的文字信息
            if string.checkState(0) == Qt.Checked:
                print(string.text(1))
                items_dict[(str(string.text(1)), str(string.text(0)))] = self.tablelist[int(string.text(1))]
        if items_dict != {}:
            # files_address=QFileDialog.getExistingDirectory(self.window, "选择保存文件夹")
            files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存","默认名称", '*.xlsx')
            criteria = '(.*)/'
            compile = re.compile(criteria, re.S)
            results = compile.findall(files_address)
            if files_address == "":
                self.result = '取消'
            # if '/' in files_address:
            if files_address != "":
                for key,self.tab in items_dict.items():
                    self.tableWidget_allrows = int(self.tab.rowCount())  # 获取总行数
                    self.table_column = int(self.tab.columnCount())
                    if str(self.tab.rowCount()) != '0':
                        file_path = os.path.join(results[0],'~$'+key[0]+key[1]+'.xlsx')
                        if os.path.exists(file_path):
                            QMessageBox.information(self.window, '温馨提示','无法保存到Excel，请先关闭或删除同名的Excel文件')
                            self.result = '文件打开了'
                            return
                        else:
                            nw = openpyxl.Workbook()  # 新建文件
                            for self.Single_rows in range(0, self.tableWidget_allrows):
                                for self.Single_colum in range(0,self.table_column):
                                    self.text= self.tab.item(self.Single_rows, self.Single_colum).text()
                                    nw.active.cell(self.Single_rows+1,self.Single_colum+1,str(self.text)).alignment = self.alig
                                    nw.active.cell(self.Single_rows + 1, self.Single_colum + 1,str(self.text)).font = self.font
                                    nw.active.cell(self.Single_rows + 1, self.Single_colum + 1,str(self.text)).border =self.border

                            if os.path.exists(file_path):
                                self.result = '文件打开了'
                            else:
                                nw.save(os.path.join(results[0], key[0]+key[1]+'.xlsx')) # 保存新建的文件
                                self.result = '执行完毕'

                    else:
                        continue
                QMessageBox.information(self.window, '温馨提示', '保存到Excel完成，请检查。')
            else:
                self.result = '取消'
        else:
            self.result = '取消'

            # fileName_choose, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", self.cwd, "*.xlsx")
            # if fileName_choose == "":
            #     print("\n取消选择")
            #     return
            # if fileName_choose!="":
            #     nw = openpyxl.Workbook()  # 新建文件
            #     nw.save(fileName_choose)  # 保存新建的文件
            #     self.new_wb = openpyxl.load_workbook(fileName_choose, read_only=False, data_only=True, keep_links=False)
            #     self.new_ws = self.new_wb.active
            # self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            # self.table_column = int(self.tableWidget.columnCount())
            # print(self.tabldict)
            # # if self.tabldict != {}:
            # for self.tab in self.tabldict.values():
            #     if str(self.tab.rowCount()) != '0':
            #         for self.Single_rows in range(0, self.tableWidget_allrows):
            #             for self.Single_colum in range(0,self.table_column):
            #                 self.text= self.tableWidget.item(self.Single_rows, self.Single_colum).text()
            #                 self.new_ws.cell(self.Single_rows+1,self.Single_colum+1,str(self.text))
            # self.bin_path = os.path.join(files_address, self.fist.text(0), self.second.text(0) + '.xlsx')
            # self.new_wb.save(fileName_choose)  # 保存新建的文件

    def write_zero_Excel(self):
        if int(self.window.tableWidget_0.rowCount())!=0:
            files_address = QFileDialog.getExistingDirectory(self.window, "选择保存文件夹")
            if '/' in files_address:
                self.tableWidget_allrows = int(self.window.tableWidget_0.rowCount())  # 获取总行数
                self.table_column = int(self.window.tableWidget_0.columnCount())
                file_path = os.path.join(files_address,'~$'+'标准清单.xlsx')
                if os.path.exists(file_path):
                    QMessageBox.information(self.window, '温馨提示', '无法保存到Excel，请先关闭或删除同名的Excel文件')
                    return
                else:
                    nw = openpyxl.Workbook()  # 新建文件
                    for self.Single_rows in range(0, self.tableWidget_allrows):
                        for self.Single_colum in range(0, self.table_column):
                            self.text = self.window.tableWidget_0.item(self.Single_rows, self.Single_colum).text()
                            nw.active.cell(self.Single_rows + 1, self.Single_colum + 1,
                                           str(self.text)).alignment = self.alig
                            nw.active.cell(self.Single_rows + 1, self.Single_colum + 1,
                                           str(self.text)).font = self.font
                            nw.active.cell(self.Single_rows + 1, self.Single_colum + 1,
                                           str(self.text)).border = self.border

                    if os.path.exists(file_path):
                        QMessageBox.information(self.window, '温馨提示','无法保存到Excel，请先关闭或删除同名的Excel文件')
                        return
                    else:
                        nw.save(os.path.join(files_address,'标准清单.xlsx'))  # 保存新建的文件
                        QMessageBox.information(self.window, '温馨提示', '保存到Excel完成，请检查。')
        else:
            return

    label_value=''
    def all_qdk(self): # 复用标准清单
        self.tab_change = '不提取'
        self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
        self.table_column = int(self.tableWidget.columnCount())
        if self.tableWidget_allrows != '0':
            # 撤销
            self.old_undo_dict = {}
            if self.tabldict != {}:
                for self.tab in self.tabldict.values():
                    # 撤销
                    self.Click_list = []
                    self.old_undo_dict[self.tab] = self.Click_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        self.Click_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            self.pwindow.progressBar.setRange(0,self.tableWidget_allrows)
            num = 0
            for self.Single_rows in range(0, self.tableWidget_allrows):
                Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Checked:
                    text_list = []
                    for self.Single_colum in range(0, self.table_column):
                        check_copytext = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(check_copytext)
                    # print(text_list[6:10])

                    if self.tabldict != {}:
                        for self.tab in self.tabldict.values():
                            # print(self.tab.rowCount())
                            if str(self.tab.rowCount()) != '0':
                                for self.other_rows in range(0, int(self.tab.rowCount())):
                                    self.others_quota = self.tab.item(self.other_rows, 5).text()
                                    self.others_system = self.tab.item(self.other_rows, 6).text()
                                    self.others_name = self.tab.item(self.other_rows, 7).text()
                                    self.others_specification = self.tab.item(self.other_rows, 8).text()
                                    self.others_unit = self.tab.item(self.other_rows, 9).text()
                                    # print([self.others_name,self.others_specification,self.others_unit])
                                    if self.others_system in text_list[6] and [self.others_name,self.others_specification,self.others_unit] == text_list[7:10] and self.others_quota!='《定额》':
                                        self.tab.setItem(self.other_rows, 0, QTableWidgetItem(str(text_list[0])))
                                        self.tab.setItem(self.other_rows, 1, QTableWidgetItem(str(text_list[1])))
                                        self.tab.setItem(self.other_rows, 2, QTableWidgetItem(str(text_list[2])))
                                        self.tab.setItem(self.other_rows, 3, QTableWidgetItem(str(text_list[3])))
                                        self.tab.setItem(self.other_rows, 4, QTableWidgetItem(str(text_list[4])))
                                        self.tab.setItem(self.other_rows, 5, QTableWidgetItem(str(text_list[5])))
                                        self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                        global label_value
                                        label_value = '第{}条清单复用完成'.format(str(self.other_rows))
                                        num+=1
                                        self.Mywork.runSing.emit(int(self.other_rows))
            # 回撤
            self.new_undo_dict = {}
            if self.tabldict != {}:
                for self.tab in self.tabldict.values():
                    new_text_list = []
                    self.new_undo_dict[self.tab] = new_text_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            if self.old_undo_dict != self.new_undo_dict:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                self.undoStack_del.push(command)
            self.tab_change = '提取'
            self.ns = '继续执行'
        else:
            self.ns = '不执行'

    choose_sys_rows=[]
    def select_allname(self):#勾选名称
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tab_change = '不提取'
            if self.window.treeWidget_Items.currentItem():
                if self.tableWidget.rowCount() != 0:
                    mytable = self.tableWidget.selectedItems()
                    self.column_dict = {}
                    for r in mytable:  # 第三种方法获取值
                        if r.column()==7:
                            print(r.row())
                            text = self.tableWidget.item(r.row(), 5).text()
                            if text!= '《定额》':
                                self.tableWidget.item(r.row(), 7).setCheckState(Qt.Checked)

                    # if self.name_row_list != []:# 按选定的材料匹配行
                    #     for self.Single_rows in self.name_row_list:  # 按表格中选中显示的行  self.tableWidget.rowCount()
                    #         self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Checked)
                    # if self.name_row_list== []:# 按系统匹配行
                    #     if self.sys_0 == '整个工程':
                    #         for self.Single_rows in range(0,self.tableWidget.rowCount()): #按表格中选中显示的行  self.tableWidget.rowCount()
                    #             text = self.tableWidget.item(self.Single_rows, 5).text()
                    #             if text!= '《定额》':
                    #                 self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Checked)
                    #     if self.sys_0 != '整个工程':
                    #         if self.choose_sys_rows!=[]:
                    #             for self.Single_rows in self.choose_sys_rows: #按表格中选中显示的行  self.tableWidget.rowCount()
                    #                 text = self.tableWidget.item(self.Single_rows, 5).text()
                    #                 if text != '《定额》':
                    #                     self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Checked)

                        # if self.window.treeWidget_system.currentIndex().row()== -1:#点击单位工程全选材料
                        #     self.table_rows = self.tableWidget.rowCount()
                        #     self.sys_row = self.window.treeWidget_system.currentIndex().row()  # 获取行
                        #     for self.Single_rows in range(0, self.table_rows):
                        #         self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Checked)
                self.tab_change = '提取'

    def unselect_allname(self):#取消全勾选
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tab_change = '不提取'
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)!=0:
            if self.tableWidget.rowCount()!=0:
                mytable = self.tableWidget.selectedItems()
                self.column_dict = {}
                for r in mytable:  # 第三种方法获取值
                    if r.column() == 7:
                        print(r.row())
                        text = self.tableWidget.item(r.row(), 5).text()
                        if text != '《定额》':
                            self.tableWidget.item(r.row(), 7).setCheckState(Qt.Unchecked)

                # for self.Single_rows in range(0, self.tableWidget.rowCount()):
                #     text = self.tableWidget.item(self.Single_rows, 5).text()
                #     if text != '《定额》':
                #         self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                self.tab_change = '提取'

    def copy_names(self):# 复用匹配名称
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.tableWidget.currentRow()!=-1:
                self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                self.row = self.tableWidget.currentRow()  # 获取单元格行数
                if str(self.row)!='-1':
                    items = ["同名称—单位", "同名称—规格—单位", "同系统—名称—规格—单位"]
                    com_text, ok = QInputDialog().getItem(self.window, "选择应用范围", "应用其它单位工程需提前打√", items, 0, True)
                    if ok:
                        if com_text == '同名称—单位':
                            self.startThread_run_copy_name()
                        if com_text == '同名称—规格—单位':
                            self.startThread_run_name_sp()
                        if com_text == '同系统—名称—规格—单位':
                            self.startThread_run_check_name()

    def copy_name(self):#复用同名称
        self.tab_change = '不提取'
        Checken_dict = {}
        if self.tableWidget.currentRow()!=-1:
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row)!='-1':
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget]=self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)

                if self.tabldict != {}:
                    for self.tab in self.tabldict.values():
                        # 撤销
                        self.Click_list = []
                        self.old_undo_dict[self.tab] = self.Click_list
                        for self.Single_rows in range(0, self.tab.rowCount()):
                            t_list = []
                            self.Click_list.append(t_list)
                            for self.Single_colum in range(0, self.tab.columnCount()):
                                self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)

                # 复用清单名称
                self.column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    self.column.append(r.column())
                if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                        and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                    copytext4=self.tableWidget.item(self.row, 4).text()#获取单元格内容
                    copytext5 = self.tableWidget.item(self.row, 5).text()  # 获取单元格内容
                    for self.Single_rows in range(0,int(self.tableWidget.rowCount())):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1 == Qt.Checked:
                            self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(copytext4)))
                            # self.tableWidget.setItem(self.Single_rows, 5, QTableWidgetItem(str(copytext5)))
                            check_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            check_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            check_copytext4 = self.tableWidget.item(self.Single_rows, 4).text()
                            check_copytext5 = self.tableWidget.item(self.Single_rows, 5).text()
                            Checken_dict[check_allname,check_unit] = [check_copytext4,check_copytext5]
                            self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                    for self.Single_rows in range(0, int(self.tableWidget.rowCount())):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1 == Qt.Unchecked:
                            uncheck_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            uncheck_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            merge = (uncheck_allname,uncheck_unit)
                            if merge in Checken_dict:
                                self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(Checken_dict[merge][0])))
                                # self.tableWidget.setItem(self.Single_rows, 5,QTableWidgetItem(str(Checken_dict[merge][1])))


                    if self.tabldict!= {}:
                        for self.tab in self.tabldict.values():
                            # 写入清单名称
                            if str(self.tab.rowCount()) != '0':
                                for self.other_rows in range(0,int(self.tab.rowCount())):
                                    self.others_name = self.tab.item(self.other_rows, 7).text()
                                    self.others_unit = self.tab.item(self.other_rows, 9).text()
                                    others_merge = (self.others_name,self.others_unit)
                                    if others_merge in Checken_dict:
                                        self.tab.setItem(self.other_rows, 4, QTableWidgetItem(str(Checken_dict[others_merge][0])))
                                        # self.tab.setItem(self.other_rows, 5,QTableWidgetItem(str(Checken_dict[others_merge][1])))

                    # 回撤
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.tabldict!= {}:
                        for self.tab in self.tabldict.values():
                            new_text_list = []
                            self.new_undo_dict[self.tab] = new_text_list
                            for self.Single_rows in range(0, self.tab.rowCount()):
                                t_list = []
                                new_text_list.append(t_list)
                                for self.Single_colum in range(0, self.tab.columnCount()):
                                    self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                    t_list.append(self.text)
                    if self.old_undo_dict != self.new_undo_dict:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)

                self.tableWidget.viewport().update()
                self.tab_change = '提取'
                self.cn='继续执行'

            else:
                self.cn = '不执行'
        else:
            self.cn = '不执行'
    def name_sp(self):#复用同名称规格
        Checkens_dict = {}
        if self.tableWidget.currentRow()!=-1:
            self.tab_change = '不提取'
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row)!='-1':
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget]=self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.tabldict != {}:
                    for self.tab in self.tabldict.values():
                        # 撤销
                        self.Click_list = []
                        self.old_undo_dict[self.tab] = self.Click_list
                        for self.Single_rows in range(0, self.tab.rowCount()):
                            t_list = []
                            self.Click_list.append(t_list)
                            for self.Single_colum in range(0, self.tab.columnCount()):
                                self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)
                # 写入清单名称
                self.column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    self.column.append(r.column())
                if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                        and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                    self.copytext0 = self.tableWidget.item(self.row, 0).text()  # 获取单元格内容
                    self.copytext1 = self.tableWidget.item(self.row, 1).text()  # 获取单元格内容
                    self.copytext2 = self.tableWidget.item(self.row, 2).text()  # 获取单元格内容
                    self.copytext3 = self.tableWidget.item(self.row, 3).text()  # 获取单元格内容
                    self.copytext4=self.tableWidget.item(self.row, 4).text()#获取单元格内容
                    self.copytext5 = self.tableWidget.item(self.row, 5).text()  # 获取单元格内容
                    self.name = self.tableWidget.item(self.row, 7).text()  # 获取单元格内容
                    self.specification = self.tableWidget.item(self.row, 8).text()  # 获取单元格内容
                    for self.Single_rows in range(0,self.tableWidget_allrows):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1 == Qt.Checked:
                            self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(self.copytext4)))
                            # self.tableWidget.setItem(self.Single_rows, 5, QTableWidgetItem(str(self.copytext5)))
                            check_copytext4 = self.tableWidget.item(self.Single_rows, 4).text()
                            check_copytext5 = self.tableWidget.item(self.Single_rows, 5).text()
                            check_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            check_allspecification = self.tableWidget.item(self.Single_rows, 8).text()
                            check_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            Checkens_dict[check_allname,check_allspecification,check_unit]=[check_copytext4,check_copytext5]
                            self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)

                    for self.Single_rows in range(0, self.tableWidget_allrows):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1 == Qt.Unchecked:
                            uncheck_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            uncheck_allspecification = self.tableWidget.item(self.Single_rows, 8).text()
                            uncheck_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            merge = (uncheck_allname, uncheck_allspecification,uncheck_unit)
                            if merge in Checkens_dict:
                                self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(Checkens_dict[merge][0])))
                                # self.tableWidget.setItem(self.Single_rows, 5,QTableWidgetItem(str(Checkens_dict[merge][1])))


                    if self.tabldict != {}:
                        for self.tab in self.tabldict.values():
                            # print(self.tab.rowCount())
                            if str(self.tab.rowCount()) != '0':
                                for self.other_rows in range(0,int(self.tab.rowCount())):
                                    self.others_name = self.tab.item(self.other_rows, 7).text()
                                    self.others_specification = self.tab.item(self.other_rows, 8).text()
                                    self.others_unit = self.tab.item(self.other_rows, 9).text()
                                    others_merge = (self.others_name, self.others_specification,self.others_unit)
                                    if others_merge in Checkens_dict:
                                        self.tab.setItem(self.other_rows, 4, QTableWidgetItem(str(Checkens_dict[others_merge][0])))
                                        # self.tab.setItem(self.other_rows, 5,QTableWidgetItem(str(Checkens_dict[others_merge][1])))


                    # 回撤
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.tabldict!= {}:
                        for self.tab in self.tabldict.values():
                            new_text_list = []
                            self.new_undo_dict[self.tab] = new_text_list
                            for self.Single_rows in range(0, self.tab.rowCount()):
                                t_list = []
                                new_text_list.append(t_list)
                                for self.Single_colum in range(0, self.tab.columnCount()):
                                    self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                    t_list.append(self.text)
                    if self.old_undo_dict != self.new_undo_dict:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                    self.tableWidget.viewport().update()
                self.tab_change = '提取'
                self.ns='继续执行'
            else:
                self.ns = '不执行'
        else:
            self.ns = '不执行'

    def check_sps_name(self):# 复用同系统同名称同规
        Checkeddict= {}
        if self.tableWidget.currentRow()!=-1:
            self.tab_change = '不提取'
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if str(self.row)!='-1':
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget]=self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.tabldict != {}:
                    for self.tab in self.tabldict.values():
                        # 撤销
                        self.Click_list = []
                        self.old_undo_dict[self.tab] = self.Click_list
                        for self.Single_rows in range(0, self.tab.rowCount()):
                            t_list = []
                            self.Click_list.append(t_list)
                            for self.Single_colum in range(0, self.tab.columnCount()):
                                self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)
                # 写入清单名称
                self.column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    self.column.append(r.column())
                if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                        and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                    copytext4 = self.tableWidget.item(self.row, 4).text()  # 获取单元格内容
                    copytext5 = self.tableWidget.item(self.row, 5).text()  # 获取单元格内容
                    for self.Single_rows in range(0,self.tableWidget_allrows):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1==Qt.Checked:
                            self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(copytext4)))
                            # self.tableWidget.setItem(self.Single_rows, 5, QTableWidgetItem(str(copytext5)))
                            check_copytext4 = self.tableWidget.item(self.Single_rows, 4).text()
                            check_copytext5 = self.tableWidget.item(self.Single_rows, 5).text()
                            check_system = self.tableWidget.item(self.Single_rows, 6).text()
                            check_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            check_allspecification = self.tableWidget.item(self.Single_rows, 8).text()
                            check_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            Checkeddict[check_system,check_allspecification,check_unit,check_allname]=[check_copytext4,check_copytext5]#用加号是为了可以在合并系统后用
                            self.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                    for self.Single_rows in range(0, self.tableWidget_allrows):
                        Item1 = self.tableWidget.item(self.Single_rows, 7).checkState()
                        if Item1==Qt.Unchecked:
                            uncheck_system = self.tableWidget.item(self.Single_rows, 6).text()
                            uncheck_allname = self.tableWidget.item(self.Single_rows, 7).text()
                            uncheck_allspecification = self.tableWidget.item(self.Single_rows, 8).text()
                            uncheck_unit = self.tableWidget.item(self.Single_rows, 9).text()
                            # print((uncheck_allspecification,uncheck_unit,uncheck_allname))
                            for k,v in Checkeddict.items():
                                if uncheck_system in k[0] and (uncheck_allspecification,uncheck_unit,uncheck_allname) == k[1:]:
                                    self.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(str(v[0])))
                                    # self.tableWidget.setItem(self.Single_rows, 5, QTableWidgetItem(str(v[1])))
                    if self.tabldict != {}:
                        for self.tab in self.tabldict.values():
                            # print(self.tab.rowCount())
                            if str(self.tab.rowCount()) != '0':
                                for self.other_rows in range(0,int(self.tab.rowCount())):
                                    self.others_system = self.tab.item(self.other_rows, 6).text()
                                    self.others_name = self.tab.item(self.other_rows, 7).text()
                                    self.others_specification = self.tab.item(self.other_rows, 8).text()
                                    self.others_unit = self.tab.item(self.other_rows, 9).text()
                                    # others_merge = (self.others_system, self.others_name, self.others_specification)
                                    for k, v in Checkeddict.items():
                                        if self.others_system in k[0] and (self.others_specification,self.others_unit,self.others_name) == k[1:]:
                                            self.tab.setItem(self.other_rows, 4, QTableWidgetItem(str(v[0])))
                                            # self.tab.setItem(self.other_rows, 5, QTableWidgetItem(str(v[1])))

                    # 回撤
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.tabldict!= {}:
                        for self.tab in self.tabldict.values():
                            new_text_list = []
                            self.new_undo_dict[self.tab] = new_text_list
                            for self.Single_rows in range(0, self.tab.rowCount()):
                                t_list = []
                                new_text_list.append(t_list)
                                for self.Single_colum in range(0, self.tab.columnCount()):
                                    self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                                    t_list.append(self.text)
                    if self.old_undo_dict != self.new_undo_dict:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                    self.tableWidget.viewport().update()
                self.tab_change = '提取'
                self.ns='继续执行'
            else:
                self.ns = '不执行'
        else:
            self.ns = '不执行'

    # 多线程合并单位工程
    def startThread_run_merge_unit(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_merge_unit)
        self.Mywork.stopSing.connect(self.stopThread_merge_unit)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行

    def stopThread_merge_unit(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源

    def merge_unit(self):#合并工程除了备注0到10列一样才能合并
        system_dict = {}
        if self.tabldict != {}:
            self.tab_change = '不提取'
            self.all_rows = 0
            for self.tab in self.tabldict.values():
                if str(self.tab.rowCount()) != '0':
                    for self.all_row in range(0,int(self.tab.rowCount())):
                        self.all_rows += 1
                        system = []
                        for self.all_colums in range(0,int(self.tab.columnCount())):
                            self.all_values = self.tab.item(self.all_row, self.all_colums).text()
                            system.append(self.all_values)
                        if tuple(system[0:7]+system[7:10]) not in system_dict and '《定额》' not in system:
                            system_dict[tuple(system[0:7]+system[7:10])] = [system[10]]
                        elif tuple(system[0:7]+system[7:10]) in system_dict and '《定额》' not in system:
                            system_dict[tuple(system[0:7]+system[7:10])].append(system[10])
            if system_dict!={}:
                self.window.tableWidget_0.clearContents()# 清空表格
                num=0
                rows=len(system_dict.items())
                self.window.tableWidget_0.setRowCount(rows)  # 设置行数
                for r in range(0,rows):
                    for c in range(0,self.window.tableWidget_0.columnCount()):
                        self.window.tableWidget_0.setItem(r, c ,QTableWidgetItem(''))
                for items,values in system_dict.items():
                    num += 1
                    if '《定额》' not in items:
                        self.window.tableWidget_0.setItem(num-1, 0, QTableWidgetItem(str(items[0])))
                        self.window.tableWidget_0.setItem(num-1, 1, QTableWidgetItem(str(items[1])))
                        self.window.tableWidget_0.setItem(num-1, 2, QTableWidgetItem(str(items[2])))
                        self.window.tableWidget_0.setItem(num-1, 3, QTableWidgetItem(str(items[3])))
                        self.window.tableWidget_0.setItem(num-1, 4, QTableWidgetItem(str(items[4])))
                        self.window.tableWidget_0.setItem(num-1, 5, QTableWidgetItem(str(items[5])))
                        self.window.tableWidget_0.setItem(num-1,6, QTableWidgetItem(str(items[6])))
                        self.item1 = QTableWidgetItem(str(items[7]))
                        self.item1.setCheckState(Qt.Unchecked)
                        self.window.tableWidget_0.setItem(num-1, 7, QTableWidgetItem(self.item1))
                        self.window.tableWidget_0.setItem(num-1, 8, QTableWidgetItem(str(items[8])))
                        self.window.tableWidget_0.setItem(num-1, 9, QTableWidgetItem(str(items[9])))
                        self.window.tableWidget_0.setItem(num - 1, 10, QTableWidgetItem('+'.join(values)))
                        self.window.tableWidget_0.setItem(num-1, 11, QTableWidgetItem(''))
                if self.tableWidget==self.window.tableWidget_0:
                    self.write_sys()
                    self.match_name()
                # 合并工程将单位工程不选中
                n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
                for i in range(0, n):
                    ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    for j in range(0, count):
                        strin = ite.child(j)  # 子节点的文字信息
                        if strin != None:
                            strin.setCheckState(0, Qt.Unchecked)
                # self.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
            self.tabldict = {}
            self.tab_change = '提取'
    # 多线程重读系统
    def startThread_run_write_sys(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_write_sys)
        self.Mywork.stopSing.connect(self.stopThread_write_sys)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行

    def stopThread_write_sys(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
    def write_sys(self):#重读系统和材料名称到树界面
        self.tab_change = '不提取'
        font = QFont()
        font.setPointSize(10)  # 设置字体大小为10像素
        font.setFamily("宋体")
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        self.window.treeWidget_system.clear()#清空
        self.window.treeWidget_name.clear()
        self.window.treeWidget_specification.clear()
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)>=0:
            system_dict = {}
            names_dict={}
            specification_dict={}
            self.table_rows = self.tableWidget.rowCount()
            if self.table_rows != '0':
                for self.all_row in range(0,self.table_rows):
                    self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                    self.all_name = self.tableWidget.item(self.all_row, 7).text()
                    self.all_specification = self.tableWidget.item(self.all_row, 8).text()
                    if self.all_sys not in system_dict:
                        system_dict[self.all_sys]=' '
                    if self.all_name not in names_dict:
                        names_dict[self.all_name]=' '
                    if self.all_specification not in specification_dict:
                        specification_dict[self.all_specification]=' '
            if system_dict != {}:
                self.root = QTreeWidgetItem(self.window.treeWidget_system)
                self.root.setText(0, '整个工程')
                self.root.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.root.setText(1, '0')
                self.root.setCheckState(0, Qt.Unchecked)

                for sys in system_dict.keys():
                    self.son = QTreeWidgetItem(self.root)
                    # self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                    self.son.setText(0, sys)
                    self.son.setSizeHint(0, QSize(0, 25))
                    self.son.setCheckState(0, Qt.Unchecked)
                    self.son.setFont(0, font)

            if names_dict != {}:
                self.r_name = QTreeWidgetItem(self.window.treeWidget_name)
                self.r_name.setText(0, '全部名称')
                self.r_name.setFlags(
                    QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                self.r_name.setText(1, '0')
                self.r_name.setCheckState(0, Qt.Unchecked)
                for name in names_dict.keys():
                    self.grandson = QTreeWidgetItem(self.r_name)
                    self.grandson.setText(0, name)
                    self.grandson.setSizeHint(0, QSize(0, 25))
                    self.grandson.setCheckState(0, Qt.Unchecked)
                    self.grandson.setFont(0, font)
                    self.grandson.setFlags(
                        QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
            if specification_dict != {}:
                self.r_specification = QTreeWidgetItem(self.window.treeWidget_specification)
                self.r_specification.setText(0, '全部规格')
                self.r_specification.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                self.r_specification.setText(1, '0')
                self.r_specification.setCheckState(0, Qt.Unchecked)
                for specifications in specification_dict.keys():
                    self.son = QTreeWidgetItem(self.r_specification)
                    self.son.setText(0, str(specifications))
                    self.son.setSizeHint(0, QSize(0, 25))
                    self.son.setCheckState(0, Qt.Unchecked)
                    self.son.setFont(0, font)
                    self.son.setFlags(
                        QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            self.window.treeWidget_system.expandAll()
            self.window.treeWidget_name.expandAll()
            self.window.treeWidget_specification.expandAll()
            self.tab_change = '提取'

    sys_0=''
    def system_Clicked(self):  # 点击系统获取对应的材料和筛选行
        self.tab_change = '不提取'
        self.choose_sys_rows = []#全选行函数系统行清零
        self.name_row_list = []#名称全选行函数材料行清零
        self.check_names_dict={}
        self.check_specification_dict={}
        # self.tab_change = '不提取'
        font = QFont()
        font.setPointSize(10)  # 设置字体大小为10像素
        font.setFamily("宋体")
        name_dict = {}
        specifications_dict={}
        manifest_dict = {}
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            self.sys_row = self.window.treeWidget_system.currentIndex().row()  # 获取行
            if int(self.sys_row) != -1:
                self.sys = self.window.treeWidget_system.currentItem()
                self.sys_0 = self.sys.text(0)  # 获取0列内容
                self.window.treeWidget_system.headerItem().setText(0, self.sys_0)  # 表头写入内容
                self.table_rows = self.tableWidget.rowCount()
                if self.table_rows != '0':
                    self.window.treeWidget_name.clear()
                    self.window.treeWidget_specification.clear()
                    self.window.listWidget_sys.clear()  # 清空列表

                    if self.sys_0 == '整个工程':
                        self.startThread_run_show_rows()
                        self.write_sys()
                    elif self.sys_0 != '整个工程':
                        for self.all_row in range(0, self.table_rows):
                            self.all_manifest = self.tableWidget.item(self.all_row, 4).text()
                            self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                            self.all_name = self.tableWidget.item(self.all_row, 7).text()
                            self.specification=self.tableWidget.item(self.all_row, 8).text()
                            self.tableWidget.hideRow(self.all_row)
                            if self.sys_0 == self.all_sys:
                                # print('第',self.all_row,self.all_name)
                                self.tableWidget.showRow(self.all_row)
                                self.choose_sys_rows.append(self.all_row)#全选名称函数使用的行
                                if self.all_name not in name_dict:#向树界面添加材料名称
                                    name_dict[self.all_name] = ''
                                if self.specification not in specifications_dict:
                                    specifications_dict[self.specification]=''
                                if self.all_manifest not in manifest_dict:#向树界面清单匹配名称
                                    manifest_dict[self.all_manifest] = ''

                        self.r_name = QTreeWidgetItem(self.window.treeWidget_name)
                        self.r_name.setText(0, '全部名称')
                        self.r_name.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                        self.r_name.setText(1, '0')
                        self.r_name.setCheckState(0, Qt.Unchecked)
                        self.window.treeWidget_name.expandAll()
                        for key, value in name_dict.items():
                            self.grandson = QTreeWidgetItem(self.r_name)
                            self.grandson.setText(0, key)
                            self.grandson.setSizeHint(0, QSize(0, 25))
                            self.grandson.setCheckState(0, Qt.Unchecked)
                            self.grandson.setFont(0, font)
                            self.grandson.setFlags(
                                QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)

                        self.r_specification = QTreeWidgetItem(self.window.treeWidget_specification)
                        self.r_specification.setText(0, '全部规格')
                        self.r_specification.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                        self.r_specification.setText(1, '0')
                        self.r_specification.setCheckState(0, Qt.Unchecked)
                        self.window.treeWidget_specification.expandAll()
                        for key, value in specifications_dict.items():
                            self.grandson = QTreeWidgetItem(self.r_specification)
                            self.grandson.setText(0, key)
                            self.grandson.setSizeHint(0, QSize(0, 25))
                            self.grandson.setCheckState(0, Qt.Unchecked)
                            self.grandson.setFont(0, font)
                            self.grandson.setFlags(
                                QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                        for text in manifest_dict.keys():
                            self.qlist_text = QListWidgetItem()  # 创建QListWidgetItem实例
                            self.qlist_text.setText(text)
                            self.window.listWidget_sys.addItem(self.qlist_text)  # 添加到列表控件中
                            self.qlist_text.setSizeHint(QSize(0, 25))
                        self.tab_change = '提取'
    def startThread_run_system_Clickede(self):
        lock = threading.Lock()
        lock.acquire()  # 设置锁
        t1=threading.Thread(target=self.system_Clicked())
        t1.start()
        t1.join()
        lock.release()  # 释放锁
    def match_name(self):  # 读取所有匹配清单
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        self.window.listWidget_sys.clear()  # 清空列表
        if int(self.Item_row) != -1:
            self.itemzero = self.window.treeWidget_Items.currentItem()
            self.item_zero = self.itemzero.text(1)  # 获取1列内容
            # if int(self.item_zero)>=0:
            self.table_rows = self.tableWidget.rowCount()
            if self.table_rows != 0:
                match_dict = {}
                for row in range(0, self.table_rows):
                    self.text = self.tableWidget.item(row, 4).text()  # 获取单元格内
                    if self.text not in match_dict:
                        match_dict[self.text] = ''
                for text in match_dict.keys():
                    # self.window.listWidget_sys.addItems(match_list)
                    self.qlist_text = QListWidgetItem()  # 创建QListWidgetItem实例
                    self.qlist_text.setText(text)
                    self.window.listWidget_sys.addItem(self.qlist_text)  # 添加到列表控件中
                    self.qlist_text.setSizeHint(QSize(0, 25))

    def D_click(self):  # 双击Qlist文本添加到table
        text = self.window.listWidget_sys.currentItem().text()  # 当前选中对象，获取文本内容
        self.row = self.tableWidget.currentRow()  # 获取单元格行数
        if str(self.row) != '-1':
            self.tableWidget.setItem(self.row, 4, QTableWidgetItem(str(text)))

    check_names_dict={}
    def check_names(self, item, cloumn):  # 提取材料名称
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            self.table_rows = self.tableWidget.rowCount()
            if self.table_rows != '0':
                if item.checkState(cloumn) == Qt.Checked:
                    self.check_names_dict[self.sys_0, item.text(0)] = ''
                if item.checkState(cloumn) == Qt.Unchecked:
                    if (self.sys_0, item.text(0)) in self.check_names_dict:
                        self.check_names_dict.pop((self.sys_0, item.text(0)))
                # if item.checkState(cloumn) == Qt.Checked:  # 全选子节点
                #     if item.text(0) == '整个工程':
                #         n = self.window.treeWidget_name.topLevelItemCount()  # 获取根节点数量
                #         for i in range(0, n):
                #             ite = self.window.treeWidget_name.topLevelItem(i)  # 循环获取根节点
                #             count = ite.childCount()  # 获取当前根节点的子节点数量
                #             for j in range(0, count):
                #                 string = item.child(j)  # 子节点的文字信息
                #                 string.setCheckState(0, Qt.Checked)
                # if item.checkState(cloumn) == Qt.Unchecked:
                #     if item.text(0) == '整个工程':
                #         n = self.window.treeWidget_name.topLevelItemCount()  # 获取根节点数量
                #         for i in range(0, n):
                #             ite = self.window.treeWidget_name.topLevelItem(i)  # 循环获取根节点
                #             count = ite.childCount()  # 获取当前根节点的子节点数量
                #             for j in range(0, count):
                #                 strin = item.child(j)  # 子节点的文字信息
                #                 if strin != None:
                #                     strin.setCheckState(0, Qt.Unchecked)
                # if self.check_names_dict != {}:
                #     for self.all_row in range(0, self.table_rows):
                #         self.tableWidget.hideRow(self.all_row)
                #     for sys_names in self.check_names_dict.keys():
                #         for self.all_row in range(0, self.table_rows):
                #             self.all_manifest = self.tableWidget.item(self.all_row, 4).text()
                #             self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                #             self.all_name = self.tableWidget.item(self.all_row, 7).text()
                #             if sys_names[0] != '整个工程':
                #                 if sys_names[0] == self.all_sys and sys_names[1] == self.all_name:
                #                     # print('第',self.all_row,self.all_sys,self.all_name)
                #                     self.tableWidget.showRow(self.all_row)
                #             if sys_names[0] == '整个工程':
                #                 if sys_names[1] == self.all_name:
                #                     # print('第',self.all_row,self.all_sys,self.all_name)
                #                     self.tableWidget.showRow(self.all_row)
                # if self.check_names_dict == {}:
                #     for self.all_row in range(0, self.table_rows):
                #         self.all_manifest = self.tableWidget.item(self.all_row, 4).text()
                #         self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                #         self.all_name = self.tableWidget.item(self.all_row, 7).text()
                #         self.tableWidget.hideRow(self.all_row)
                #         if self.sys_0 == self.all_sys:
                #             # print('第',self.all_row,self.all_name)
                #             self.tableWidget.showRow(self.all_row)
                #         if self.sys_0 == '整个工程':
                #             self.tableWidget.showRow(self.all_row)

    def up_check(self):#按输入的文本选中材料名称
        text=self.window.lineEdit.text()#获取文本内容
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            n = self.window.treeWidget_name.topLevelItemCount()  # 获取根节点数量
            for i in range(0, n):
                ite = self.window.treeWidget_name.topLevelItem(i)  # 循环获取根节点
                count = ite.childCount()  # 获取当前根节点的子节点数量
                for j in range(0, count):
                    string = ite.child(j)  # 子节点的文字信息
                    check_name = string.text(0)  # 子节点的文字信息
                    if text!='' and text in check_name:
                        string.setCheckState(0,Qt.Checked)
    def specification_check(self):#按输入的文本选中规格
        text=self.window.lineEdit_specification.text()#获取文本内容
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            n = self.window.treeWidget_specification.topLevelItemCount()  # 获取根节点数量
            for i in range(0, n):
                ite = self.window.treeWidget_specification.topLevelItem(i)  # 循环获取根节点
                count = ite.childCount()  # 获取当前根节点的子节点数量
                for j in range(0, count):
                    string = ite.child(j)  # 子节点的文字信息
                    check_name = string.text(0)  # 子节点的文字信息
                    if text!='' and text in check_name:
                        string.setCheckState(0,Qt.Checked)

    check_specification_dict={}
    def check_specification(self, item, cloumn):
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            self.table_rows = self.tableWidget.rowCount()
            if self.table_rows != '0':
                if item.checkState(cloumn) == Qt.Checked:
                    self.check_specification_dict[item.text(0)] = ''
                if item.checkState(cloumn) == Qt.Unchecked:
                    if item.text(0) in self.check_specification_dict:
                        self.check_specification_dict.pop(item.text(0))
    # 多线程显示名称
    def startThread_run_show_name(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_show_name)
        self.Mywork.stopSing.connect(self.stopThread_show_name)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行

    def stopThread_show_name(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
        print('结束线程')

    def show_name_row(self):# 点击按钮显示隐藏的行
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            for show_row in self.name_row_list:
                self.tableWidget.showRow(show_row)

    name_row_list = []#全选名称行列表
    row_values_dict={}
    def show_name(self):#按选择的材料显示
        try:
            #
            # self.sys_0 = self.window.treeWidget_system.topLevelItem(0).text(0)  # 获取0列内容
            # print(self.sys_0)
            self.tab_change = '不提取'
            self.name_row_list=[]
            self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
                self.check_specification_dict = {} # 对规格字典清零
                self.row_values_dict={}
                self.table_rows = self.tableWidget.rowCount()
                if self.check_names_dict != {}:# 按筛选的材料匹配行
                    for self.all_row in range(0, self.table_rows):
                        self.tableWidget.hideRow(self.all_row)
                    for sys_names in self.check_names_dict.keys():
                        print(sys_names)
                        for self.all_row in range(0, self.table_rows):
                            self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                            self.all_name = self.tableWidget.item(self.all_row, 7).text()
                            self.all_specification = self.tableWidget.item(self.all_row, 8).text()
                            if sys_names[0] != '整个工程':#筛选行
                                if sys_names[0] == self.all_sys and sys_names[1] == self.all_name:
                                    self.tableWidget.showRow(self.all_row)
                                    self.name_row_list.append(self.all_row)
                                    self.row_values_dict[str(self.all_row)]=self.all_specification
                            if sys_names[0]== '整个工程' or sys_names[0]=='':#筛选行
                                if sys_names[1] == self.all_name:
                                    self.tableWidget.showRow(self.all_row)
                                    self.name_row_list.append(self.all_row)
                                    self.row_values_dict[str(self.all_row)]=self.all_specification
                if self.row_values_dict!={}:
                    font = QFont()
                    font.setPointSize(10)  # 设置字体大小为10像素
                    font.setFamily("宋体")
                    self.window.treeWidget_specification.clear()
                    # 将筛选的材料对应的规格筛选出来
                    self.r_specification = QTreeWidgetItem(self.window.treeWidget_specification)
                    self.r_specification.setText(0, '全部规格')
                    self.r_specification.setFlags(
                        QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                    self.r_specification.setText(1, '0')
                    self.r_specification.setCheckState(0, Qt.Unchecked)
                    self.r_specification.setFont(0, font)

                    # 将筛选的材料对应的规格筛选出来
                    for self.all_specification in set(self.row_values_dict.values()):
                        self.grandson = QTreeWidgetItem(self.r_specification)
                        self.grandson.setText(0, str(self.all_specification))
                        self.grandson.setSizeHint(0, QSize(0, 25))
                        self.grandson.setCheckState(0, Qt.Unchecked)
                        self.grandson.setFont(0, font)
                        self.grandson.setFlags(
                            QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                    self.window.treeWidget_specification.expandAll()
                # if self.check_names_dict == {}:#按整个工程或系统筛选行
                #     for self.all_row in range(0, self.table_rows):
                #         self.all_manifest = self.tableWidget.item(self.all_row, 4).text()
                #         self.all_sys = self.tableWidget.item(self.all_row, 6).text()
                #         self.all_name = self.tableWidget.item(self.all_row, 7).text()
                #         self.tableWidget.hideRow(self.all_row)
                #         if self.sys_0 == self.all_sys:
                #             self.name_row_list.append(self.all_row)
                #             self.tableWidget.showRow(self.all_row)
                #         if self.sys_0 == '整个工程':
                #             self.tableWidget.showRow(self.all_row)
                #             self.name_row_list.append(self.all_row)
                # if self.name_row_list!=[]:#多线程显示行
                #     self.startThread_run_show_name()
                self.tab_change = '提取'
        except:
            pass
    def filter_sp(self):
        if self.row_values_dict!={} and self.check_specification_dict!={}:
            # print(self.check_specification_dict.keys())
            for row,values in self.row_values_dict.items():
                self.tableWidget.hideRow(int(row))
                for specification in self.check_specification_dict.keys():
                    print(values,specification)
                    if values==specification:
                        self.tableWidget.showRow(int(row))
                        print(row,values)

    unit_sys_list = []#合并系统第一步提取树界面系统
    def check_unit_sys(self,item, cloumn):
        if item.checkState(cloumn) == Qt.Checked:
            if item.text(1) != '0':
                self.unit_sys_list.append(item.text(0))
        if item.checkState(cloumn) == Qt.Unchecked:
            if item.text(0) in self.unit_sys_list:
                self.unit_sys_list.remove(item.text(0))
        if item.checkState(cloumn) == Qt.Checked:  # 全选子节点
            if item.text(1) == '0':
                n = self.window.treeWidget_system.topLevelItemCount()  # 获取根节点数量
                for i in range(0, n):
                    ite = self.window.treeWidget_system.topLevelItem(i)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        string.setCheckState(0, Qt.Checked)

        if item.checkState(cloumn) == Qt.Unchecked:
            if item.text(1) == '0':
                n = self.window.treeWidget_system.topLevelItemCount()  # 获取根节点数量
                for i in range(0, n):
                    ite = self.window.treeWidget_system.topLevelItem(i)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    for j in range(0, count):
                        strin = ite.child(j)  # 子节点的文字信息
                        if strin != None:
                            strin.setCheckState(0, Qt.Unchecked)
    def unit_sys(self):#合并系统第二步将符合的系统汇总
        if self.unit_sys_list!=[]:
            self.tab_change = '不提取'
            value_st='、'.join(self.unit_sys_list)
            self.unit_table_list=[]# table行构建的列表
            self.unit_sys_dict = {}# 符合系统构建的字典
            self.remainder_list=[]# 不符合系统构建的列表
            self.unit_list=[]# 符合系统构建的字典转列表
            self.table_rows = self.window.tableWidget_0.rowCount()
            self.table_column = self.window.tableWidget_0.columnCount()
            for i in range(0, int(self.table_rows)):
                unit_table=[]
                self.unit_table_list.append(unit_table)
                for j in range(0, int(self.table_column)):
                    self.text = self.window.tableWidget_0.item(i, j).text()  # 获取单元格内容
                    unit_table.append(self.text)
            for unit in self.unit_table_list:
                if unit[6] in value_st and tuple(unit[0:6]+unit[7:10]) not in self.unit_sys_dict:
                    self.unit_sys_dict[tuple(unit[0:6]+unit[7:10])] = [value_st,unit[10]]
                elif unit[6] in value_st and tuple(unit[0:6]+unit[7:10]) in self.unit_sys_dict:
                    self.unit_sys_dict[tuple(unit[0:6] + unit[7:10])].append(unit[10])
                if unit[6] not in value_st:
                    self.remainder_list.append(unit[0:6]+unit[7:10]+unit[6:7]+unit[10:])

            for key,value in self.unit_sys_dict.items():
                key_list=list(key)
                key_list+=value[0:1]+['+'.join(value[1:])]
                print(key_list)
                self.unit_list.append(key_list)
            all_list=self.unit_list+self.remainder_list#总列表
            self.window.tableWidget_0.clearContents()#可以清除表格所有的内容
            # self.window.tableWidget_0.setRowCount(len(all_list))  # 设置行
            for row,v in enumerate(all_list):
                self.window.tableWidget_0.setRowCount(row+1)  # 设置行
                self.window.tableWidget_0.setItem(row, 0, QTableWidgetItem(v[0]))
                self.window.tableWidget_0.setItem(row, 1, QTableWidgetItem(v[1]))
                self.window.tableWidget_0.setItem(row, 2, QTableWidgetItem(v[2]))
                self.window.tableWidget_0.setItem(row, 3, QTableWidgetItem(v[3]))
                self.window.tableWidget_0.setItem(row, 4, QTableWidgetItem(v[4]))
                self.window.tableWidget_0.setItem(row, 5, QTableWidgetItem(v[5]))
                self.window.tableWidget_0.setItem(row, 6, QTableWidgetItem(str(v[9])))#写入系统
                self.item1 = QTableWidgetItem(v[6])
                self.item1.setCheckState(Qt.Unchecked)
                self.window.tableWidget_0.setItem(row, 7, QTableWidgetItem(self.item1))#写入名称
                self.window.tableWidget_0.setItem(row, 8, QTableWidgetItem(str(v[7])))  # 写入规格
                self.window.tableWidget_0.setItem(row, 9, QTableWidgetItem(str(v[8])))#写入单位
                self.window.tableWidget_0.setItem(row, 10, QTableWidgetItem(str(v[10])))  # 工程量
                self.window.tableWidget_0.setItem(row, 11, QTableWidgetItem(''))  # 备注
            self.unit_sys_list.clear()
            self.write_sys()
            self.match_name()
            self.tab_change = '提取'

    def shear_text(self): # 剪切文本
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            if self.tableWidget.currentRow() != -1:
                self.tab_change = '不提取'
                # 撤销
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                # 剪切文本
                row = self.tableWidget.currentRow()  # 获取单元格行数
                column = self.tableWidget.currentColumn()
                self.text_dict = {}
                if str(row) != '-1':
                    mytable = self.tableWidget.selectedItems()
                    for r in mytable:  # 第三种方法获取值
                        contents = self.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                        # print('第',r.row(),'行','第',r.column(),'列',contents)
                        if r.row() not in self.text_dict:
                            self.text_dict[r.row()] = [contents]
                        elif r.row() in self.text_dict:
                            self.text_dict[r.row()].append(contents)
                    for r in mytable:  # 第三种方法获取值
                        self.tableWidget.setItem(r.row(), r.column(), QTableWidgetItem(''))
                    self.go = '执行'
                    # 恢复
                    if self.table_do == '执行':
                        self.new_undo_dict = {}
                        new_text_list = []
                        self.new_undo_dict[self.tableWidget] = new_text_list
                        for self.Single_rows in range(0, self.tableWidget.rowCount()):
                            t_list = []
                            new_text_list.append(t_list)
                            for self.Single_colum in range(0, self.tableWidget.columnCount()):
                                self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)
                        if self.Click_list != new_text_list:
                            print(new_text_list)
                            command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict, self.new_undo_dict)
                            self.undoStack_del.push(command)
                    self.tab_change = '提取'
    def Clipboard(self):
        self.go='剪切板'
    go='不执行'
    def tool_copy(self):#复制文本
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            if self.tableWidget.currentRow() != -1:
                self.tab_change = '不提取'
                row = self.tableWidget.currentRow()  # 获取单元格行数
                column = self.tableWidget.currentColumn()
                self.text_dict={}
                if str(row) != '-1':
                    mytable = self.tableWidget.selectedItems()
                    for r in mytable:  # 第三种方法获取值
                        contents = self.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                        # print('第',r.row(),'行','第',r.column(),'列',contents)
                        if r.row() not in self.text_dict:
                            self.text_dict[r.row()]=[contents]
                        elif r.row() in self.text_dict:
                            self.text_dict[r.row()].append(contents)
                    self.go = '执行'
                    self.tab_change = '提取'

    # 粘贴
    def startThread_run_tool_paste(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_tool_paste)
        self.Mywork.stopSing.connect(self.stopThread_tool_paste)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行

    def stopThread_tool_paste(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
    def tool_paste(self):#粘贴文本
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            if self.tableWidget.currentRow() != -1:
                self.tab_change = '不提取'
                # 撤销
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.go=='执行':
                    mytable = self.tableWidget.selectedItems()
                    for row,values in enumerate(self.text_dict.values()):
                        for column, value in enumerate(values):
                            if len(self.text_dict.values())==1 and len(values)==1:
                                for r in mytable:#第三种方法获取值
                                    self.tableWidget.setItem(r.row(),r.column(), QTableWidgetItem(str(value)))
                            else:
                                self.tableWidget.setItem(self.tableWidget.currentRow()+row, self.tableWidget.currentColumn()+column,QTableWidgetItem(str(value)))
                try:
                    if self.go == '剪切板':
                        mimeData = self.clipboard.mimeData()
                        if mimeData.hasFormat('text/plain'):

                            all_text=[]
                            test_list = []
                            values_list=[]
                            texts = mimeData.text()
                            print(texts)
                            if '\t' in texts:
                                criteria = '"(.*?)"'
                                compile = re.compile(criteria, re.S)
                                self.results = compile.findall(texts)
                                sp_n = texts.replace('"', '').split('\n')#拆n
                                for sp in sp_n:
                                    sp_t=str(sp).split('\t')#拆t
                                    test_list.append(sp_t)
                                for new_list in test_list:
                                    for new in new_list:
                                        all_text.append(str(new))#构建新列表

                                for result in self.results:
                                    res=str(result).split('\n')[1:]
                                    for rs in res:
                                        all_text.remove(str(rs))#删除新列表同一个单元格内多余的换行后内容
                                if '' in all_text[-1]:
                                    all_text.pop(-1)

                                row=len(all_text)-texts.count('\t')#获取行
                                column=int(len(all_text)/row-1)#获取列
                                num=int(len(all_text)/row)

                                split_t = texts.replace('"', '').split('\t')#拆t
                                for i in range(0, row):  # 循环行数
                                    split_n = split_t[column].split('\n')[:-1]#在一个单元格内有多余的换行符，需要截取有用的内容
                                    join_n='\n'.join(split_n)
                                    last = split_t[column].split('\n')[-1]#截取最后的换行符内容
                                    split_t[column]=join_n#修改列表内容
                                    split_t.insert(column+1,last)#插入列表内容
                                    column+=num

                                if '' in split_t[-1]:
                                    split_t.pop(-1)#删除最后一个元素‘’

                                for r in range(row):#按行循环
                                    values_list.append(split_t[0:num])#按列添加到列表
                                    for c in range(num):#按列删除列表内容
                                        split_t.pop(0)
                                    for row, values in enumerate(values_list):#每次循环一行写入到表格里一行内容
                                        for column, value in enumerate(values):
                                            self.tableWidget.setItem(self.tableWidget.currentRow()+row,self.tableWidget.currentColumn() + column,QTableWidgetItem(str(value)))
                            if '\t' not in texts and '"\n' in texts:
                                sp_n = texts.split('"')  # 拆n
                                text_list=[]
                                for sp in sp_n:
                                    if sp!='':
                                        if sp!='\n':
                                            a=sp.split('\n')
                                            if a[0]!='' and a [-1]!='':
                                                join_text='\n'.join(a)
                                                text_list.append(str(join_text))

                                            if a[-1]=='':
                                                for b in a:
                                                    if b!='':
                                                        text_list.append(str(b))

                                for row, value in enumerate(text_list):  # 每次循环一行写入到表格里一行内容
                                    self.tableWidget.setItem(self.tableWidget.currentRow()+row,
                                                             self.tableWidget.currentColumn(),
                                                             QTableWidgetItem(str(value)))

                            if '\t' not in texts and '"\n' not in texts:
                                sp_n = texts.split('\n')
                                if '' in sp_n[-1]:
                                    sp_n.pop(-1)#删除最后一个元素‘’
                                for row, value in enumerate(sp_n):  # 每次循环一行写入到表格里一行内容
                                    self.tableWidget.setItem(self.tableWidget.currentRow() + row,
                                                             self.tableWidget.currentColumn(),
                                                             QTableWidgetItem(str(value)))
                except:
                    pass

                self.tableWidget.viewport().update()
                # 回撤
                if self.table_do == '执行':
                    self.new_undo_dict = {}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'


    rows_values=[]
    def copy_row(self):#复制行
        self.rows_values=[]
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.table_column = int(self.tableWidget.columnCount())
            if self.tableWidget_allrows!=0:
                self.rows = []
                self.column=[]
                mytable = self.tableWidget.selectedItems()
                for r in mytable:#第三种方法获取值
                    if r.row() not in self.rows:
                        self.rows.append(r.row())
                    self.column.append(r.column())
                self.rows.sort(reverse=False)
                for self.Single_rows in self.rows:
                    text_list = []
                    if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                            and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                        self.rows_values.append(text_list)
                        for self.Single_colum in range(0, self.table_column):
                            copytext = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            text_list.append(copytext)
    def paste_row(self):#粘贴行
        if self.rows_values!=[]:
            self.tab_change = '不提取'
            self.row = self.tableWidget.currentRow()  # 获取单元格行数
            if self.row != -1:
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)

                for insert_row in range(0,len(self.rows)):
                    self.tableWidget.insertRow(self.row+insert_row+1)
                    for j in range(0, self.tableWidget.columnCount()):
                        item = QTableWidgetItem(self.rows_values[insert_row][j])
                        self.tableWidget.setItem(self.row+insert_row+1, j, item)
                        if j == 7:
                            item.setCheckState(Qt.Unchecked)
                            self.tableWidget.setItem(self.row+insert_row+1, 7, item)
                # 恢复
                if self.table_do == '执行':
                    self.new_undo_dict={}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'

    def shear_row(self):#剪切行
        self.tab_change = '不提取'
        self.rows_values = []
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.table_column = int(self.tableWidget.columnCount())
            if self.tableWidget_allrows !=0:
                # 撤销
                self.old_undo_dict={}
                self.Click_list = []
                self.old_undo_dict[self.tableWidget] = self.Click_list
                for self.Single_rows in range(0, self.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, self.tableWidget.columnCount()):
                        self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)

                self.rows = []
                self.column=[]
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    if r.row() not in self.rows:
                        self.rows.append(r.row())
                    self.column.append(r.column())
                self.rows.sort(reverse=False)
                for self.Single_rows in self.rows:
                    text_list = []
                    self.rows_values.append(text_list)
                    for self.Single_colum in range(0, self.table_column):
                        copytext = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(copytext)
                mytable = self.tableWidget.selectedItems()
                num=0
                for i in range(0,len(self.rows)):
                    if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column\
                            and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                        self.tableWidget.removeRow(mytable[num].row())
                        num+=self.table_column
                # 回撤
                if self.table_do == '执行':
                    self.new_undo_dict={}
                    new_text_list = []
                    self.new_undo_dict[self.tableWidget] = new_text_list
                    for self.Single_rows in range(0, self.tableWidget.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tableWidget.columnCount()):
                            self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    if self.Click_list != new_text_list:
                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                        self.undoStack_del.push(command)
                self.tab_change = '提取'

    def hide_rows(self):# 隐藏行
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.table_rows = self.tableWidget.currentRow()
            if self.table_rows!=-1:
                rows=[]
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if r.row()!=0 and r.row() not in rows:
                        rows.append(r.row())
                if rows!=[]:
                    rows.sort(reverse=False)
                    for hide_row in rows:
                        if hide_row!=self.tableWidget.rowCount()-1:
                            self.tableWidget.hideRow(hide_row)

    def startThread_run_tool_show_row(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_tool_show_row)
        self.Mywork.stopSing.connect(self.stopThread_tool_show_row)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
    def stopThread_tool_show_row(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
    def tool_show_row(self):# 点击按钮显示隐藏的行
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.table_rows =self.tableWidget.currentRow()
            if self.table_rows !=-1:
                rows = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if self.table_rows!=0:
                        if r.row() not in rows:
                            rows.append(r.row())
                rows.sort(reverse=False)
                if rows!=[]:
                    for show_row in range(min(rows),max(rows)):
                        self.tableWidget.showRow(show_row)
    def hide_column(self):#隐藏列
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.table_rows =self.tableWidget.currentRow()
            if self.table_rows !=-1:
                column= []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if r.column()!=0 and r.column() not in column:
                        column.append(r.column())
                if column!=[]:
                    column.sort(reverse=False)
                    for hide_column in column:
                        if hide_column!=self.tableWidget.columnCount()-1:
                            self.tableWidget.hideColumn(hide_column)
    def show_column(self):# 点击按钮显示隐藏的列
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.table_rows=self.tableWidget.currentRow()  # 获取单元格行数
            if self.table_rows !=-1:
                column = []
                mytable = self.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if r.column() not in column:
                        column.append(r.column())
                column.sort(reverse=False)
                if column!=[]:
                    for show_column in range(min(column),max(column)):
                        self.tableWidget.showColumn(show_column)

    def startThread_run_show_rows(self):# 点击整个工程显示所有行
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_show_rows)
        self.Mywork.stopSing.connect(self.stopThread_show_rows)  # 停止信号连接到stopThread方法
        self.thread.finished.connect(self.thread_finished)
        # if self.Mywork.is_running == True:
        self.thread.start()  # 开始线程的运行
    def stopThread_show_rows(self):
        self.thread.quit()  # 退出
        self.thread.wait()  # 回收资源
        self.thread.terminate()

    def thread_finished(self):
        print('结束线程')
        pass
    def show_rows(self):# 点击整个工程显示所有行
        self.table_rows = self.tableWidget.rowCount()
        if self.table_rows != '0':
            for self.all_row in range(0, self.table_rows):
                self.tableWidget.showRow(self.all_row)

    def sum_quantity(self):#工程量求和
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            self.table_column = int(self.tableWidget.columnCount())
            if str(self.tableWidget_allrows) != '0':
                for self.Single_rows in range(0, self.tableWidget_allrows):
                    self.text = self.tableWidget.item(self.Single_rows,10).text()
                    # print(self.Single_rows,self.text)
                    w = []
                    float_num = re.compile(r'\D+')
                    float_re =float_num.findall(self.text)
                    for i in float_re:
                        if '.' != i:
                            if '+' != i:
                                if ' '!=i:
                                    w.append(i)
                    num = 0
                    if w == []:
                        rp=str(self.text).replace(' ','')
                        if rp!='':
                            s = rp.split('+')
                            for t in s:
                                if t != '':
                                    num += float(t)
                                    print(num)
                    if num!=0:
                        self.tableWidget.setItem(self.Single_rows,10, QTableWidgetItem(str(round(num, 3))))
    def part_sum(self):#局部求和
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
            row = self.tableWidget.currentRow()
            if str(self.tableWidget_allrows) != '0':
                values_list = []
                if str(row) != '-1':
                    mytable = self.tableWidget.selectedItems()
                    for r in mytable:  # 第三种方法获取值
                        if r.column()==10:
                            value_list = []
                            contents = self.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                            float_num = re.compile(r"\D+")
                            float_re = float_num.findall(str(contents).replace(' ',''))
                            for i in float_re:
                                if '.' != i:
                                    value_list.append(i)
                            if value_list==[]:
                                content=str(contents).replace(' ','')
                                if content!='':
                                    values_list.append(float(content))
                    text=int(round(sum(values_list), 0))
                    QMessageBox.information(self.window, '工程量', '汇总工程量是:{}  '.format(str(text)))
                    # print(sum(values_list))


    def image_text(self):
        self.image_text =image_text_window()
        self.image_text.window.show()
    def pdf_tool(self):# pdf工具
        self.pdf_tool_windows =pdf_tool_windows()
        self.pdf_tool_windows.window.show()

    def speck_tool(self):# 语音工具窗口
        self.speck_tool_windows =text_speck_window()
        self.speck_tool_windows.window.show()

    def open_Text_window(self,Item=None):# 加载文本修改窗口
        self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            if self.tableWidget.currentColumn()==2:
                self.Text_window =text_Window()
                if Item == None:
                    return
                self.va = Item.text()
                self.Text_window.window.plainTextEdit.setPlainText(str(self.va))
                self.Text_window.window.show()

    def open_qdk_window(self):# 加载清单库窗口
        self.qdk_window =QDK_Window()
        self.qdk_window.window.show()


    def find_replace(self):#加载查找替换窗口
        self.my_window=f_r_window()
        self.my_window.window.show()

    def all_checks(self):
        self.check_window=check_Window()
        self.check_window.window.show()

    def startThread_run_time_save(self):
        self.thread = QThread()  # 实例化一个线程
        self.Mywork = Work()  # 实例化工作类,并传入参数到工作线程进行加工
        self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
        self.thread.started.connect(self.Mywork.run_time_save)
        self.Mywork.stopSing.connect(self.stopThread_time_save)  # 停止信号连接到stopThread方法
        self.thread.start()  # 开始线程的运行
        print('第一次多线程运行')
    def stopThread_time_save(self):
        print('多线程结束')
        if self.json_save_list!=[]:
            print(self.json_save_list)
            messageBox = QMessageBox()
            messageBox.resize(300, 600)
            messageBox.setWindowTitle('保存')
            messageBox.setText('"10分钟提示", "请保存工程！"')
            messageBox.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
            buttonYes = messageBox.button(QMessageBox.Yes)
            buttonYes.setText("确定")
            buttoncancel = messageBox.button(QMessageBox.Cancel)
            buttoncancel.setText("取消")
            messageBox.exec_()
            if messageBox.clickedButton() == buttonYes:
                save_json = []
                self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
                if int(self.Item_row) != -1:
                    n = self.window.treeWidget_Items.topLevelItemCount()  # 获取根节点数量
                    for i in range(0, n):
                        ite = self.window.treeWidget_Items.topLevelItem(i)  # 循环获取根节点
                        name = ite.text(0)
                        num_max = ite.text(2)
                        self.table0_json = []
                        save_json.append(self.table0_json)
                        table0_alldict = {}
                        table0_allvalues = []
                        for self.Single_rows in range(0, self.window.tableWidget_0.rowCount()):
                            table0_values = []
                            table0_allvalues.append(table0_values)
                            for self.Single_colum in range(0, self.window.tableWidget_0.columnCount()):
                                self.text0 = self.window.tableWidget_0.item(self.Single_rows, self.Single_colum).text()
                                table0_values.append(self.text0)
                        table0_alldict[name + '$' + num_max] = table0_allvalues
                        self.table0_json.append(table0_alldict)
                        count = ite.childCount()  # 获取当前根节点的子节点数量
                        for j in range(0, count):
                            self.table_json = []
                            save_json.append(self.table_json)
                            table_alldict = {}
                            table_allvalues = []
                            string = ite.child(j)  # 子节点的文字信息
                            self.item_zero = string.text(0)  # 获取0列内容
                            self.item_zer1 = string.text(1)  # 获取1列内容
                            self.item = self.item_zero, self.item_zer1
                            self.tableWidget = self.tablelist[int(self.item_zer1)]
                            self.tableWidget_allrows = int(self.tableWidget.rowCount())  # 获取总行数
                            self.table_column = int(self.tableWidget.columnCount())
                            for self.Single_rows in range(0, self.tableWidget_allrows):
                                table_values = []
                                table_allvalues.append(table_values)
                                for self.Single_colum in range(0, self.table_column):
                                    self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                    table_values.append(self.text)
                            table_alldict[self.item_zero + '$' + self.item_zer1] = table_allvalues
                            self.table_json.append(table_alldict)
                    self.data_json = json.dumps(save_json, ensure_ascii=False)  # python转json文。
                    with open(self.fileName_choose, "w") as f:
                        json.dump(self.data_json, f)  # 禁止ascii转换，这样就可以打印中文。

            if messageBox.clickedButton() == buttoncancel:
                print('取消')
        else:
            self.json_save()
            print('执行未保存循环')
        self.thread.quit()  # 退出
        self.timer.singleShot(600000, self.startThread_run_time_save)

    def closeEvent(self, event):
        messageBox = QMessageBox()
        messageBox.resize(300, 600)
        messageBox.setWindowTitle('确认')
        messageBox.setText('"温馨提示", "是否保存当前工程"')
        messageBox.setStandardButtons(QMessageBox.Yes | QMessageBox.Ok | QMessageBox.Close)
        buttonYes = messageBox.button(QMessageBox.Yes)
        buttonYes.setText("保存")
        buttonOk = messageBox.button(QMessageBox.Ok)
        buttonOk.setText("不保存")
        buttonC = messageBox.button(QMessageBox.Close)
        buttonC.setText('取消')
        messageBox.exec_()
        if messageBox.clickedButton() == buttonYes:
            self.json_save()
            event.accept()
        elif messageBox.clickedButton() == buttonOk:
            event.accept()
        elif messageBox.clickedButton() == buttonC:
            event.ignore()
    def movie(self):# 播放视频
        # 定义一个播放器对象
        player = QMediaPlayer()
        # 设置播放器的窗口
        video_widget = QVideoWidget()
        # 设置播放器输出窗口
        player.setVideoOutput(video_widget)
        # 播放器绑定视频路径
        video_url = QUrl(r"D:\2023年广联达培训\山西二建\山西二建安装5月29号录屏.mp4")
        player.setMedia(QMediaContent(video_url))
        gridLayout.addWidget(self.window.widget_movie, 0, 0)
        # 开始播放视频
        player.play()
        # 视频暂停
        # player.pause()

class Thread_load_qdc(QThread):# 加载清单池
    qmut = QMutex()
    stopSing=pyqtSignal(list)
    def __init__(self):
        super(Thread_load_qdc, self).__init__()
    def run(self):
        self.qmut.lock()  # 加锁
        BASE_DIR = os.path.dirname(__file__)  # 清单池
        files_qdc = BASE_DIR + '/' + "清单数据库" + '/' + "清单池"
        self.all_qdc_list = []
        for dirpath, dirnames, filenames in os.walk(files_qdc):
            for filename in filenames:
                files = os.path.join(dirpath, filename)
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    for self.ws in self.wb.worksheets:
                        self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        for value in self.values:
                            value_list=[]
                            for v in value:
                                value_list.append(str(v).replace('None',''))
                            self.all_qdc_list.append(value_list)
        self.qmut.unlock()  # 解锁
        self.stopSing.emit(self.all_qdc_list)

class Thread_load_dek(QThread):# 多线程加载定额库
    qmut = QMutex()
    runSing = pyqtSignal(list)
    stopSing=pyqtSignal(list)
    def __init__(self,msg,files):
        super(Thread_load_dek, self).__init__()
        self.text=msg
        self.files=files
    def run(self):
        self.qmut.lock()  # 加锁
        # for dirpath, dirnames, filenames in os.walk(self.files):
        #     for filename in filenames:
        #         files = os.path.join(dirpath, filename)
        if '~$' not in self.files and 'xlsx' == self.files.split('.')[-1]:
            # print(filename)
            self.wb = openpyxl.load_workbook(self.files, read_only=False, data_only=True, keep_links=False)
            self.num = 0
            for i in range(len(self.wb.sheetnames)):
                if self.text == self.wb.sheetnames[i]: #or self.item.text(0)==str(filename).split('.')[0]
                    # self.num = 0
                    self.ws = self.wb[self.wb.sheetnames[i]]
                    # print(self.ws)
                    self.values = list(
                        self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                    self.qmut.unlock()  # 解锁
                    self.stopSing.emit(self.values)

class Work_qd(QObject):# 多线程
    qmut = QMutex()
    startSing = pyqtSignal()
    stopSing=pyqtSignal(list)
    def __init__(self,text,files_qd):
        super().__init__()
        self.files_qdk=files_qd
        self.files_qdc=files_qd
        self.files_quota = files_qd
        self.text=text
    def run_Search_qdk(self):# 搜索清单库
        self.qmut.lock()  # 加锁
        all_text_list=[]
        for dirpath, dirnames, filenames in os.walk(self.files_qdk):
            for filename in filenames:
                files = os.path.join(dirpath, filename)
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    for self.ws in self.wb.worksheets:
                        self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        for value in self.values:
                            value_list=[]
                            for v in value:
                                value_list.append(str(v).replace('None',''))
                            if self.text in ''.join(value_list):
                                all_text_list.append(value_list)

        self.qmut.unlock()  # 解锁
        self.stopSing.emit(all_text_list)
    def run_Search_qdc(self):# 搜索清单池
        self.qmut.lock()  # 加锁
        all_text_list=[]
        for dirpath, dirnames, filenames in os.walk(self.files_qdc):
            for filename in filenames:
                files = os.path.join(dirpath, filename)
                print(files)
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    for self.ws in self.wb.worksheets:
                        self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        for value in self.values:
                            value_list=[]
                            for v in value:
                                value_list.append(str(v).replace('None',''))
                            if self.text in ''.join(value_list):
                                all_text_list.append(value_list)
        self.qmut.unlock()  # 解锁
        self.stopSing.emit(all_text_list)

    def run_Search_quota(self):# 搜索定额库
        self.qmut.lock()  # 加锁
        all_text_list=[]
        for dirpath, dirnames, filenames in os.walk(self.files_quota):
            for filename in filenames:
                files = os.path.join(dirpath, filename)
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    for self.ws in self.wb.worksheets:
                        self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        for value in self.values:
                            value_list=[]
                            for v in value:
                                value_list.append(str(v).replace('None',''))
                            if self.text in ''.join(value_list):
                                all_text_list.append(value_list)
        self.qmut.unlock()  # 解锁
        self.stopSing.emit(all_text_list)

class  text_speck_window():
    def __init__(self):
        super().__init__()
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "text_speck_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('文字转语音窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.4)
        self.width = int(self.screenwidth * 0.4)
        self.window.resize(self.width, self.height)
        self.window.pushButton_save.clicked.connect(self.save_voice)  # 语音工具

        # self.window.textEdit.append(self.python_data['task_id'])
        # self.window.lineEdit.returnPressed.connect(self.mp3)  # 绑定enter键
        self.API_KEY = "YhMifyFMzxGmBHywvYx6iAZL"  # ak,控制台内创建app获取
        self.SECRET_KEY = "GIDgHSc2vQpSDHREIsAL7XFPiiWDXdkU"  # sk,控制台内创建app获取
    task_id_list = []
    def text_speck(self,location):
        info = self.window.textEdit.toPlainText()  # 通过 toPlainText 方法获取编辑框内的文本内容
        if info!='':
            self.engine = pyttsx4.init()  # 初始化语音引擎
            self.engine.setProperty('rate', 170)  # 设置语速
            volume = self.engine.getProperty('volume')
            self.engine.setProperty('volume',str(volume))
            voices = self.engine.getProperty('voices')
            print(voices)
            self.engine.setProperty('voice', voices[0].id)  # 设置第一个语音合成器
            if location == '开始':
                # self.engine.stop()
                self.engine.say(info)
                self.engine.runAndWait()
                self.engine.stop()
            if location == '停止':
                self.engine.stop()
                return


    def thrad_run(self):
        t1=threading.Thread(target=self.text_speck,args=('开始',))#args=(参数,)后边要加逗号
        t1.start()
    #     self.task_id_list=[]
    #     # self.window.textEdit.setPlainText('')
    #     # self.window.textEdit.append(self.task_id_list[0])
    #     info = self.window.textEdit.toPlainText()  # 通过 toPlainText 方法获取编辑框内的文本内容
    #     url = "https://aip.baidubce.com/oauth/2.0/token"
    #     params = {"grant_type": "client_credentials", "client_id": self.API_KEY, "client_secret": self.SECRET_KEY}
    #     self.get_access_token= str(requests.post(url, params=params).json().get("access_token"))
    #     url_token = "https://aip.baidubce.com/rpc/2.0/tts/v1/create?access_token=" + self.get_access_token
    #     payload = json.dumps({
    #         "text": info,  # 待合成的文本
    #         "format": "wav",  # 音频格式
    #         "voice": 106,  # 音库
    #         "lang": "zh",  # 语言，固定zh
    #         "speed": 5,  # 语速
    #         "pitch": 5,  # 音调
    #         "volume": 5,  # 音量
    #         "enable_subtitle": 2,  # 是否开启字幕时间戳，取值范围0, 1, 2
    #         "break": 5000  # 段落间隔
    #     })
    #     headers = {
    #         'Content-Type': 'application/json',
    #         'Accept': 'application/json'
    #     }
    #     response = requests.request("POST", url_token, headers=headers, data=payload)
    #     self.python_data = json.loads(response.text)  # json转python
    #     # self.task_id_list.append(self.python_data['task_id'])
    #     # self.window.textEdit.append(self.python_data['task_id'])
    #     self.window.lineEdit.setText(self.python_data['task_id'])
    #
    def save_voice(self):
        files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存",'未命名', '*.mp3')
        print(files_address)
        info = self.window.textEdit.toPlainText()  # 通过 toPlainText 方法获取编辑框内的文本内容
        if info!='':
            engine = pyttsx4.init()  # 初始化语音引擎
            engine.setProperty('rate', 170)  # 设置语速
            volume = engine.getProperty('volume')
            engine.setProperty('volume',str(volume))
            voices = engine.getProperty('voices')
            engine.setProperty('voice', voices[0].id)  # 设置第一个语音合成器
            engine.save_to_file(info,files_address)

            engine.say('文本转语音完成')
            engine.runAndWait()
            engine.stop()
        # info=self.window.lineEdit.text()
        # url = "https://aip.baidubce.com/oauth/2.0/token"
        # params = {"grant_type": "client_credentials", "client_id": self.API_KEY, "client_secret": self.SECRET_KEY}
        # get_access_token= str(requests.post(url, params=params).json().get("access_token"))
        # url = "https://aip.baidubce.com/rpc/2.0/tts/v1/query?access_token=" + get_access_token
        # payload = json.dumps({
        #     "task_ids": [
        #         info  # create获取的task_id
        #     ]
        # })
        # print(payload)
        # headers = {
        #     'Content-Type': 'application/json',
        #     'Accept': 'application/json'
        # }
        #
        # response = requests.request("POST", url, headers=headers, data=payload)
        #
        # print(response.text)
    def thread_stop(self):
        t2=threading.Thread(target=self.text_speck,args=('停止',))#args=(参数,)后边要加逗号
        t2.start()

class image_text_window():#通用文字识别窗口
    def __init__(self):
        super().__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "image_text_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('图片文字识别窗口(限时免费)')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.6)
        self.width = int(self.screenwidth * 0.6)
        self.window.resize(self.width, self.height)
        self.window.pushButton_images.clicked.connect(self.image_files)
        self.window.pushButton_images_text.clicked.connect(self.images_text)
        self.window.pushButton_images_table.clicked.connect(self.images_table)
        self.window.pushButton_save_word.clicked.connect(self.save_word)
        self.window.tableWidget.setRowCount(1000)
        self.window.tableWidget.setColumnCount(100)
        self.rows = self.window.tableWidget.rowCount()  # 获取所有行数
        self.cols =self.window.tableWidget.columnCount()  # 获取所有列数
        for row in range(0,self.rows):
            for column in range(0,self.cols):
                items = QTableWidgetItem('')
                self.window.tableWidget.setItem(row, column, items)
        self.window.pushButton_create_excel.clicked.connect(self.create_excel)
        self.font = Font(
            name="宋体",  # 字体linkActivated
            size=9,  # 字体大小
            color="000000",  # 字体颜色，用16进制rgb表示
            bold=False,  # 是否加粗，True/False
            italic=False,  # 是否斜体，True/False
            strike=None,  # 是否使用删除线，True/False
            underline=None, )  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
        self.alig = Alignment()
        self.alig.wrap_text = True  # 自动换行
        self.alig.vertical = 'center'  # 垂直方向居中对齐
        # 设置边框样式
        # l_side =openpyxl.styles.borders.Side(style='dashDot', color=None)
        # r_side = openpyxl.styles.borders.Side(style='dashDotDot', color=None)
        # t_side = openpyxl.styles.borders.Side(style='dashed', color=None)
        # b_side = openpyxl.styles.borders.Side(style='mediumDashDot', color=None)
        # style  == 'dashDot'左四划线, 'dashDotDot'左三划线, 'dashed'左六划线, 'mediumDashDot'右三划线, 'double'右七双划线, 'slantDashDot',右二划线 'thin'左六实线, 'hair',左一划线 'dotted',左二划线
        # 'thick',右六粗线 'mediumDashed',右四点划线 'mediumDashDotDot'右一点划线, 'medium'右五中粗实线
        l_side = Side(style='thin', color=None)
        r_side = Side(style='thin', color=None)
        t_side = Side(style='thin', color=None)
        b_side = Side(style='thin', color=None)
        self.border = Border(left=l_side, right=r_side, top=t_side, bottom=b_side, vertical=l_side)

    image_list=[]
    def image_files(self):
        self.image_list=[]
        self.import_filePath, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '所有图片(*);;(*.png);;(*.ico);;(*.jpg);;(*.jpeg);;(*.bmp);;(*.gif)')
        if self.import_filePath != []:
            self.image_list+=(self.import_filePath)
            QMessageBox.information(self.window, '提示', '图片加载完成，可以识别了')
        else:
            return
    def images_text(self):
        try:
            if self.image_list!=[]:
                BASE_DIR = os.path.dirname(__file__)
                self.progressBar = QProgressBar(self.window)
                self.progressBar.setValue(0)
                self.progressBar.resize(300, 30)
                self.progressBar.move(400, 400)
                self.progressBar.setRange(0, len(self.image_list))
                self.progressBar.show()
                files = os.path.join(BASE_DIR, "image.ask")
                self.window.textEdit.setPlainText('')
                rows = self.window.tableWidget.rowCount()  # 获取所有行数
                cols = self.window.tableWidget.columnCount()  # 获取所有列数
                for row in range(0, rows):
                    for column in range(0, cols):
                        items = QTableWidgetItem('')
                        self.window.tableWidget.setItem(row, column, items)
                time.sleep(1)
                with open(files, 'r') as file:
                    json_data = json.load(file)
                    python_data=json.loads(json_data)
                    url = "https://aip.baidubce.com/oauth/2.0/token"
                    params = {"grant_type": "client_credentials", "client_id":str(python_data[0]['A_K']),"client_secret":str(python_data[0]['S_K'])}
                    access_token = str(requests.post(url, params=params).json().get("access_token"))
                    request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"
                    column=0
                    for image_file in self.image_list:
                        time.sleep(1)
                        if '.png' in image_file or '.jpg' in image_file or '.jpeg' in image_file or '.bmp' in image_file or '.gif' in image_file or '.ico' in image_file:
                            column+=1
                            f = open(image_file, 'rb')
                            img = base64.b64encode(f.read())
                            params = {"image": img}
                            access_token = access_token
                            request_url =request_url+ "?access_token=" + str(access_token)
                            headers = {'content-type': 'application/x-www-form-urlencoded'}
                            response = requests.post(request_url, data=params, headers=headers)
                            if response:
                                data = (response.json())
                                row=0
                                for i in (data['words_result']):
                                    conent = i['words']
                                    row+=1
                                    print(row-1, column-1, conent)
                                    self.window.textEdit.append(str(conent))
                                    # self.window.textEdit.insertPlainText(str(conent))
                                    item = QTableWidgetItem(str(conent))
                                    self.window.tableWidget.setItem(row-1, column-1, item)
                                self.progressBar.setValue(column)
                    self.progressBar.close()
        except:
            pass
    def create_excel(self):
        try:
            files_address = QFileDialog.getExistingDirectory(self.window, "选择保存文件夹")
            if '/' in files_address:
                self.rows = []
                self.column = []
                mytable = self.window.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值
                    if r.row() not in self.rows:
                        self.rows.append(r.row())
                    if r.column() not in self.column:
                        self.column.append(r.column())
                self.rows.sort(reverse=False)
                self.column.sort(reverse=False)
                if self.column!=[] and self.rows!=[]:
                    for self.Single_colum in self.column:
                        text_list = []
                        for self.Single_rows in self.rows:
                            text = self.window.tableWidget.item(self.Single_rows,self.Single_colum).text()
                            text=str(text).replace(' ','')
                            if str(text)!='':
                                text_list.append(str(text))

                        if text_list != [] and len(text_list)>=2:
                            file_path = os.path.join(files_address, '~$' + str(text_list[0]) + '.xlsx')
                            if os.path.exists(file_path):
                                QMessageBox.information(self.window, '提示', '请先关闭打开的同名称Excel')
                                break
                            nw=openpyxl.Workbook()
                            for sheet_name in text_list[1:]:
                                nw.create_sheet(sheet_name)
                            nw.remove(nw['Sheet'])
                            nw.save(os.path.join(files_address, str(text_list[0])+'.xlsx')) # 保存新建的文件
                        if text_list != [] and len(text_list)==1:
                            file_path = os.path.join(files_address, '~$' + str(text_list[0]) + '.xlsx')
                            if os.path.exists(file_path):
                                QMessageBox.information(self.window, '提示', '同名称的Excel打开了，请先关闭文件。')
                                break
                            nw=openpyxl.Workbook()
                            for sheet_name in text_list[1:]:
                                nw.create_sheet(sheet_name)
                            nw.save(os.path.join(files_address, str(text_list[0])+'.xlsx')) # 保存新建的文件
                    QMessageBox.information(self.window, '温馨提示', '创建Excel完成，请检查。')
                if self.column == [] and self.rows == []:
                    QMessageBox.information(self.window, '温馨提示', '请选中表格内文字。')
        except:
            pass
    def save_word(self):
        doc_word = []
        self.info=self.window.textEdit.toPlainText()#通过 toPlainText 方法获取编辑框内的文本内容
        if self.info!='':
            doc_word.append(self.info)
            files_address = QFileDialog.getExistingDirectory(self.window, "选择保存文件夹")
            if '/' in files_address:
                t = time.localtime()  # 获取当前本地时间
                strtime = time.strftime("%Y-%m-%d %H-%M-%S", t)
                doc = Document()# 创建Word
                doc.save(files_address + '/' + '文字识别{}.docx'.format(str(strtime)))
                doc = Document(files_address + '/' + '文字识别{}.docx'.format(str(strtime)))#读取Word
                doc.add_paragraph(doc_word)# 写入内容
                doc.save(files_address + '/' + '文字识别{}.docx'.format(str(strtime)))
                QMessageBox.information(self.window, '温馨提示', '图片文字保存完成，请检查。')
            else:
                return
    def images_table(self):
        if self.image_list!=[]:
            files_address=QFileDialog.getExistingDirectory(self.window, "选择保存文件夹")
            if '/' in files_address:

                BASE_DIR = os.path.dirname(__file__)
                files = os.path.join(BASE_DIR, "image.ask")
                with open(files, 'r') as file:
                    json_data = json.load(file)
                    python_data = json.loads(json_data)
                    url = "https://aip.baidubce.com/oauth/2.0/token"
                    params = {"grant_type": "client_credentials", "client_id":str(python_data[0]['A_K']),"client_secret":str(python_data[0]['S_K'])}
                    access_token = str(requests.post(url, params=params).json().get("access_token"))
                    num=0
                    for image_file in self.image_list:
                        time.sleep(1)
                        if '.png' in image_file or '.jpg' in image_file or '.jpeg' in image_file or '.bmp' in image_file or '.gif' in image_file or '.ico' in image_file:
                            t = time.localtime()  # 获取当前本地时间
                            strtime = time.strftime("%Y-%m-%d-%H-%M-%S", t)
                            nw = openpyxl.Workbook()  # 新建文件
                            num+=1
                            f = open(image_file, 'rb')
                            request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/table"
                            img = base64.b64encode(f.read())
                            params = {"image": img}
                            access_token = access_token
                            request_url = request_url + "?access_token=" +str(access_token)
                            headers = {'content-type': 'application/x-www-form-urlencoded'}
                            response = requests.post(request_url, data=params, headers=headers)
                            if response:
                                print(response.json()["tables_result"][0]['body'])
                                d=response.json()["tables_result"][0]['body']
                                for key in d:
                                    row=key['row_start']
                                    col=key['col_start']
                                    self.text=key['words']
                                    nw.active.cell(int(row)+1,int(col)+1,
                                                   str(self.text)).alignment = self.alig
                                    nw.active.cell(int(row)+1,int(col)+1,
                                                   str(self.text)).font = self.font
                                    nw.active.cell(int(row)+1,int(col)+1,
                                                   str(self.text)).border = self.border

                                nw.save(os.path.join(files_address,'第'+str(num)+'张'+'{}.xlsx'.format(strtime))) # 保存新建的文件
                    QMessageBox.information(self.window, '温馨提示', '表格文字识别完成，请检查。')
            else:
                return
class QDK_Window():#加载清单库窗口
    def __init__(self):
        super().__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "Manifest_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('数据库窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.6)
        self.width = int(self.screenwidth * 0.6)
        self.window.resize(self.width, self.height)
        self.window.treeWidget_qdk.setColumnWidth(0, 600)#列宽
        self.window.treeWidget_qdc.setColumnWidth(0, 600)#列宽
        # self.window.treeWidget_qdk.setStyleSheet("QHeaderView::section{background:rgb(196, 223, 255);}")  # 列表头颜色
        # self.window.treeWidget_qdc.setStyleSheet("QHeaderView::section{background:rgb(196, 223, 255);}")  # 列表头颜色
        self.window.splitter.setStretchFactor(0, 40)
        self.window.splitter.setStretchFactor(1, 60)
        self.window.splitter_2.setStretchFactor(0, 40)
        self.window.splitter_2.setStretchFactor(1, 60)
        self.window.splitter_3.setStretchFactor(0, 40)
        self.window.splitter_3.setStretchFactor(1, 60)
        self.window.treeWidget_qdc.itemClicked.connect(self.treeWidget_qdc_Clicked)#清单池
        self.window.treeWidget_qdk.itemClicked.connect(self.treeWidget_qdk_Clicked)# 清单库
        self.window.treeWidget_quota.itemClicked.connect(self.treeWidget_quota_Clicked)  # 定额库
        # self.window.tableWidget_qdc.setSelectionBehavior(QAbstractItemView.SelectRows)# 选择整行
        # self.window.tableWidget_qdk.setSelectionBehavior(QAbstractItemView.SelectRows)# 选择整行
        self.window.pushButton_c_insert_line.clicked.connect(self.c_insert_line)#清单池插入行
        self.window.pushButton_k_insert_line.clicked.connect(self.k_insert_line)#清单库插入行
        self.window.pushButton_intert_quota.clicked.connect(self.quota_insert_line)#定额库插入行

        self.window.pushButton_c_delrow.clicked.connect(self.c_delrow)#清单池删除行
        self.window.pushButton_k_delrow.clicked.connect(self.k_delrow)#清单库删除行
        self.window.pushButton_del_quota.clicked.connect(self.quota_delrow)  # 定额库删除行
        self.window.pushButton_write_qdk.clicked.connect(self.write_qdk)#写入清单库
        self.window.pushButton_k_cut.clicked.connect(self.k_cut)#剪切行
        self.window.pushButton_k_copy.clicked.connect(self.k_copy)#复制行
        self.window.pushButton_k_stickup_rows.clicked.connect(self.k_stickup_rows)#粘贴行
        self.window.pushButton_c_copy.clicked.connect(self.c_copy)  # 复制行
        self.window.pushButton_c_cut.clicked.connect(self.c_cut)#剪切行
        self.window.pushButton_c_stickup_rows.clicked.connect(self.c_stickup_rows)#粘贴行
        self.window.pushButton_save_qdc.clicked.connect(self.save_qdc)  # 保存清单池
        self.window.pushButton_save_qdk.clicked.connect(self.save_qdk)  # 保存清单库
        self.window.pushButton_save_quota.clicked.connect(self.save_quota)  # 保存定额库
        self.window.pushButton_Batch_quota.clicked.connect(self.batch_quota)  # 批量套定额
        self.window.tableWidget_qdk.itemDoubleClicked.connect(self.Double_click) # 双击清单库
        self.window.tableWidget_quota.itemDoubleClicked.connect(self.quota_Double_click)  # 双击定额库
        self.window.pushButton_parse_quota.clicked.connect(self.parse_quota)  # 解析窗
        self.window.dockWidget.close()
        self.window.pushButton_quota_specification.clicked.connect(self.quota_specification)  # 解析规格提取

        self.window.tableWidget_qdk.itemClicked.connect(self.startThread_run_features)#单击
        self.window.tableWidget_features.horizontalHeader().resizeSection(0, 280)  # 调整第2列的大小为500像素
        self.window.tableWidget_features.itemDoubleClicked.connect(self.Double_features)  # 双击项目特征表格
        self.window.lineEdit_Search_qdk.setPlaceholderText('查找清单')  # 提示文本
        self.window.lineEdit_Search_qdc.setPlaceholderText('查找清单')  # 提示文本
        self.window.toolButton_Search_qdk.clicked.connect(self.startThread_run_Search_qdk)  # 搜索清单库
        self.window.lineEdit_Search_qdk.returnPressed.connect(self.startThread_run_Search_qdk)  # 绑定enter键
        self.window.toolButton_Search_qdc.clicked.connect(self.startThread_run_Search_qdc)  # 搜索清单池
        self.window.lineEdit_Search_qdc.returnPressed.connect(self.startThread_run_Search_qdc)  # 绑定enter键
        self.window.toolButton_Search_quota.clicked.connect(self.startThread_run_Search_quota)  # 搜索定额库
        self.window.lineEdit_Search_quota.returnPressed.connect(self.startThread_run_Search_quota)  # 绑定enter键

        self.clipboard_qd = QApplication.clipboard()
        self.clipboard_qd.dataChanged.connect(self.clipboard)# 剪切板内容发生变化连接函数
        self.window.pushButton_c_copy_clipboard.clicked.connect(self.c_copy_clipboard)  # 复制外部数据
        self.window.pushButton_c_copy_clipboard.setShortcut('ctrl+V')

        self.window.pushButton_k_copy_clipboard.clicked.connect(self.k_copy_clipboard)  # 复制外部数据
        self.window.pushButton_k_copy_clipboard.setShortcut('ctrl+V')
        self.window.pushButton_copy_quota.clicked.connect(self.quota_copy_clipboard)  # 复制外部数据
        self.window.pushButton_copy_quota.setShortcut('ctrl+V')

        # self.window.setWindowModality(Qt.ApplicationModal)#阻塞主窗口不能点击WindowModal
        self.font = Font(
            name="宋体",  # 字体
            size=9,  # 字体大小
            color="000000",  # 字体颜色，用16进制rgb表示
            bold=False,  # 是否加粗，True/False
            italic=False,  # 是否斜体，True/False
            strike=None,  # 是否使用删除线，True/False
            underline=None, )  # 下划线, 可选'singleAccounting', 'double', 'single', 'doubleAccounting'
        self.alig = Alignment()
        self.alig.wrap_text = True
        self.alig.vertical = 'center'
        l_side = Side(style='thin', color=None)
        r_side = Side(style='thin', color=None)
        t_side = Side(style='thin', color=None)
        b_side = Side(style='thin', color=None)
        self.border = Border(left=l_side, right=r_side, top=t_side, bottom=b_side, vertical=l_side)
        # 加载清单池树界面
        # BASE_DIR = os.path.dirname(__file__)  # 清单池
        # files_address = BASE_DIR + '/' + "清单数据库" + '/' + "清单池" + '/'
        # self.tree_font = QFont()
        # self.tree_font.setPointSize(9)  # 设置字体大小为9像素
        # self.tree_font.setFamily("宋体")
        # self.window.treeWidget_qdc.expandAll()
        # for file_name in os.listdir(files_address):
        #     self.root = QTreeWidgetItem(self.window.treeWidget_qdc)
        #     self.root.setFlags(
        #         QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
        #     self.root.setText(0, file_name)
        #     self.root.setSizeHint(0, QSize(0, 25))
        #     self.root.setFont(0, self.tree_font)
        #     files = files_address + file_name
        #     for file in os.listdir(files):
        #         if '~$' not in file and '.xlsx' in file:
        #             self.son = QTreeWidgetItem(self.root)
        #             self.son.setText(0, str(file.split('.')[0]))
        #             self.son.setSizeHint(0, QSize(0, 25))
        #             self.son.setFont(0, self.tree_font)
        #             f = files + '/' + file
        #             self.wb = openpyxl.load_workbook(f, read_only=False, data_only=True, keep_links=False)
        #             for i in range(len(self.wb.sheetnames)):
        #                 self.grandson = QTreeWidgetItem(self.son)
        #                 self.grandson.setText(0, self.wb.sheetnames[i])
        #                 self.grandson.setSizeHint(0, QSize(0, 25))
        #                 self.grandson.setFont(0, self.tree_font)

        comboBox_list = []  # 清单库
        BASE_DIR = os.path.dirname(__file__)
        self.files_qdk = BASE_DIR + '/' + "清单数据库" + '/' + "清单库" + '/'
        for file_name in os.listdir(self.files_qdk):
            comboBox_list.append(file_name)
        self.window.comboBox.addItems(comboBox_list)
        self.window.comboBox.currentIndexChanged.connect(self.handleSelectionChange)

        # 多线程加载清单池
        self.worker = Thread_load_qdc()
        self.worker.stopSing.connect(self.stopThread_load_qdc)
        self.worker.start()
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint)  # 窗口始终在前面
    #     加载定额库

        comboBox2_list = []  # 定额库
        BASE_DIR = os.path.dirname(__file__)
        self.files_dek = BASE_DIR + '/' +"定额库" + '/'
        for file_name in os.listdir(self.files_dek):
            comboBox2_list.append(file_name)
        self.window.comboBox_2.addItems(comboBox2_list)
        self.window.comboBox_2.activated.connect(self.region_dek_Change)

    def load_qdc(self):
        BASE_DIR = os.path.dirname(__file__) # 清单池
        files_address=BASE_DIR+'/'+"清单数据库"+'/'+"清单池" + '/'
        self.tree_font = QFont()
        self.tree_font.setPointSize(9)  # 设置字体大小为9像素
        self.tree_font.setFamily("宋体")
        self.window.treeWidget_qdc.expandAll()
        for file_name in os.listdir(files_address):
            self.root = QTreeWidgetItem(self.window.treeWidget_qdc)
            self.root.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            self.root.setText(0, file_name)
            self.root.setSizeHint(0, QSize(0, 25))
            self.root.setFont(0, self.tree_font)
            files = files_address +file_name
            for file in os.listdir(files):
                if '~$' not in file and 'xlsx' == file.split('.')[-1]:
                    self.son = QTreeWidgetItem(self.root)
                    self.son.setText(0,str(file.split('.')[0]))
                    self.son.setSizeHint(0, QSize(0, 25))
                    self.son.setFont(0, self.tree_font)
                    f=files + '/' + file
                    self.wb = openpyxl.load_workbook(f, read_only=False, data_only=True, keep_links=False)
                    for i in range(len(self.wb.sheetnames)):
                        self.grandson = QTreeWidgetItem(self.son)
                        self.grandson.setText(0, self.wb.sheetnames[i])
                        self.grandson.setSizeHint(0, QSize(0, 25))
                        self.grandson.setFont(0, self.tree_font)
    def stopThread_load_qdc(self,all_qdc_list):
        self.all_qdc_list=all_qdc_list
        self.worker.quit()  # 退出
        self.worker.wait()# 回收资源
        self.worker.deleteLater()
        self.load_qdc()


    qdc_text_list=[]
    qdc_save='不保存'
    def treeWidget_qdc_Clicked(self):#点击树窗口显示清单池
        self.qdc_text_list = []
        self.root_num = self.window.treeWidget_qdc.topLevelItemCount()  # 获取根节点数量.currentIndex().row()
        self.item = self.window.treeWidget_qdc.currentItem()
        # self.save_qdk(self.item)
        # for i in range(0, self.root_num):
        #     item = self.window.treeWidget_qdc.topLevelItem(i)  # 循环获取根节点
        #     text = item.text(0)  # 根节点文字信息（默认一列）
        #     if self.item.text(0) == text:
        self.window.tableWidget_qdc.setRowCount(0)  # 设置行数
        BASE_DIR = os.path.dirname(__file__)
        files_address = BASE_DIR + '/' + "清单数据库" + '/' + "清单池"
        if self.item.parent():
            self.second = self.item.parent()
            if self.second.parent():
                self.window.pushButton_save_qdc.setEnabled(True)
        #         self.fist = self.second.parent()
        #         # print(self.second.text(0),self.fist.text(0),self.item.text(0))
        #         self.bin_path_qdc = os.path.join(files_address, self.fist.text(0), self.second.text(0) + '.xlsx')
        #         self.file_path_qdc = os.path.join(files_address, self.fist.text(0),'~$' + self.second.text(0) + '.xlsx')

        # for dirpath, dirnames, filenames in os.walk(files_address):
        #     for filename in filenames:
        #         files = os.path.join(dirpath, filename)
                files = os.path.join(files_address, self.second.parent().text(0), self.second.text(0) + '.xlsx')
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    # print(filename)
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    self.num = 0
                    for i in range(len(self.wb.sheetnames)):
                        if self.item.text(0) == self.wb.sheetnames[i]:  # or self.item.text(0)==str(filename).split('.')[0]
                            # self.num = 0
                            self.ws = self.wb[self.wb.sheetnames[i]]
                            # print(self.ws)
                            self.values = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,values_only=True))
                            max_num = len(self.values)
                            # print(max_num)
                            for wb_ws_row in range(0, max_num):
                                data = self.values[wb_ws_row]
                                self.num += 1
                                # print(self.num,data)
                                self.window.tableWidget_qdc.setRowCount(self.num)  # 设置行数
                                self.qdc_text = QPlainTextEdit()
                                self.qdc_text_list.append(self.qdc_text)
                                self.window.tableWidget_qdc.setItem(int(self.num) - 1, 0,QTableWidgetItem(''))
                                self.window.tableWidget_qdc.setItem(int(self.num) - 1, 1, QTableWidgetItem(''))
                                self.window.tableWidget_qdc.setItem(int(self.num) - 1, 2, QTableWidgetItem(''))
                                self.window.tableWidget_qdc.setCellWidget(int(self.num) - 1, 3, self.qdc_text)
                                self.window.tableWidget_qdc.setItem(int(self.num) - 1, 4, QTableWidgetItem(''))

                                for j in range(0, len(data)):
                                    self.item1 = str(data[j]).replace('None', '')
                                    if j == 3:
                                        self.qdc_text_list[int(self.num) - 1].setPlainText(self.item1)
                                    if j != 3:
                                        self.window.tableWidget_qdc.setItem(int(self.num) - 1, int(j),
                                                                            QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdc.verticalHeader().resizeSection(int(self.num) - 1, 100)  # 调整每一行的大小为100像素
                                    # self.window.tableWidget_qdc.verticalHeader().setSectionResizeMode(
                                    #     int(self.num) - 1, QHeaderView.ResizeToContents)
                            self.window.tableWidget_qdc.setStyleSheet("gridline-color: rgb(257, 1, 0)")
                            # self.window.tableWidget_qdc.horizontalHeader().setStyleSheet(
                            #     "color: rgb(0, 83, 128);border:1px solid rgb(210, 210, 210);")
                            self.window.tableWidget_qdc.horizontalHeader().setSectionResizeMode(
                                QHeaderView.Interactive)  # 设置列宽，列宽可调
                            self.window.tableWidget_qdc.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
                            self.window.tableWidget_qdc.horizontalHeader().resizeSection(4, 100)  # 调整第2列的大小为500像素
                            # it=self.window.tableWidget_qdc.item(12,3)
                            # self.window.tableWidget_qdc.scrollToItem(it)# 滚轮定位

    def save_qdc(self):#保存清单池
        BASE_DIR = os.path.dirname(__file__)
        files_address = BASE_DIR + '/' + "清单数据库" + '/' + "清单池"
        row_c = self.window.treeWidget_qdc.currentIndex().row()  # 获取行
        if int(row_c) != -1:
            self.item = self.window.treeWidget_qdc.currentItem()
            if self.item.parent():
                self.second = self.item.parent()
                if self.second.parent():
                    self.fist = self.second.parent()
                    # print(self.second.text(0),self.fist.text(0),self.item.text(0))
                    self.bin_path_qdc = os.path.join(files_address, self.fist.text(0), self.second.text(0) + '.xlsx')
                    self.file_path_qdc = os.path.join(files_address, self.fist.text(0), '~$' + self.second.text(0) + '.xlsx')
                    print(self.bin_path_qdc)
                    if os.path.exists(self.file_path_qdc):
                        QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')
                    else:
                        self.wb = openpyxl.load_workbook(self.bin_path_qdc, read_only=False, data_only=True, keep_links=False)
                        self.ws=self.wb[self.item.text(0)]
                        self.rows = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        self.table_rows = self.window.tableWidget_qdc.rowCount()
                        self.table_column = self.window.tableWidget_qdc.columnCount()
                        # print(self.ws.max_row)
                        self.ws.delete_rows(1, self.ws.max_row+1)  # 删除行，前边是起始行，后边是行数
                        for i in range(0, int(self.table_rows)):
                            for j in range(0, int(self.table_column)):
                                if j!=3:
                                    self.text = self.window.tableWidget_qdc.item(i, j).text()  # 获取单元格内容
                                    self.ws.cell(i + 1, j + 1, str(self.text)).alignment = self.alig
                                    self.ws.cell(i + 1, j + 1, str(self.text)).font = self.font
                                if j==3:
                                    text3=str(self.qdc_text_list[i].toPlainText())
                                    self.ws.cell(i + 1, j + 1, str(text3)).alignment = self.alig
                                    self.ws.cell(i + 1, j + 1, str(text3)).font = self.font
                                print(i,j, str(self.text))
                        if os.path.exists(self.file_path_qdc):
                            QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件') # 防止在写入时打开文件，软件崩溃。
                        else:
                            self.wb.save(self.bin_path_qdc)  # 保存新建的文件


    def write_qdk(self):#写入清单库
        self.Item_row = self.window.treeWidget_qdc.currentIndex().row()  # 获取行
        if int(self.Item_row)!=-1:
            self.Item_rowk = self.window.treeWidget_qdk.currentIndex().row()  # 获取行
            if int(self.Item_rowk) != -1:
                self.item = self.window.treeWidget_qdc.currentItem()
                BASE_DIR = os.path.dirname(__file__)
                files_address = BASE_DIR + '/' + "清单数据库" + '/' + "清单池"
                if self.item.parent():
                    self.second = self.item.parent()
                    if self.second.parent():
                        self.fist = self.second.parent()
                        # print(self.second.text(0), self.fist.text(0), self.item.text(0))
                        self.bin_path = os.path.join(files_address, self.fist.text(0), self.second.text(0) + '.xlsx')
                        self.file_path = os.path.join(files_address, self.fist.text(0), '~$' + self.second.text(0) + '.xlsx')
                        #清单库
                        self.itemk = self.window.treeWidget_qdk.currentItem()
                        if self.itemk.parent():
                            self.secondk = self.itemk.parent()
                            if self.secondk.parent():
                                self.fistk = self.second.parent()
                                # print(self.second.text(0),self.fist.text(0),self.item.text(0))
                                self.bin_pathk = os.path.join(self.files, self.fistk.text(0),self.secondk.text(0) + '.xlsx')
                                self.file_pathk = os.path.join(self.files, self.fistk.text(0), '~$' + self.secondk.text(0) + '.xlsx')
                                self.tableWidget_allrows = self.window.tableWidget_qdc.rowCount()  # 获取总行数
                                if self.tableWidget_allrows != 0:
                                    items = ["应用当前分部", "应用当前工程专业"]
                                    com_text, ok = QInputDialog().getItem(self.window, "选择应用范围",
                                                                          "请选择写入清单库的范围", items, 0, True)
                                    if ok:
                                        print(com_text)
                                        if com_text == '应用当前分部':
                                            # 清单池
                                            self.wbc = openpyxl.load_workbook(self.bin_path, read_only=False, data_only=True, keep_links=False)
                                            self.wbk = openpyxl.load_workbook(self.bin_pathk, read_only=False,data_only=True, keep_links=False)
                                            for c in range(len(self.wbc.sheetnames)):
                                                if self.item.text(0) == self.wbc.sheetnames[c]:
                                                    self.wsc = self.wbc[self.wbc.sheetnames[c]]
                                                    self.values_qdc = list(self.wsc.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,values_only=True))
                                                    for value_qdc in self.values_qdc:
                                            #            # 清单库
                                                        for k in range(len(self.wbk.sheetnames)):
                                                            self.num = 0
                                                            self.wsk = self.wbk[self.wbk.sheetnames[k]]
                                                            self.values_qdk = list(self.wsk.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None,values_only=True))
                                                            for n, value_qdk in enumerate(self.values_qdk):
                                                                if value_qdc[2] in value_qdk and value_qdc[4] in value_qdk:
                                                                    print(n+1,value_qdc[0],value_qdc[3])
                                                                    self.wsk.cell(n+1, 4,value_qdc[3]).alignment = self.alig
                                                                    self.wsk.cell(n+1, 4,value_qdc[3]).font = self.font
                                                                    self.wsk.cell(n+1, 4,value_qdc[3]).border = self.border
                                                                    self.wsk.cell(n+1, 1,value_qdc[0]).alignment = self.alig
                                                                    self.wsk.cell(n+1, 1,value_qdc[0]).font = self.font
                                                                    self.wsk.cell(n+1, 1,value_qdc[0]).border = self.border
                                            self.wbk.save(self.bin_pathk)
                                        if com_text =='应用当前工程专业':
                                            self.wbc = openpyxl.load_workbook(self.bin_path, read_only=False,data_only=True, keep_links=False)
                                            self.wbk = openpyxl.load_workbook(self.bin_pathk, read_only=False,data_only=True, keep_links=False)
                                            for c in range(len(self.wbc.sheetnames)):
                                                self.wsc = self.wbc[self.wbc.sheetnames[c]]
                                                self.values_qdc = list(self.wsc.iter_rows(min_row=None, max_row=None, min_col=None,max_col=None, values_only=True))
                                                for value_qdc in self.values_qdc:
                                                    # 清单池
                                                    for k in range(len(self.wbk.sheetnames)):
                                                        self.num = 0
                                                        self.wsk = self.wbk[self.wbk.sheetnames[k]]
                                                        self.values_qdk = list(
                                                            self.wsk.iter_rows(min_row=None, max_row=None,min_col=None, max_col=None,values_only=True))
                                                        for n, value_qdk in enumerate(self.values_qdk):
                                                            if value_qdc[2] in value_qdk and value_qdc[4] in value_qdk:
                                                                self.wsk.cell(n + 1, 4,value_qdc[3]).alignment = self.alig
                                                                self.wsk.cell(n + 1, 4, value_qdc[3]).font = self.font
                                                                self.wsk.cell(n + 1, 4,value_qdc[3]).border = self.border
                                                                self.wsk.cell(n + 1, 1, value_qdc[0]).alignment = self.alig
                                                                self.wsk.cell(n + 1, 1,value_qdc[0]).font = self.font
                                                                self.wsk.cell(n + 1, 1,value_qdc[0]).border = self.border
                                            self.wbk.save(self.bin_pathk)
    def region_dek_Change(self):#加载定额库
        city_name = self.window.comboBox_2.currentText()  # 获取当前选中的选项的文本
        self.files_de = self.files_dek + city_name
        self.window.treeWidget_quota.clear()
        for file_name in os.listdir(self.files_de):  # 文件夹作为一级目录
            self.root = QTreeWidgetItem(self.window.treeWidget_quota)
            self.root.setFlags(
                QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            self.root.setText(0, file_name)
            self.root.setSizeHint(0, QSize(0, 25))
            self.files_son = os.path.join(self.files_de, file_name)
            for file in os.listdir(self.files_son):  # excel名称作为二级目录
                if '~$' not in file and 'xlsx' == file.split('.')[-1]:
                    self.son = QTreeWidgetItem(self.root)
                    self.son.setText(0, str(file.split('.')[0]))
                    self.son.setSizeHint(0, QSize(0, 25))
                    f = os.path.join(self.files_son, file)
                    self.wb = openpyxl.load_workbook(f, read_only=False, data_only=True, keep_links=False)
                    for i in range(len(self.wb.sheetnames)):  # excel里的sheet作为三级目录
                        self.grandson = QTreeWidgetItem(self.son)
                        self.grandson.setText(0, self.wb.sheetnames[i])
                        self.grandson.setSizeHint(0, QSize(0, 25))


    def treeWidget_quota_Clicked(self):#单击定额库树界面显示定额
        # self.root_num = self.window.treeWidget_quota.topLevelItemCount()  # 获取根节点数量.currentIndex().row()
        self.item=self.window.treeWidget_quota.currentItem()
        self.window.tableWidget_quota.clearContents()
        self.window.tableWidget_quota.setRowCount(0)  # 设置行数
        if self.item.parent():
            self.second=self.item.parent()
            if self.second.parent():
                files_de = os.path.join(self.files_de,self.second.parent().text(0),self.item.parent().text(0)+'.xlsx')
                print(files_de)
                # 多线程加载清单池
                self.worker = Thread_load_dek(str(self.item.text(0)),files_de)
                self.worker.stopSing.connect(self.stopThread_load_dek)
                self.worker.start()
                self.window.pushButton_save_quota.setEnabled(True)
                self.window.tableWidget_quota.setStyleSheet("gridline-color: rgb(257, 1, 0)")
                self.window.tableWidget_quota.horizontalHeader().setSectionResizeMode(
                    QHeaderView.Interactive)  # 设置列宽，列宽可调
                self.window.tableWidget_quota.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
                # self.window.pushButton_parse_quota.setEnabled(True)
    def stopThread_load_dek(self,msg):
        self.text_list=msg
        max_num = len(self.text_list)
        # print(max_num)
        self.num=0
        for wb_ws_row in range(0, max_num):
            data = self.text_list[wb_ws_row]
            self.num += 1
            # print(self.num,data)
            self.window.tableWidget_quota.setRowCount(self.num)  # 设置行数
            self.window.tableWidget_quota.setItem(int(self.num) - 1, 0, QTableWidgetItem(''))
            self.window.tableWidget_quota.setItem(int(self.num) - 1, 1, QTableWidgetItem(''))
            self.window.tableWidget_quota.setItem(int(self.num) - 1, 2, QTableWidgetItem(''))
            self.window.tableWidget_quota.setItem(int(self.num) - 1, 3, QTableWidgetItem(''))
            self.window.tableWidget_quota.setItem(int(self.num) - 1, 4, QTableWidgetItem(''))
            for j in range(0, len(data)):
                self.item1 = str(data[j]).replace('None', '')
                self.window.tableWidget_quota.setItem(int(self.num) - 1, int(j), QTableWidgetItem(self.item1))
            self.window.tableWidget_quota.verticalHeader().resizeSection(int(self.num) - 1, 50)  # 调整每一行的大小为100像素
        self.worker.quit()  # 退出
        self.worker.wait()# 回收资源
    def save_quota(self):#保存定额库
        try:
            row_quota = self.window.treeWidget_quota.currentIndex().row()  # 获取行
            if int(row_quota) != -1:
                self.item = self.window.treeWidget_quota.currentItem()
                if self.item.parent():
                    self.second = self.item.parent()
                    if self.second.parent():
                        self.fist = self.second.parent()
                        self.bin_path_dek = os.path.join(self.files_de, self.fist.text(0), self.second.text(0) + '.xlsx')
                        self.file_path_dek = os.path.join(self.files_de, self.fist.text(0),'~$' + self.second.text(0) + '.xlsx')

                        if os.path.exists(self.file_path_dek):
                            QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')
                        else:
                            self.wb = openpyxl.load_workbook(self.bin_path_dek, read_only=False, data_only=True,
                                                             keep_links=False)
                            self.ws = self.wb[self.item.text(0)]
                            self.rows = list(
                                self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                            self.table_rows = self.window.tableWidget_quota.rowCount()
                            self.table_column = self.window.tableWidget_quota.columnCount()
                            self.ws.delete_rows(1, self.ws.max_row + 1)  # 删除行，前边是起始行，后边是行数
                            for i in range(0, int(self.table_rows)):
                                for j in range(0, int(self.table_column)):
                                    self.text = self.window.tableWidget_quota.item(i, j).text()  # 获取单元格内容
                                    self.ws.cell(i + 1, j + 1, str(self.text)).alignment = self.alig
                                    self.ws.cell(i + 1, j + 1, str(self.text)).font = self.font
                                    # print(i,j, str(self.text))
                            if os.path.exists(self.file_path_dek):
                                QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')  # 防止在写入时打开文件，软件崩溃。
                            else:
                                self.wb.save(self.bin_path_dek)  # 保存新建的文件
                                win.import_dek('重新加载定额库')

        except:
            pass
    def quota_specification(self):# 解析定额规格
        try:
            # self.Item_row = self.window.treeWidget_quota.currentIndex().row()  # 获取行
            # if int(self.Item_row) != -1:
            row = self.window.tableWidget_quota.currentRow()  # 获取行数
            if row != -1:
                text1 = self.window.comboBox_3.currentText()  # 获取当前选中的选项的文本
                text2 = self.window.comboBox_4.currentText()  # 获取当前选中的选项的文本
                print(text1,text2)
                if text1 != '' and text2 != '':
                    criteria = '{}(.*?){}'.format(text1, text2)
                    compile = re.compile(criteria, re.S)
                    mytable = self.window.tableWidget_quota.selectedItems()
                    self.column_dict = {}
                    for r in mytable:  # 第三种方法获取值
                        if r.column() not in self.column_dict:
                            self.column_dict[r.column()] = ''
                    if len(list(self.column_dict.keys())) == 1 and 1 not in list(self.column_dict.keys()):
                        for r in mytable:  # 第三种方法获取值
                            all_names =self.window.tableWidget_quota.item(r.row(), r.column()).text()  # 获取单元格内容
                            if r.column() == 3:
                                result = compile.findall(all_names)
                                if result != []:
                                    self.window.tableWidget_quota.setItem(r.row(), 1,QTableWidgetItem(str(result[0]).replace('(','').replace(')','').replace(' ','')))

                elif text1 != '' and text2 == '':
                    criteria = '{}(.*)'.format(text1)
                    print(criteria)
                    compile = re.compile(criteria)
                    mytable =self.window.tableWidget_quota.selectedItems()
                    self.column_dict = {}
                    for r in mytable:  # 第三种方法获取值
                        if r.column() not in self.column_dict:
                            self.column_dict[r.column()] = ''
                    if len(list(self.column_dict.keys())) == 1 and 1 not in list(self.column_dict.keys()):
                        for r in mytable:  # 第三种方法获取值
                            all_names = self.window.tableWidget_quota.item(r.row(), r.column()).text()  # 获取单元格内容
                            if r.column() == 3:
                                result = compile.findall(all_names)
                                if result != []:
                                    self.window.tableWidget_quota.setItem(r.row(), 1, QTableWidgetItem(str(result[0]).replace('(','').replace(')','').replace(' ','')))

                elif text1 == '' and text2 != '':
                    criteria = '(.*){}'.format(text2)
                    compile = re.compile(criteria)
                    mytable = self.window.tableWidget_quota.selectedItems()
                    self.column_dict = {}
                    for r in mytable:  # 第三种方法获取值
                        if r.column() not in self.column_dict:
                            self.column_dict[r.column()] = ''
                    if len(list(self.column_dict.keys())) == 1 and 1 not in list(self.column_dict.keys()):
                        for r in mytable:  # 第三种方法获取值
                            all_names = self.window.tableWidget_quota.item(r.row(), r.column()).text()  # 获取单元格内容
                            if r.column() == 3:
                                result = compile.findall(all_names)
                                if result != []:
                                    self.window.tableWidget_quota.setItem(r.row(), 1, QTableWidgetItem(str(result[0]).replace('(','').replace(')','').replace(' ','')))
                elif text1 == '' and text2 == '':
                    criteria = '(.*)'
                    compile = re.compile(criteria, re.S)
                    mytable = self.window.tableWidget_quota.selectedItems()
                    self.column_dict = {}
                    for r in mytable:  # 第三种方法获取值
                        if r.column() not in self.column_dict:
                            self.column_dict[r.column()] = ''
                    if len(list(self.column_dict.keys())) == 1 and 1 not in list(self.column_dict.keys()):
                        for r in mytable:  # 第三种方法获取值
                            all_names =self.window.tableWidget_quota.item(r.row(), r.column()).text()  # 获取单元格内容
                            if r.column() == 3:
                                result = compile.findall(all_names)
                                if result != []:
                                    self.window.tableWidget_quota.setItem(r.row(), 1, QTableWidgetItem(str(result[0])))
                self.window.tableWidget_quota.viewport().update()  # 刷新tab内容
        except:
            pass
    def batch_quota(self):#批量套定额
        try:
            self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:
                # if self.window.treeWidget_quota.currentIndex().row()!=-1:
                quota_row=self.window.tableWidget_quota.currentRow()
                if quota_row!=-1:
                    win.tab_change = '不提取'
                    # 撤销
                    self.old_undo_dict = {}
                    self.Click_list = []
                    self.old_undo_dict[win.tableWidget] = self.Click_list
                    for self.Single_rows in range(0, win.tableWidget.rowCount()):
                        t_list = []
                        self.Click_list.append(t_list)
                        for self.Single_colum in range(0, win.tableWidget.columnCount()):
                            self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
                    # 插入定额
                    mytable = self.window.tableWidget_quota.selectedItems()
                    self.rows_dict = {}
                    values_dict={}
                    for r in mytable:  # 第三种方法获取值
                        if r.row() not in self.rows_dict:
                            self.rows_dict[r.row()] = ''
                    if self.rows_dict!={}:
                        if win.tableWidget.rowCount() != 0:
                            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                                if Item1 == Qt.Checked:
                                    num = 0
                                    text_5 = win.tableWidget.item(self.Single_rows, 5).text()
                                    sys_text = win.tableWidget.item(self.Single_rows, 6).text()
                                    name_text = win.tableWidget.item(self.Single_rows, 7).text()
                                    specification_text = win.tableWidget.item(self.Single_rows, 8).text()
                                    quantity_text = win.tableWidget.item(self.Single_rows, 10).text()
                                    print(str(text_5).split('.')[0])
                                    float_num = re.compile(r'\D')
                                    float_re = float_num.findall(str(text_5).split('.')[0])
                                    if float_re == [] and text_5!='':
                                        for row in self.rows_dict.keys():  # 第三种方法获取值
                                            all_name1 =self.window.tableWidget_quota.item(row,1).text()  # 获取单元格内容
                                            all_name2 = self.window.tableWidget_quota.item(row, 2).text()  # 获取单元格内容
                                            all_name3 = self.window.tableWidget_quota.item(row, 3).text()  # 获取单元格内容
                                            all_name4 = self.window.tableWidget_quota.item(row, 4).text()  # 获取单元格内容
                                            all_list=[all_name1,all_name2,all_name3,all_name4,sys_text,name_text,specification_text,quantity_text]
                                            float_quota = re.compile(r'\D')
                                            float_q = float_quota.findall(str(all_name1).split('.')[0])
                                            if float_q==[] and all_name1!='':
                                                if float(text_5)<=float(all_list[0]):
                                                    num+=1
                                                    if num==1:
                                                        values_dict[self.Single_rows]=all_list
                                                        # inster_rows = self.Single_rows + 1
                                                        # win.tableWidget.insertRow(inster_rows)
                                                        break
                    if values_dict!={}:
                        print(list(values_dict.values()))
                        for insert_row,row in enumerate(values_dict.keys()):
                            win.tableWidget.insertRow(insert_row+row+1)
                            for column in range(0, win.tableWidget.columnCount()):
                                win.tableWidget.setItem(insert_row+row+1, column,QTableWidgetItem(''))
                            value=list(values_dict.values())[insert_row]
                            print(value)
                            win.tableWidget.setItem(insert_row+row+1, 0,QTableWidgetItem(str(value[1])))
                            win.tableWidget.setItem(insert_row+row+1, 1,QTableWidgetItem(str(value[2])))
                            win.tableWidget.setItem(insert_row+row+1, 3,QTableWidgetItem(str(value[3])))
                            win.tableWidget.setItem(insert_row+row+1, 6,QTableWidgetItem(str(value[4])))
                            win.tableWidget.setItem(insert_row + row + 1, 7, QTableWidgetItem(str(value[5])))
                            win.tableWidget.setItem(insert_row + row + 1, 8, QTableWidgetItem(str(value[6])))
                            win.tableWidget.setItem(insert_row + row + 1, 10, QTableWidgetItem(str(value[7])))
                            win.tableWidget.setItem(insert_row+row+1, 5,QTableWidgetItem('《定额》'))

                        # 撤回
                        if win.table_do == '执行':
                            self.new_undo_dict = {}
                            new_text_list = []
                            self.new_undo_dict[win.tableWidget] = new_text_list
                            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                                t_list = []
                                new_text_list.append(t_list)
                                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                    t_list.append(self.text)
                            if self.Click_list != new_text_list:
                                command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict,
                                                             self.new_undo_dict)
                                win.undoStack_del.push(command)
                        win.tab_change = '提取'
        except:
            pass

    def handleSelectionChange(self):#加载清单库树界面
        city_name=self.window.comboBox.currentText()#获取当前选中的选项的文本
        self.files=self.files_qdk+city_name
        self.window.treeWidget_qdk.clear()
        for file_name in os.listdir(self.files):#文件夹作为一级目录
            self.root = QTreeWidgetItem(self.window.treeWidget_qdk)
            self.root.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            self.root.setText(0, file_name)
            self.root.setSizeHint(0, QSize(0, 25))
            self.files_son =os.path.join(self.files,file_name)
            for file in os.listdir(self.files_son):#excel名称作为二级目录
                if '~$' not in file and 'xlsx' == file.split('.')[-1]:
                    self.son = QTreeWidgetItem(self.root)
                    self.son.setText(0, str(file.split('.')[0]))
                    self.son.setSizeHint(0, QSize(0, 25))
                    f = os.path.join(self.files_son,file)
                    self.wb = openpyxl.load_workbook(f, read_only=False, data_only=True, keep_links=False)
                    for i in range(len(self.wb.sheetnames)):#excel里的sheet作为三级目录
                        self.grandson = QTreeWidgetItem(self.son)
                        self.grandson.setText(0, self.wb.sheetnames[i])
                        self.grandson.setSizeHint(0, QSize(0, 25))
        # child_item = self.window.treeWidget_qdk.topLevelItem(2).child(3)
        # self.window.treeWidget_qdk.scrollToItem(child_item)  # 滚轮定位
        # self.window.treeWidget_qdk.expandAll()
    qp_text_list=[]
    qdk_save = '不保存'
    def treeWidget_qdk_Clicked(self):#单击清单库树界面显示清单库
        self.window.tableWidget_features.clearContents()  # 可以清除表格所有的内容
        self.window.tableWidget_features.setRowCount(0)
        self.qp_text_list = []
        self.root_num = self.window.treeWidget_qdk.topLevelItemCount()  # 获取根节点数量.currentIndex().row()
        self.item=self.window.treeWidget_qdk.currentItem()

        self.window.tableWidget_qdk.setRowCount(0)  # 设置行数
        if self.item.parent():
            self.second=self.item.parent()
            if self.second.parent():
                self.window.pushButton_save_qdk.setEnabled(True)
        #         self.fist=self.second.parent()
        #         self.bin_path_qdk = os.path.join(self.files,self.fist.text(0), self.second.text(0)+'.xlsx')
        #         self.file_path_qdk = os.path.join(self.files,self.fist.text(0),'~$'+self.second.text(0)+'.xlsx')

        # for dirpath, dirnames, filenames in os.walk(self.files):
        #     for filename in filenames:
        #         files = os.path.join(dirpath, filename)
                files = os.path.join(self.files,self.second.parent().text(0), self.second.text(0)+'.xlsx')
                if '~$' not in files and 'xlsx' == files.split('.')[-1]:
                    # print(filename)
                    self.wb = openpyxl.load_workbook(files, read_only=False, data_only=True, keep_links=False)
                    self.num = 0
                    for i in range(len(self.wb.sheetnames)):
                        if self.item.text(0) == self.wb.sheetnames[i]: #or self.item.text(0)==str(filename).split('.')[0]
                            # self.num = 0
                            self.ws = self.wb[self.wb.sheetnames[i]]
                            # print(self.ws)
                            self.values = list(
                                self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                            max_num = len(self.values)
                            # print(max_num)
                            for wb_ws_row in range(0, max_num):
                                data = self.values[wb_ws_row]
                                self.num += 1
                                # print(self.num,data)
                                self.window.tableWidget_qdk.setRowCount(self.num)  # 设置行数
                                self.qp_text = QPlainTextEdit()
                                self.qp_text_list.append(self.qp_text)
                                self.window.tableWidget_qdk.setItem(int(self.num) - 1, 0,QTableWidgetItem(''))
                                self.window.tableWidget_qdk.setItem(int(self.num) - 1, 1, QTableWidgetItem(''))
                                self.window.tableWidget_qdk.setItem(int(self.num) - 1, 2, QTableWidgetItem(''))
                                self.window.tableWidget_qdk.setCellWidget(int(self.num) - 1, 3, self.qp_text)
                                self.window.tableWidget_qdk.setItem(int(self.num) - 1, 4, QTableWidgetItem(''))
                                for j in range(0, len(data)):
                                    # self.item1 = QTableWidgetItem(str(data[j]).replace('None', ''))
                                    self.item1 =str(data[j]).replace('None', '')
                                    if j ==3:
                                        self.qp_text_list[int(self.num) - 1].setPlainText(self.item1)
                                    if j !=3:
                                        self.window.tableWidget_qdk.setItem(int(self.num) - 1, int(j),QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdk.verticalHeader().resizeSection(int(self.num)-1, 100)  # 调整每一行的大小为100像素
                                    # self.window.tableWidget_qdk.verticalHeader().setSectionResizeMode(int(self.num) - 1,QHeaderView.ResizeToContents)
                            self.window.tableWidget_qdk.setStyleSheet("gridline-color: rgb(257, 1, 0)")
                            # self.window.tableWidget_qdk.horizontalHeader().setStyleSheet("color: rgb(0, 83, 128);border:1px solid rgb(210, 210, 210);")
                            self.window.tableWidget_qdk.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)  # 设置列宽，列宽可调
                            self.window.tableWidget_qdk.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
                            self.window.tableWidget_qdk.horizontalHeader().resizeSection(4, 100)  # 调整第2列的大小为500像素
    def parse_quota(self):#解析窗
        self.window.dockWidget.show()
    def quota_Double_click(self):# 双击插入定额库
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.checken_list = []
            if win.tableWidget.rowCount() != 0:
                for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                    Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                    if Item1 == Qt.Checked:
                        check_allname = win.tableWidget.item(self.Single_rows, 7).text()
                        self.checken_list.append(self.Single_rows)

                if self.checken_list != []:
                    # "同名称—单位清单插入定额", "同名称—规格—单位清单插入定额",
                    # "同系统—名称—规格—单位清单插入定额"
                    items = ["当前工程选中清单下插入定额"]
                    com_text, ok = QInputDialog().getItem(self.window, "选择应用范围", "应用其它单位工程需提前打√",
                                                          items, 0, True)
                    if ok:
                        # if com_text == '同名称—单位清单插入定额':
                        #     pass
                        # if com_text == '同名称—单位清单插入定额':
                        #     pass
                        # if com_text == '同名称—规格—单位清单插入定额':
                        #     pass
                        if com_text == '当前工程选中清单下插入定额':
                            # 撤销
                            self.old_undo_dict = {}
                            self.Click_list = []
                            self.old_undo_dict[win.tableWidget] = self.Click_list
                            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                                t_list = []
                                self.Click_list.append(t_list)
                                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                    t_list.append(self.text)
                            win.tab_change = '不提取'
                            row=self.window.tableWidget_quota.currentRow()
                            quota0 = self.window.tableWidget_quota.item(row, 0).text()
                            quota1 = self.window.tableWidget_quota.item(row, 1).text()
                            quota2 = self.window.tableWidget_quota.item(row, 2).text()
                            quota3 = self.window.tableWidget_quota.item(row, 3).text()
                            quota4 = self.window.tableWidget_quota.item(row, 4).text()
                            quota5 = self.window.tableWidget_quota.item(row, 5).text()
                            for num,row in enumerate(self.checken_list):
                                sys_text = win.tableWidget.item(num+row, 6).text()
                                name_text = win.tableWidget.item(num + row, 7).text()
                                specification_text = win.tableWidget.item(num + row, 8).text()
                                quantity_text = win.tableWidget.item(num+row, 10).text()
                                print(sys_text)
                                win.tableWidget.insertRow(num+row+1)
                                for column in range(0,win.tableWidget.columnCount()):
                                    item = QtWidgets.QTableWidgetItem('')
                                    item.setFlags(Qt.ItemIsEnabled)
                                    # item.setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)
                                    brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                    brush.setStyle(QtCore.Qt.SolidPattern)
                                    item.setBackground(brush)# 背景颜色

                                    # brush = QtGui.QBrush(QtGui.QColor(85, 85, 255))
                                    # brush.setStyle(QtCore.Qt.SolidPattern)
                                    # item.setForeground(brush)# 字体颜色QtCore.Qt.NoBrush
                                    win.tableWidget.setItem(num + row + 1, column,item)
                                item2 = QtWidgets.QTableWidgetItem(str(quota2))
                                # item2.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                # brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                # brush.setStyle(QtCore.Qt.SolidPattern)
                                # item2.setBackground(brush)  # 背景颜色
                                win.tableWidget.setItem(num+row+1, 0, item2)

                                item3 = QtWidgets.QTableWidgetItem(str(quota3))
                                # item3.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                # brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                # brush.setStyle(QtCore.Qt.SolidPattern)
                                # item3.setBackground(brush)  # 背景颜色
                                win.tableWidget.setItem(num + row + 1, 1,item3)

                                item4 = QtWidgets.QTableWidgetItem(str(quota4))
                                # item4.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                # brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                # brush.setStyle(QtCore.Qt.SolidPattern)
                                # item4.setBackground(brush)  # 背景颜色
                                win.tableWidget.setItem(num + row + 1, 3,item4)

                                item6 = QtWidgets.QTableWidgetItem('《定额》')
                                item6.setFlags(Qt.ItemIsEnabled) # 内容不能编辑
                                # item6.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                # brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                # brush.setStyle(QtCore.Qt.SolidPattern)
                                # item6.setBackground(brush)  # 背景颜色
                                win.tableWidget.setItem(num + row + 1, 5, item6)

                                item_sys = QtWidgets.QTableWidgetItem(str(sys_text))
                                # item_sys.setTextAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                                # brush = QtGui.QBrush(QtGui.QColor(244, 245, 255))
                                # brush.setStyle(QtCore.Qt.SolidPattern)
                                # item_sys.setBackground(brush)  # 背景颜色
                                win.tableWidget.setItem(num + row + 1, 6, item_sys)
                                win.tableWidget.setItem(num + row + 1, 7, QtWidgets.QTableWidgetItem(str(name_text)))  # 写入工程量
                                win.tableWidget.setItem(num + row + 1, 8,QtWidgets.QTableWidgetItem(str(specification_text)))  # 写入工程量
                                win.tableWidget.setItem(num + row + 1, 10, QtWidgets.QTableWidgetItem(str(quantity_text)))# 写入工程量
                                win.tableWidget.item(num+row, 7).setCheckState(Qt.Unchecked)
                            # 撤回
                            if win.table_do == '执行':
                                self.new_undo_dict = {}
                                new_text_list = []
                                self.new_undo_dict[win.tableWidget] = new_text_list
                                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                                    t_list = []
                                    new_text_list.append(t_list)
                                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                        t_list.append(self.text)
                                if self.Click_list != new_text_list:
                                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                                    win.undoStack_del.push(command)
                            win.tab_change = '提取'


    def Double_click(self):# 双击插入清单
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.checken_list = []
            if win.tableWidget.rowCount() != 0:
                for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                    Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                    if Item1 == Qt.Checked:
                        check_allname = win.tableWidget.item(self.Single_rows, 7).text()
                        self.checken_list.append(check_allname)
                if self.checken_list != []:
                    items = ["应用当前工程选中清单项", "应用同名称—单位", "应用同名称—规格—单位", "应用同系统—名称—规格—单位"]
                    com_text, ok = QInputDialog().getItem(self.window, "选择应用范围", "应用其它单位工程需提前打√", items, 0, True)
                    if ok:
                        if com_text=='应用同名称—单位':
                            self.qdk_name()
                        if com_text=='应用同名称—规格—单位':
                            self.qdk_name_sp()
                        if com_text=='应用同系统—名称—规格—单位':
                            self.qdk_sys_name_sp()
                        if com_text == '应用当前工程选中清单项':
                            try:
                                win.tab_change = '不提取'
                                # 撤销
                                self.old_undo_dict = {}
                                self.Click_list = []
                                self.old_undo_dict[win.tableWidget] = self.Click_list
                                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                                    t_list = []
                                    self.Click_list.append(t_list)
                                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                        t_list.append(self.text)
                                # 写入清单
                                self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
                                self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                self.text3 = str(self.qp_text_list[self.row].toPlainText())
                                for self.win_row in range(0, int(win.tableWidget.rowCount())):
                                    Item1 = win.tableWidget.item(self.win_row, 7).checkState()
                                    if Item1 == Qt.Checked:
                                        self.fre = win.tableWidget.item(self.win_row, 7).text()  # 获取单元格内容
                                        self.eight = str(win.tableWidget.item(self.win_row, 8).text()).split('&')
                                        if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                            self.sre = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[0]
                                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                            win.tableWidget.setItem(self.win_row,4, QTableWidgetItem(self.text0.split('、')[0]))

                                            for j in range(1,5):
                                                if j!=3:
                                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(self.text))
                                                if j ==3:
                                                    text=self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                                    win.tableWidget.setItem(self.win_row, j-1,QTableWidgetItem(text))
                                            win.tableWidget.verticalHeader().setSectionResizeMode(self.win_row,QHeaderView.ResizeToContents)
                                        if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                            self.sre = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[0]
                                            self.four = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[1]
                                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                            win.tableWidget.setItem(self.win_row,4, QTableWidgetItem(self.text0.split('、')[0]))
                                            for j in range(1,5):
                                                if j != 3:
                                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(self.text))
                                                if j == 3:
                                                    text=self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four)
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(text))
                                            win.tableWidget.verticalHeader().setSectionResizeMode(self.win_row,QHeaderView.ResizeToContents)
                                        if len(self.eight) == 3 and  'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                                            self.sre = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[0]
                                            self.four = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[1]
                                            self.five = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[2]
                                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                            win.tableWidget.setItem(self.win_row,4, QTableWidgetItem(self.text0.split('、')[0]))
                                            for j in range(1,5):
                                                if j != 3:
                                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(self.text))
                                                if j == 3:
                                                    text=self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4',self.five)
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(text))
                                            win.tableWidget.verticalHeader().setSectionResizeMode(self.win_row,QHeaderView.ResizeToContents)

                                        if len(self.eight) == 4 and  'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                                            self.sre = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[0]
                                            self.four = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[1]
                                            self.five = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[2]
                                            self.six = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[3]
                                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                            win.tableWidget.setItem(self.win_row,4, QTableWidgetItem(self.text0.split('、')[0]))
                                            for j in range(1,5):
                                                if j != 3:
                                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(self.text))
                                                if j == 3:
                                                    text=self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4',self.five).replace('NP5',self.six)
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(text))
                                            win.tableWidget.verticalHeader().setSectionResizeMode(self.win_row,QHeaderView.ResizeToContents)
                                        elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                            self.sre = str(win.tableWidget.item(self.win_row, 8).text()).split('&')[0]
                                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                            win.tableWidget.setItem(self.win_row,4, QTableWidgetItem(self.text0.split('、')[0]))
                                            for j in range(1,5):
                                                if j != 3:
                                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                                    win.tableWidget.setItem(self.win_row, j - 1, QTableWidgetItem(self.text))
                                                if j == 3:
                                                    win.tableWidget.setItem(self.win_row, j-1, QTableWidgetItem(self.text3))
                                            win.tableWidget.verticalHeader().setSectionResizeMode(self.win_row,QHeaderView.ResizeToContents)
                                        win.tableWidget.item(self.win_row, 7).setCheckState(Qt.Unchecked)
                                # 撤回
                                if win.table_do == '执行':
                                    self.new_undo_dict = {}
                                    new_text_list = []
                                    self.new_undo_dict[win.tableWidget] = new_text_list
                                    for self.Single_rows in range(0, win.tableWidget.rowCount()):
                                        t_list = []
                                        new_text_list.append(t_list)
                                        for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                            self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                            t_list.append(self.text)
                                    if self.Click_list != new_text_list:
                                        command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                                        win.undoStack_del.push(command)
                                win.tab_change = '提取'
                            except:
                                pass

    def qdk_name(self):# 复用同名称
        try:
            win.tab_change = '不提取'
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[win.tableWidget] = self.Click_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)

            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    # 撤销
                    self.Click_list = []
                    self.old_undo_dict[self.tab] = self.Click_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        self.Click_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            # 执行写入
            Checken_dict = {}
            self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
            self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
            self.text3 = str(self.qp_text_list[self.row].toPlainText())
            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Checked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    check_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    check_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    check_copytext4 = win.tableWidget.item(self.Single_rows, 4).text()
                    Checken_dict[check_allname, check_unit] = check_copytext4
                    win.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                    if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                    if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                    if len(self.eight) == 3 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                    if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                    elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Unchecked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    uncheck_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    uncheck_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    merge = (uncheck_allname,uncheck_unit)
                    if merge in Checken_dict:
                        if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)
                        if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                        if len(self.eight) == 3 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                        if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)
                        elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)
            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    if str(self.tab.rowCount()) != '0':
                        for self.other_rows in range(0,int(self.tab.rowCount())):
                            self.fre = self.tab.item(self.other_rows, 7).text()  # 获取单元格内容
                            self.eight = str(self.tab.item(self.other_rows, 8).text()).split('&')
                            self.others_name = self.tab.item(self.other_rows, 7).text()
                            self.others_unit = self.tab.item(self.other_rows, 9).text()
                            others_merge = (self.others_name, self.others_unit)
                            if others_merge in Checken_dict:
                                if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 3  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =  self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 4  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    self.six = str(self.tab.item(self.other_rows, 8).text()).split('&')[3]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows, QHeaderView.ResizeToContents)
                                elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text3))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
          # 回撤

            self.new_undo_dict = {}
            new_text_list = []
            self.new_undo_dict[win.tableWidget] = new_text_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                new_text_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if win.tabldict!= {}:
                for self.tab in win.tabldict.values():
                    new_text_list = []
                    self.new_undo_dict[self.tab] = new_text_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            if self.old_undo_dict != self.new_undo_dict:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                win.undoStack_del.push(command)
            win.tab_change = '提取'
        except:
            pass
    def qdk_name_sp(self):# 清单应用同名称同规格同单位
        try:
            win.tab_change = '不提取'
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[win.tableWidget] = self.Click_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)

            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    # 撤销
                    self.Click_list = []
                    self.old_undo_dict[self.tab] = self.Click_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        self.Click_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            # 写入清单
            Checken_dict = {}
            self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
            self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
            self.text3 = str(self.qp_text_list[self.row].toPlainText())
            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Checked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    check_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    check_specs = win.tableWidget.item(self.Single_rows, 8).text()
                    check_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    check_copytext4 = win.tableWidget.item(self.Single_rows, 4).text()
                    Checken_dict[check_allname,check_specs, check_unit] = check_copytext4
                    win.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                    if len(self.eight) == 1  and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                              QHeaderView.ResizeToContents)
                    if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',                                                                                                           self.four)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                              QHeaderView.ResizeToContents)
                    if len(self.eight) == 3 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4', self.five)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                              QHeaderView.ResizeToContents)
                    if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',
                                                                                                                 self.four).replace(
                                'NP4', self.five).replace('NP5', self.six)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                    elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Unchecked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    uncheck_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    uncheck_specs=win.tableWidget.item(self.Single_rows, 8).text()
                    uncheck_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    merge = (uncheck_allname,uncheck_specs, uncheck_unit)
                    if merge in Checken_dict:
                        if len(self.eight) == 1  and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                        if len(self.eight) == 2  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                        if len(self.eight) == 3  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                                win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                        if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3', self.four).replace('NP4', self.five).replace('NP5', self.six)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)
                        elif  'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    if str(self.tab.rowCount()) != '0':
                        for self.other_rows in range(0,int(self.tab.rowCount())):
                            self.fre = self.tab.item(self.other_rows, 7).text()  # 获取单元格内容
                            self.eight = str(self.tab.item(self.other_rows, 8).text()).split('&')
                            self.others_name = self.tab.item(self.other_rows, 7).text()
                            self.others_specs =self.tab.item(self.other_rows, 8).text()
                            self.others_unit = self.tab.item(self.other_rows, 9).text()
                            others_merge = (self.others_name, self.others_specs,self.others_unit)
                            if others_merge in Checken_dict:
                                if len(self.eight) == 1  and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 2  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 3  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 4  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    self.six = str(self.tab.item(self.other_rows, 8).text()).split('&')[3]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text = self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows, QHeaderView.ResizeToContents)
                                elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text3))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
            # 回撤
            self.new_undo_dict = {}
            new_text_list = []
            self.new_undo_dict[win.tableWidget] = new_text_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                new_text_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    new_text_list = []
                    self.new_undo_dict[self.tab] = new_text_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            if self.old_undo_dict != self.new_undo_dict:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                win.undoStack_del.push(command)
            win.tab_change = '提取'
        except:
            pass
    def qdk_sys_name_sp(self):# 清单应用同名称同规格同单位
        try:
            win.tab_change = '不提取'
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[win.tableWidget] = self.Click_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)

            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    # 撤销
                    self.Click_list = []
                    self.old_undo_dict[self.tab] = self.Click_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        self.Click_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            # 写入清单
            Checken_dict = {}
            self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
            self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
            self.text3 = str(self.qp_text_list[self.row].toPlainText())
            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Checked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    check_allsys = win.tableWidget.item(self.Single_rows, 6).text()
                    check_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    check_specs = win.tableWidget.item(self.Single_rows, 8).text()
                    check_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    check_copytext4 = win.tableWidget.item(self.Single_rows, 4).text()
                    Checken_dict[check_allsys,check_allname, check_specs, check_unit] = check_copytext4
                    # Checken_dict[check_allsys+check_specs+check_allname+check_unit] = check_copytext4
                    win.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                    if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)

                    if len(self.eight) == 2  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',                                                                                                            self.four)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                              QHeaderView.ResizeToContents)
                    if len(self.eight) == 3 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text ))
                            if j == 3:
                                text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',
                                                                                                                 self.four).replace(
                                'NP4', self.five)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                              QHeaderView.ResizeToContents)
                    if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                        self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                        self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows, QHeaderView.ResizeToContents)

                    elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                        self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                        # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                        win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                        for j in range(1, 5):
                            if j != 3:
                                self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                            if j == 3:
                                win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                        win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

            for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                Item1 = win.tableWidget.item(self.Single_rows, 7).checkState()
                if Item1 == Qt.Unchecked:
                    self.fre = win.tableWidget.item(self.Single_rows, 7).text()  # 获取单元格内容
                    self.eight = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')
                    uncheck_sys = win.tableWidget.item(self.Single_rows, 6).text()
                    uncheck_allname = win.tableWidget.item(self.Single_rows, 7).text()
                    uncheck_specs = win.tableWidget.item(self.Single_rows, 8).text()
                    uncheck_unit = win.tableWidget.item(self.Single_rows, 9).text()
                    # for k, v in Checken_dict.items():
                    #     if uncheck_sys in k and uncheck_specs+uncheck_allname+uncheck_unit in k:
                    merge = (uncheck_sys,uncheck_allname,uncheck_specs, uncheck_unit)
                    if merge in Checken_dict:
                        print(merge)
                        if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                        if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',                                                                                                                   self.four)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)
                        if len(self.eight) == 3  and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text = self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',
                                                                                                                     self.four).replace(
                                    'NP4', self.five)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,
                                                                                  QHeaderView.ResizeToContents)
                        if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            self.four = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[1]
                            self.five = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[2]
                            self.six = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[3]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',
                                                                                                                     self.four).replace(
                                    'NP4', self.five).replace('NP5', self.six)
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(text))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

                        elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                            self.sre = str(win.tableWidget.item(self.Single_rows, 8).text()).split('&')[0]
                            # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                            win.tableWidget.setItem(self.Single_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                            for j in range(1, 5):
                                if j != 3:
                                    self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text))
                                if j == 3:
                                    win.tableWidget.setItem(self.Single_rows, j - 1, QTableWidgetItem(self.text3))
                            win.tableWidget.verticalHeader().setSectionResizeMode(self.Single_rows,QHeaderView.ResizeToContents)

            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    if str(self.tab.rowCount()) != '0':
                        for self.other_rows in range(0,int(self.tab.rowCount())):
                            self.fre = self.tab.item(self.other_rows, 7).text()  # 获取单元格内容
                            self.eight = str(self.tab.item(self.other_rows, 8).text()).split('&')
                            self.others_sys = self.tab.item(self.other_rows, 6).text()
                            self.others_name = self.tab.item(self.other_rows, 7).text()
                            self.others_specs =self.tab.item(self.other_rows, 8).text()
                            self.others_unit = self.tab.item(self.other_rows, 9).text()
                            # for k, v in Checken_dict.items():
                            #     if self.others_sys in k and self.others_specs + self.others_name +self.others_unit in k:
                            merge = (self.others_sys,self.others_name,self.others_specs,self.others_unit)
                            if merge in Checken_dict:
                                if len(self.eight) == 1 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 2 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =self.text3.replace('NP1', str(self.fre)).replace('NP2', self.sre).replace('NP3',self.four)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 3 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
                                if len(self.eight) == 4 and 'NP2' in self.text3 and 'NP3' in self.text3 and 'NP4' in self.text3 and 'NP5' in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    self.four = str(self.tab.item(self.other_rows, 8).text()).split('&')[1]
                                    self.five = str(self.tab.item(self.other_rows, 8).text()).split('&')[2]
                                    self.six = str(self.tab.item(self.other_rows, 8).text()).split('&')[3]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            text =self.text3.replace('NP1', str(self.fre)).replace('NP2',self.sre).replace('NP3',self.four).replace('NP4', self.five).replace('NP5', self.six)
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(text))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows, QHeaderView.ResizeToContents)
                                elif 'NP1' not in self.text3 and 'NP2' not in self.text3 and 'NP3' not in self.text3 and 'NP4' not in self.text3 and 'NP5' not in self.text3:
                                    self.sre = str(self.tab.item(self.other_rows, 8).text()).split('&')[0]
                                    # self.text0 = self.window.tableWidget_qdk.item(self.row, 0).text()  # 获取单元格内容
                                    self.tab.setItem(self.other_rows, 4, QTableWidgetItem(self.text0.split('、')[0]))
                                    for j in range(1, 5):
                                        if j != 3:
                                            self.text = self.window.tableWidget_qdk.item(self.row, j).text()  # 获取单元格内容
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text))
                                        if j == 3:
                                            self.tab.setItem(self.other_rows, j - 1, QTableWidgetItem(self.text3))
                                    self.tab.verticalHeader().setSectionResizeMode(self.other_rows,QHeaderView.ResizeToContents)
            # 回撤
            self.new_undo_dict = {}
            new_text_list = []
            self.new_undo_dict[win.tableWidget] = new_text_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                new_text_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            if win.tabldict != {}:
                for self.tab in win.tabldict.values():
                    new_text_list = []
                    self.new_undo_dict[self.tab] = new_text_list
                    for self.Single_rows in range(0, self.tab.rowCount()):
                        t_list = []
                        new_text_list.append(t_list)
                        for self.Single_colum in range(0, self.tab.columnCount()):
                            self.text = self.tab.item(self.Single_rows, self.Single_colum).text()
                            t_list.append(self.text)
            if self.old_undo_dict != self.new_undo_dict:
                command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                win.undoStack_del.push(command)
            win.tab_change = '提取'
        except:
            pass

    def save_qdk(self): #保存清单库
        try:
            row_k = self.window.treeWidget_qdk.currentIndex().row()  # 获取行
            if int(row_k) != -1:
                self.item = self.window.treeWidget_qdk.currentItem()
                if self.item.parent():
                    self.second=self.item.parent()
                    if self.second.parent():
                        self.fist=self.second.parent()
                        self.bin_path_qdk = os.path.join(self.files,self.fist.text(0), self.second.text(0)+'.xlsx')
                        self.file_path_qdk = os.path.join(self.files,self.fist.text(0),'~$'+self.second.text(0)+'.xlsx')

                    if os.path.exists(self.file_path_qdk):
                        QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件')
                    else:
                        self.wb = openpyxl.load_workbook(self.bin_path_qdk, read_only=False, data_only=True, keep_links=False)
                        self.ws=self.wb[self.item.text(0)]
                        self.rows = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
                        self.table_rows = self.window.tableWidget_qdk.rowCount()
                        self.table_column = self.window.tableWidget_qdk.columnCount()
                        # print(self.ws.max_row)
                        self.ws.delete_rows(1, self.ws.max_row+1)  # 删除行，前边是起始行，后边是行数
                        for i in range(0, int(self.table_rows)):
                            for j in range(0, int(self.table_column)):
                                if j !=3:
                                    self.text = self.window.tableWidget_qdk.item(i, j).text()  # 获取单元格内容
                                    self.ws.cell(i + 1, j + 1, str(self.text)).alignment = self.alig
                                    self.ws.cell(i + 1, j + 1, str(self.text)).font = self.font
                                if j==3:
                                    text3=str(self.qp_text_list[i].toPlainText())
                                    self.ws.cell(i + 1, j + 1, str(text3)).alignment = self.alig
                                    self.ws.cell(i + 1, j + 1, str(text3)).font = self.font
                                # print(i,j, str(self.text))
                        if os.path.exists(self.file_path_qdk):
                            QMessageBox.information(self.window, '温馨提示', '数据无法写入，请先关闭Excel文件') # 防止在写入时打开文件，软件崩溃。
                        else:
                            self.wb.save(self.bin_path_qdk)  # 保存新建的文件
                            win.Excel_qdk('重新加载清单库')
        except:
            pass
    go = '不剪切板'
    def clipboard(self):
        self.go='剪切板'
    def k_copy_clipboard(self):# 清单库剪切板复制外部数据
        try:
            if self.go == '剪切板':
                mimeData = self.clipboard_qd.mimeData()
                if mimeData.hasFormat('text/plain'):
                    all_text = []
                    test_list = []
                    values_list = []
                    texts = mimeData.text()
                    if '\t' in texts:
                        criteria = '"(.*?)"'
                        compile = re.compile(criteria, re.S)
                        self.results = compile.findall(texts)
                        sp_n = texts.replace('"', '').split('\n')  # 拆n
                        for sp in sp_n:
                            sp_t = str(sp).split('\t')  # 拆t
                            test_list.append(sp_t)
                        for new_list in test_list:
                            for new in new_list:
                                all_text.append(str(new))  # 构建新列表

                        for result in self.results:
                            res = str(result).split('\n')[1:]
                            for rs in res:
                                all_text.remove(str(rs))  # 删除新列表同一个单元格内多余的换行后内容
                        if '' in all_text[-1]:
                            all_text.pop(-1)

                        row = len(all_text) - texts.count('\t')  # 获取行
                        column = int(len(all_text) / row - 1)  # 获取列
                        num = int(len(all_text) / row)

                        split_t = texts.replace('"', '').split('\t')  # 拆t
                        for i in range(0, row):  # 循环行数
                            split_n = split_t[column].split('\n')[:-1]  # 在一个单元格内有多余的换行符，需要截取有用的内容
                            join_n = '\n'.join(split_n)
                            last = split_t[column].split('\n')[-1]  # 截取最后的换行符内容
                            split_t[column] = join_n  # 修改列表内容
                            split_t.insert(column + 1, last)  # 插入列表内容
                            column += num

                        if '' in split_t[-1]:
                            split_t.pop(-1)  # 删除最后一个元素‘’

                        for r in range(row):  # 按行循环
                            values_list.append(split_t[0:num])  # 按列添加到列表
                            for c in range(num):  # 按列删除列表内容
                                split_t.pop(0)
                        if self.window.tableWidget_qdk.currentRow()!=-1:
                            for i, values in enumerate(values_list):  # 每次循环一行写入到表格里一行内容
                                if i ==0:
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        if self.window.tableWidget_qdk.currentColumn() + j == 3:
                                            self.qp_text_list[self.window.tableWidget_qdk.currentRow()].setPlainText(self.item1)
                                        if self.window.tableWidget_qdk.currentColumn() + j != 3:
                                            self.window.tableWidget_qdk.setItem(
                                                self.window.tableWidget_qdk.currentRow(),
                                                self.window.tableWidget_qdk.currentColumn() + j,
                                                QTableWidgetItem(self.item1))
                                if i>=1:
                                    self.window.tableWidget_qdk.insertRow(self.window.tableWidget_qdk.currentRow()+i)
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i,0,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 1,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 2,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 4,QTableWidgetItem(''))

                                    self.qdk_text = QPlainTextEdit()
                                    self.window.tableWidget_qdk.setCellWidget(self.window.tableWidget_qdk.currentRow() + i,3, self.qdk_text)
                                    self.qp_text_list.insert(self.window.tableWidget_qdk.currentRow() + i,self.qdk_text)
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        if self.window.tableWidget_qdk.currentColumn() + j == 3:
                                            self.qp_text_list[self.window.tableWidget_qdk.currentRow()+i].setPlainText(self.item1)
                                        if self.window.tableWidget_qdk.currentColumn() + j != 3:
                                            self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow()+ i,
                                                             self.window.tableWidget_qdk.currentColumn() + j,
                                                                                QTableWidgetItem(self.item1))
                                self.window.tableWidget_qdk.verticalHeader().resizeSection(self.window.tableWidget_qdk.currentRow()+ i,80)  # 调整每一行的大小为100像素

                    if '\t' not in texts and '"\n' in texts:# 复制整列内有单元格内有换行
                        sp_n = texts.split('"')  # 拆n
                        text_list = []
                        for sp in sp_n:
                            if sp != '':
                                if sp != '\n':
                                    a = sp.split('\n')
                                    if a[0] != '' and a[-1] != '':
                                        join_text = '\n'.join(a)
                                        text_list.append(str(join_text))

                                    if a[-1] == '':
                                        for b in a:
                                            if b != '':
                                                text_list.append(str(b))
                        if self.window.tableWidget_qdk.currentRow() != -1:
                            for i, value in enumerate(text_list):  # 每次循环一行写入到表格里一行内容
                                if i ==0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdk.setItem(
                                        self.window.tableWidget_qdk.currentRow(),
                                        self.window.tableWidget_qdk.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i>=1:
                                    self.window.tableWidget_qdk.insertRow(self.window.tableWidget_qdk.currentRow()+i)
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i,0,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 1,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 2,QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 4,QTableWidgetItem(''))

                                    self.qdk_text = QPlainTextEdit()
                                    self.window.tableWidget_qdk.setCellWidget(self.window.tableWidget_qdk.currentRow() + i,3, self.qdk_text)
                                    self.qp_text_list.insert(self.window.tableWidget_qdk.currentRow() + i,self.qdk_text)
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow()+ i,
                                                         self.window.tableWidget_qdk.currentColumn(),
                                                                            QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdk.verticalHeader().resizeSection(self.window.tableWidget_qdk.currentRow()+ i,80)  # 调整每一行的大小为100像素


                    if '\t' not in texts and '"\n' not in texts:# 复制整列内有单元格内没有换行
                        sp_n = texts.split('\n')
                        if '' in sp_n[-1]:
                            sp_n.pop(-1)  # 删除最后一个元素‘’
                        if self.window.tableWidget_qdk.currentRow() != -1:
                            for i, value in enumerate(sp_n):  # 每次循环一行写入到表格里一行内容
                                if i == 0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdk.setItem(
                                        self.window.tableWidget_qdk.currentRow(),
                                        self.window.tableWidget_qdk.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i >= 1:
                                    self.window.tableWidget_qdk.insertRow(self.window.tableWidget_qdk.currentRow() + i)
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 0,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 1,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 2,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i, 4,
                                                                        QTableWidgetItem(''))

                                    self.qdk_text = QPlainTextEdit()
                                    self.window.tableWidget_qdk.setCellWidget(
                                        self.window.tableWidget_qdk.currentRow() + i, 3, self.qdk_text)
                                    self.qp_text_list.insert(self.window.tableWidget_qdk.currentRow() + i,
                                                              self.qdk_text)
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdk.setItem(self.window.tableWidget_qdk.currentRow() + i,
                                                                        self.window.tableWidget_qdk.currentColumn(),
                                                                        QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdk.verticalHeader().resizeSection(
                                        self.window.tableWidget_qdk.currentRow() + i, 100)

        except:
            pass
    def quota_copy_clipboard(self):# 定额库剪切板复制外部数据
        try:
            if self.go == '剪切板':
                mimeData = self.clipboard_qd.mimeData()
                if mimeData.hasFormat('text/plain'):
                    all_text = []
                    test_list = []
                    values_list = []
                    texts = mimeData.text()
                    if '\t' in texts:
                        criteria = '"(.*?)"'
                        compile = re.compile(criteria, re.S)
                        self.results = compile.findall(texts)
                        sp_n = texts.replace('"', '').split('\n')  # 拆n
                        for sp in sp_n:
                            sp_t = str(sp).split('\t')  # 拆t
                            test_list.append(sp_t)
                        for new_list in test_list:
                            for new in new_list:
                                all_text.append(str(new))  # 构建新列表

                        for result in self.results:
                            res = str(result).split('\n')[1:]
                            for rs in res:
                                all_text.remove(str(rs))  # 删除新列表同一个单元格内多余的换行后内容
                        if '' in all_text[-1]:
                            all_text.pop(-1)

                        row = len(all_text) - texts.count('\t')  # 获取行
                        column = int(len(all_text) / row - 1)  # 获取列
                        num = int(len(all_text) / row)

                        split_t = texts.replace('"', '').split('\t')  # 拆t
                        for i in range(0, row):  # 循环行数
                            split_n = split_t[column].split('\n')[:-1]  # 在一个单元格内有多余的换行符，需要截取有用的内容
                            join_n = '\n'.join(split_n)
                            last = split_t[column].split('\n')[-1]  # 截取最后的换行符内容
                            split_t[column] = join_n  # 修改列表内容
                            split_t.insert(column + 1, last)  # 插入列表内容
                            column += num

                        if '' in split_t[-1]:
                            split_t.pop(-1)  # 删除最后一个元素‘’

                        for r in range(row):  # 按行循环
                            values_list.append(split_t[0:num])  # 按列添加到列表
                            for c in range(num):  # 按列删除列表内容
                                split_t.pop(0)
                        if self.window.tableWidget_quota.currentRow() != -1:
                            for i, values in enumerate(values_list):  # 每次循环一行写入到表格里一行内容
                                if i == 0:
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow(),self.window.tableWidget_quota.currentColumn() + j,QTableWidgetItem(self.item1))
                                if i >= 1:
                                    self.window.tableWidget_quota.insertRow(
                                        self.window.tableWidget_quota.currentRow() + i)
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 0,
                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 1,
                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 2,
                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 3,
                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 4,
                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 5,
                                        QTableWidgetItem(''))
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i,
                                                                                self.window.tableWidget_quota.currentColumn() + j,
                                                                                QTableWidgetItem(self.item1))
                                self.window.tableWidget_quota.verticalHeader().resizeSection(self.window.tableWidget_quota.currentRow() + i, 50)
                    if '\t' not in texts and '"\n' in texts:# 复制整列内有单元格内有换行
                        sp_n = texts.split('"')  # 拆n
                        text_list = []
                        for sp in sp_n:
                            if sp != '':
                                if sp != '\n':
                                    a = sp.split('\n')
                                    if a[0] != '' and a[-1] != '':
                                        join_text = '\n'.join(a)
                                        text_list.append(str(join_text))

                                    if a[-1] == '':
                                        for b in a:
                                            if b != '':
                                                text_list.append(str(b))
                        if self.window.tableWidget_quota.currentRow() != -1:
                            for i, value in enumerate(text_list):  # 每次循环一行写入到表格里一行内容
                                if i ==0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow(),
                                        self.window.tableWidget_quota.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i>=1:
                                    self.window.tableWidget_quota.insertRow(self.window.tableWidget_quota.currentRow()+i)
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i,0,QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 1,QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 2,QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 3, QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow() + i, 4, QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 5,QTableWidgetItem(''))

                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow()+ i,
                                                         self.window.tableWidget_quota.currentColumn(),
                                                                            QTableWidgetItem(self.item1))
                                    self.window.tableWidget_quota.verticalHeader().resizeSection(self.window.tableWidget_quota.currentRow()+ i,50)

                    if '\t' not in texts and '"\n' not in texts:# 复制整列内有单元格内没有换行
                        sp_n = texts.split('\n')
                        if '' in sp_n[-1]:
                            sp_n.pop(-1)  # 删除最后一个元素‘’
                        if self.window.tableWidget_quota.currentRow() != -1:
                            for i, value in enumerate(sp_n):  # 每次循环一行写入到表格里一行内容
                                if i == 0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_quota.setItem(
                                        self.window.tableWidget_quota.currentRow(),
                                        self.window.tableWidget_quota.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i >= 1:
                                    self.window.tableWidget_quota.insertRow(self.window.tableWidget_quota.currentRow() + i)
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 0,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 1,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 2,QTableWidgetItem(''))


                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 3,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 4,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i, 5,
                                                                        QTableWidgetItem(''))
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_quota.setItem(self.window.tableWidget_quota.currentRow() + i,
                                                                        self.window.tableWidget_quota.currentColumn(),
                                                                        QTableWidgetItem(self.item1))
                                    self.window.tableWidget_quota.verticalHeader().resizeSection(self.window.tableWidget_quota.currentRow() + i, 50)


        except:
            pass

    def c_copy_clipboard(self):# 清单池剪切板复制外部数据
        try:
            if self.go == '剪切板':
                mimeData = self.clipboard_qd.mimeData()
                if mimeData.hasFormat('text/plain'):
                    all_text = []
                    test_list = []
                    values_list = []
                    texts = mimeData.text()
                    if '\t' in texts:
                        criteria = '"(.*?)"'
                        compile = re.compile(criteria, re.S)
                        self.results = compile.findall(texts)
                        sp_n = texts.replace('"', '').split('\n')  # 拆n
                        for sp in sp_n:
                            sp_t = str(sp).split('\t')  # 拆t
                            test_list.append(sp_t)
                        for new_list in test_list:
                            for new in new_list:
                                all_text.append(str(new))  # 构建新列表

                        for result in self.results:
                            res = str(result).split('\n')[1:]
                            for rs in res:
                                all_text.remove(str(rs))  # 删除新列表同一个单元格内多余的换行后内容
                        if '' in all_text[-1]:
                            all_text.pop(-1)

                        row = len(all_text) - texts.count('\t')  # 获取行
                        column = int(len(all_text) / row - 1)  # 获取列
                        num = int(len(all_text) / row)

                        split_t = texts.replace('"', '').split('\t')  # 拆t
                        for i in range(0, row):  # 循环行数
                            split_n = split_t[column].split('\n')[:-1]  # 在一个单元格内有多余的换行符，需要截取有用的内容
                            join_n = '\n'.join(split_n)
                            last = split_t[column].split('\n')[-1]  # 截取最后的换行符内容
                            split_t[column] = join_n  # 修改列表内容
                            split_t.insert(column + 1, last)  # 插入列表内容
                            column += num

                        if '' in split_t[-1]:
                            split_t.pop(-1)  # 删除最后一个元素‘’

                        for r in range(row):  # 按行循环
                            values_list.append(split_t[0:num])  # 按列添加到列表
                            for c in range(num):  # 按列删除列表内容
                                split_t.pop(0)
                        # self.qdc_text_list.remove(self.window.tableWidget_qdc.currentRow())  # 删除指定元素
                        # self.window.tableWidget_qdc.removeRow(self.window.tableWidget_qdc.currentColumn())
                        if self.window.tableWidget_qdc.currentRow()!=-1:
                            for i, values in enumerate(values_list):  # 每次循环一行写入到表格里一行内容
                                if i ==0:
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        if self.window.tableWidget_qdc.currentColumn() + j == 3:
                                            self.qdc_text_list[self.window.tableWidget_qdc.currentRow()].setPlainText(self.item1)
                                        if self.window.tableWidget_qdc.currentColumn() + j != 3:
                                            self.window.tableWidget_qdc.setItem(
                                                self.window.tableWidget_qdc.currentRow(),
                                                self.window.tableWidget_qdc.currentColumn() + j,
                                                QTableWidgetItem(self.item1))
                                if i>=1:
                                    self.window.tableWidget_qdc.insertRow(self.window.tableWidget_qdc.currentRow()+i)
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i,0,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 1,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 2,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 4,QTableWidgetItem(''))

                                    self.qdc_text = QPlainTextEdit()
                                    self.window.tableWidget_qdc.setCellWidget(self.window.tableWidget_qdc.currentRow() + i,3, self.qdc_text)
                                    self.qdc_text_list.insert(self.window.tableWidget_qdc.currentRow() + i,self.qdc_text)
                                    for j, value in enumerate(values):
                                        self.item1 = str(value).replace('None', '')
                                        if self.window.tableWidget_qdc.currentColumn() + j == 3:
                                            self.qdc_text_list[self.window.tableWidget_qdc.currentRow()+i].setPlainText(self.item1)
                                        if self.window.tableWidget_qdc.currentColumn() + j != 3:
                                            self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow()+ i,
                                                             self.window.tableWidget_qdc.currentColumn() + j,
                                                                                QTableWidgetItem(self.item1))
                                self.window.tableWidget_qdc.verticalHeader().resizeSection(self.window.tableWidget_qdc.currentRow()+ i,80)  # 调整每一行的大小为100像素

                    if '\t' not in texts and '"\n' in texts:# 复制整列内有单元格内有换行
                        sp_n = texts.split('"')  # 拆n
                        text_list = []
                        for sp in sp_n:
                            if sp != '':
                                if sp != '\n':
                                    a = sp.split('\n')
                                    if a[0] != '' and a[-1] != '':
                                        join_text = '\n'.join(a)
                                        text_list.append(str(join_text))

                                    if a[-1] == '':
                                        for b in a:
                                            if b != '':
                                                text_list.append(str(b))
                        if self.window.tableWidget_qdc.currentRow() != -1:
                            for i, value in enumerate(text_list):  # 每次循环一行写入到表格里一行内容
                                if i ==0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdc.setItem(
                                        self.window.tableWidget_qdc.currentRow(),
                                        self.window.tableWidget_qdc.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i>=1:
                                    self.window.tableWidget_qdc.insertRow(self.window.tableWidget_qdc.currentRow()+i)
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i,0,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 1,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 2,QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 4,QTableWidgetItem(''))

                                    self.qdc_text = QPlainTextEdit()
                                    self.window.tableWidget_qdc.setCellWidget(self.window.tableWidget_qdc.currentRow() + i,3, self.qdc_text)
                                    self.qdc_text_list.insert(self.window.tableWidget_qdc.currentRow() + i,self.qdc_text)
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow()+ i,
                                                         self.window.tableWidget_qdc.currentColumn(),
                                                                            QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdc.verticalHeader().resizeSection(self.window.tableWidget_qdc.currentRow()+ i,80)  # 调整每一行的大小为100像素


                    if '\t' not in texts and '"\n' not in texts:# 复制整列内有单元格内没有换行
                        sp_n = texts.split('\n')
                        if '' in sp_n[-1]:
                            sp_n.pop(-1)  # 删除最后一个元素‘’
                        if self.window.tableWidget_qdc.currentRow() != -1:
                            for i, value in enumerate(sp_n):  # 每次循环一行写入到表格里一行内容
                                if i == 0:
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdc.setItem(
                                        self.window.tableWidget_qdc.currentRow(),
                                        self.window.tableWidget_qdc.currentColumn(),
                                        QTableWidgetItem(self.item1))
                                if i >= 1:
                                    self.window.tableWidget_qdc.insertRow(self.window.tableWidget_qdc.currentRow() + i)
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 0,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 1,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 2,
                                                                        QTableWidgetItem(''))
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i, 4,
                                                                        QTableWidgetItem(''))

                                    self.qdc_text = QPlainTextEdit()
                                    self.window.tableWidget_qdc.setCellWidget(
                                        self.window.tableWidget_qdc.currentRow() + i, 3, self.qdc_text)
                                    self.qdc_text_list.insert(self.window.tableWidget_qdc.currentRow() + i,
                                                              self.qdc_text)
                                    self.item1 = str(value).replace('None', '')
                                    self.window.tableWidget_qdc.setItem(self.window.tableWidget_qdc.currentRow() + i,
                                                                        self.window.tableWidget_qdc.currentColumn(),
                                                                        QTableWidgetItem(self.item1))
                                    self.window.tableWidget_qdc.verticalHeader().resizeSection(
                                        self.window.tableWidget_qdc.currentRow() + i, 100)

        except:
            pass
    def quota_insert_line(self):#定额库插入行
        self.row = self.window.tableWidget_quota.currentRow()  # 获取单元格行数
        colums = self.window.tableWidget_quota.columnCount()
        if self.row == -1:
            row_quota = self.window.treeWidget_quota.currentIndex().row()  # 获取行
            if int(row_quota) != -1:
                self.item = self.window.treeWidget_quota.currentItem()
                if self.item.parent():
                    self.second = self.item.parent()
                    if self.second.parent():
                        self.fist = self.second.parent()
                        print(self.fist)
                        self.window.tableWidget_quota.insertRow(0)
                        for j in range(0, colums):
                            self.item1 = ''
                            self.window.tableWidget_quota.setItem(0, int(j), QTableWidgetItem(self.item1))
                        self.window.tableWidget_quota.verticalHeader().resizeSection(0, 50)
        if self.row != -1:
            self.window.tableWidget_quota.insertRow(self.row + 1)
            for j in range(0, colums):
                self.item1 = ''
                self.window.tableWidget_quota.setItem(self.row + 1, int(j), QTableWidgetItem(self.item1))
            self.window.tableWidget_quota.verticalHeader().resizeSection(self.row + 1, 50)

    def c_insert_line(self):#插入行
        self.row =self.window.tableWidget_qdc.currentRow()  # 获取单元格行数
        colums=self.window.tableWidget_qdc.columnCount()
        if self.row==-1:
            row_c = self.window.treeWidget_qdc.currentIndex().row()  # 获取行
            if int(row_c) != -1:
                self.item = self.window.treeWidget_qdc.currentItem()
                if self.item.parent():
                    self.second = self.item.parent()
                    if self.second.parent():
                        self.fist = self.second.parent()
                        self.window.tableWidget_qdc.insertRow(0)
                        for j in range(0,colums):
                            self.item1 =''
                            if j == 3:
                                self.qdc_text = QPlainTextEdit()
                                self.qdc_text.setPlainText(self.item1)
                                self.window.tableWidget_qdc.setCellWidget(0, int(j), self.qdc_text)
                                self.qdc_text_list.insert(0,self.qdc_text)
                            if j != 3:
                                self.window.tableWidget_qdc.setItem(0, int(j), QTableWidgetItem(self.item1))
                        self.window.tableWidget_qdc.verticalHeader().resizeSection(0,100)
        if self.row!=-1:
            self.window.tableWidget_qdc.insertRow(self.row+1)
            for j in range(0,colums):
                self.item1 =''
                if j == 3:
                    self.qdc_text = QPlainTextEdit()
                    self.qdc_text.setPlainText(self.item1)
                    self.window.tableWidget_qdc.setCellWidget(self.row+1, int(j), self.qdc_text)
                    self.qdc_text_list.insert(self.row+1,self.qdc_text)
                if j != 3:
                    self.window.tableWidget_qdc.setItem(self.row+1, int(j), QTableWidgetItem(self.item1))
            self.window.tableWidget_qdc.verticalHeader().resizeSection(self.row+1,100)

    def quota_delrow(self):
        self.row = self.window.tableWidget_quota.currentRow()  # 获取单元格行数
        if self.row != -1:
            self.rows = []
            self.column = []
            mytable = self.window.tableWidget_quota.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                # print(self.window.tableWidget_qdc.item(r.row(), r.column()).text())
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            print(self.column)
            num = 0
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column:
                    self.window.tableWidget_quota.removeRow(mytable[num].row())
                    num += 6
    def c_delrow(self):#删除行qdc_text_list
        self.row = self.window.tableWidget_qdc.currentRow()  # 获取单元格行数
        if self.row != -1:
            self.rows = []
            self.column = []
            mytable = self.window.tableWidget_qdc.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                # print(self.window.tableWidget_qdc.item(r.row(), r.column()).text())
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            print(self.column)
            num = 0
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column:
                    del_text = self.qdc_text_list[mytable[num].row()]
                    self.qdc_text_list.remove(del_text)  # 删除指定元素
                    self.window.tableWidget_qdc.removeRow(mytable[num].row())
                    num+=4
    c_values = []
    def c_copy(self):  # 复制行
        self.c_values = []
        self.tableWidget_allrows = int(self.window.tableWidget_qdc.rowCount())  # 获取总行数
        self.table_column = int(self.window.tableWidget_qdc.columnCount())
        if self.tableWidget_allrows != -1:
            self.rows = []
            self.column = []
            mytable = self.window.tableWidget_qdc.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            for self.Single_rows in self.rows:
                text_list = []
                self.c_values.append(text_list)
                for self.Single_colum in range(0, self.table_column):
                    if self.Single_colum != 3:
                        copytext = self.window.tableWidget_qdc.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(copytext)
                    if self.Single_colum == 3:
                        self.text3 = str(self.qdc_text_list[self.Single_rows].toPlainText())
                        text_list.append(self.text3)
            print(self.c_values)
    def c_cut(self):#剪切行
        self.c_values = []
        self.tableWidget_allrows = int(self.window.tableWidget_qdc.rowCount())  # 获取总行数
        self.table_column = int(self.window.tableWidget_qdc.columnCount())
        if self.tableWidget_allrows != -1:
            self.rows = []
            self.column = []
            mytable =self.window.tableWidget_qdc.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            for self.Single_rows in self.rows:
                text_list = []
                self.c_values.append(text_list)
                for self.Single_colum in range(0, self.table_column):
                    if self.Single_colum!=3:
                        copytext =self.window.tableWidget_qdc.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(copytext)
                    if self.Single_colum == 3:
                        self.text3 = str(self.qdc_text_list[self.Single_rows].toPlainText())
                        text_list.append(self.text3)
            mytable = self.window.tableWidget_qdc.selectedItems()
            num = 0
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column:
                    del_text = self.qdc_text_list[mytable[num].row()]
                    self.qdc_text_list.remove(del_text)  # 删除指定元素
                    self.window.tableWidget_qdc.removeRow(mytable[num].row())
                    num+=4

    def c_stickup_rows(self):  # 粘贴行
        if self.c_values != []:
            self.row = self.window.tableWidget_qdc.currentRow()  # 获取单元格行数
            if self.row != -1:
                for insert_row in range(0, len(self.rows)):
                    self.window.tableWidget_qdc.insertRow(self.row + insert_row + 1)
                    for j in range(0, self.window.tableWidget_qdc.columnCount()):
                        item = QTableWidgetItem(self.c_values[insert_row][j])
                        if j == 3:
                            self.qdc_text = QPlainTextEdit()
                            self.qdc_text.setPlainText(self.c_values[insert_row][j])
                            self.window.tableWidget_qdc.setCellWidget(self.row + insert_row + 1, j, self.qdc_text)
                            self.qdc_text_list.insert(self.row + insert_row + 1, self.qdc_text)
                        if j != 3:
                            self.window.tableWidget_qdc.setItem(self.row + insert_row + 1, j, item)
                    self.window.tableWidget_qdc.verticalHeader().resizeSection(self.row + insert_row + 1, 100)  # 调整每一行的大小为100像素

    def k_insert_line(self):#插入行
        self.row =self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
        colums=self.window.tableWidget_qdk.columnCount()
        if self.row==-1:
            row_k = self.window.treeWidget_qdk.currentIndex().row()  # 获取行
            if int(row_k) != -1:
                self.item = self.window.treeWidget_qdk.currentItem()
                if self.item.parent():
                    self.second = self.item.parent()
                    if self.second.parent():
                        self.fist = self.second.parent()
                        print(self.fist)
                        self.window.tableWidget_qdk.insertRow(0)
                        for j in range(0,colums):
                            self.item1 =''
                            if j == 3:
                                self.qdk_text = QPlainTextEdit()
                                self.qdk_text.setPlainText(self.item1)
                                self.window.tableWidget_qdk.setCellWidget(0, int(j), self.qdk_text)
                                self.qp_text_list.insert(0,self.qdk_text)
                            if j != 3:
                                self.window.tableWidget_qdk.setItem(0, int(j), QTableWidgetItem(self.item1))
                        self.window.tableWidget_qdk.verticalHeader().resizeSection(0,100)
        if self.row!=-1:
            self.window.tableWidget_qdk.insertRow(self.row+1)
            for j in range(0,colums):
                self.item1 = ''
                if j == 3:
                    self.qdk_text = QPlainTextEdit()
                    self.qdk_text.setPlainText(self.item1)
                    self.window.tableWidget_qdk.setCellWidget(self.row + 1, int(j), self.qdk_text)
                    self.qp_text_list.insert(self.row+1,self.qdk_text)
                if j != 3:
                    self.window.tableWidget_qdk.setItem(self.row + 1, int(j), QTableWidgetItem(self.item1))
            self.window.tableWidget_qdk.verticalHeader().resizeSection(self.row+1,100)

    def k_delrow(self):#删除行
        self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
        if self.row != -1:
            self.rows = []
            self.column = []
            mytable = self.window.tableWidget_qdk.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                # print(self.window.tableWidget_qdk.item(r.row(), r.column()).text())
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                    print(r.column())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            num = 0
            print(self.column)
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column:
                    del_text = self.qp_text_list[mytable[num].row()]
                    self.qp_text_list.remove(del_text)  # 删除指定元素
                    self.window.tableWidget_qdk.removeRow(mytable[num].row())
                    num+=4


    k_values = []
    def k_copy(self):#复制行
        self.k_values = []
        self.tableWidget_allrows = int(self.window.tableWidget_qdk.rowCount())  # 获取总行数
        self.table_column = int(self.window.tableWidget_qdk.columnCount())
        if self.tableWidget_allrows != -1:
            self.rows = []
            self.column = []
            mytable = self.window.tableWidget_qdk.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            for self.Single_rows in self.rows:
                text_list = []
                self.k_values.append(text_list)
                for self.Single_colum in range(0, self.table_column):
                    if self.Single_colum != 3:
                        copytext = self.window.tableWidget_qdk.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(copytext)
                    if self.Single_colum == 3:
                        self.text3 = str(self.qp_text_list[self.Single_rows].toPlainText())
                        text_list.append(self.text3)
    def k_cut(self):#剪切行
        self.k_values = []
        self.tableWidget_allrows = int(self.window.tableWidget_qdk.rowCount())  # 获取总行数
        self.table_column = int(self.window.tableWidget_qdk.columnCount())
        if self.tableWidget_allrows != -1:
            self.rows = []
            self.column = []
            mytable =self.window.tableWidget_qdk.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            for self.Single_rows in self.rows:
                text_list = []
                self.k_values.append(text_list)
                for self.Single_colum in range(0, self.table_column):
                    if self.Single_colum!=3:
                        copytext =self.window.tableWidget_qdk.item(self.Single_rows, self.Single_colum).text()
                        text_list.append(copytext)
                    if self.Single_colum == 3:
                        self.text3 = str(self.qp_text_list[self.Single_rows].toPlainText())
                        text_list.append(self.text3)
            mytable = self.window.tableWidget_qdk.selectedItems()
            num = 0
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column:
                    del_text = self.qp_text_list[mytable[num].row()]
                    print(del_text.toPlainText())
                    self.qp_text_list.remove(del_text)  # 删除指定元素
                    self.window.tableWidget_qdk.removeRow(mytable[num].row())
                    num +=4

    def k_stickup_rows(self):# 粘贴行
        if self.k_values!= []:
            self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
            if self.row != -1:
                for insert_row in range(0, len(self.rows)):
                    self.window.tableWidget_qdk.insertRow(self.row + insert_row + 1)
                    for j in range(0, self.window.tableWidget_qdk.columnCount()):
                        item = QTableWidgetItem(self.k_values[insert_row][j])
                        if j == 3:
                            self.qdk_text = QPlainTextEdit()
                            self.qdk_text.setPlainText(self.k_values[insert_row][j])
                            self.window.tableWidget_qdk.setCellWidget(self.row + insert_row + 1, j, self.qdk_text)
                            self.qp_text_list.insert(self.row+insert_row+1, self.qdk_text)
                        if j!= 3:
                            self.window.tableWidget_qdk.setItem(self.row + insert_row + 1, j, item)
                    self.window.tableWidget_qdk.verticalHeader().resizeSection(self.row + insert_row + 1, 100)  # 调整每一行的大小为100像素
            # self.over='多线程结束'

    def startThread_run_features(self):# 显示项目特征
        # try:
        self.Item_rowk = self.window.treeWidget_qdk.currentIndex().row()  # 获取行
        if int(self.Item_rowk) != -1:
            self.window.tableWidget_features.clearContents()  # 可以清除表格所有的内容
            self.window.tableWidget_features.setRowCount(0)
            self.row = self.window.tableWidget_qdk.currentRow()  # 获取单元格行数
            if self.row != -1:
                text_2 = self.window.tableWidget_qdk.item(self.row, 2).text()  # 获取单元格内容
                text_4 = self.window.tableWidget_qdk.item(self.row, 4).text()  # 获取单元格内容
                num=0
                for rows,values in enumerate(self.all_qdc_list):
                    if text_2==values[2] and text_4==values[4]:
                        num+=1
                        self.window.tableWidget_features.setRowCount(num)  # 设置行数
                        item=QTableWidgetItem(str(values[3]))
                        item.setFlags(Qt.ItemIsEnabled)
                        self.window.tableWidget_features.setItem(num-1, 0, item)
                        self.window.tableWidget_features.verticalHeader().resizeSection(num-1, 100)  # 调整每一行的大小为100像素
                self.window.tableWidget_features.setStyleSheet("gridline-color: rgb(257, 1, 0)")


    def Double_features(self):# 双击项目特征写入到清单库
        self.row = self.window.tableWidget_features.currentRow()  # 获取单元格行数
        if self.row != -1:
            text_0 = self.window.tableWidget_features.item(self.row, 0).text()  # 获取单元格内容
            write_row = self.window.tableWidget_qdk.currentRow()
            self.qp_text_list[write_row].setPlainText(text_0)

    def startThread_run_Search_quota(self):#搜索定额库
        self.text = self.window.lineEdit_Search_quota.text()  # 获取文本内容
        self.window.tableWidget_quota.clearContents()#可以清除表格所有的内容
        self.window.tableWidget_quota.setRowCount(0)  # 设置行数
        if self.text!='':
            self.city_name = self.window.comboBox_2.currentText()  # 获取当前选中的选项的文本
            files_de = self.files_dek+self.city_name
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work_qd(self.text,files_de)  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_Search_quota)
            self.Mywork.stopSing.connect(self.stopThread_Search_quota)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
            self.window.pushButton_save_quota.setEnabled(False)
            # self.window.pushButton_parse_quota.setEnabled(False)

        else:
            return
    def stopThread_Search_quota(self,all_text_list):
        self.thread.quit()  # 退出
        self.Search_quota(all_text_list)
    def Search_quota(self,all_text_list):# 搜索定额库写入表格
        for rows,values in enumerate(all_text_list):
            self.window.tableWidget_quota.setRowCount(rows+1)  # 设置行数
            for colums,value in enumerate(values):
                self.item1 =str(value).replace('None', '')
                self.window.tableWidget_quota.setItem(rows,colums,QTableWidgetItem(self.item1))
                self.window.tableWidget_quota.verticalHeader().resizeSection(rows, 50)  # 调整每一行的大小为100像素
        self.window.tableWidget_quota.setStyleSheet("gridline-color: rgb(257, 1, 0)")
        self.window.tableWidget_quota.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive) # 设置列宽，列宽可调
        self.window.tableWidget_quota.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
        self.window.tableWidget_quota.viewport().update()  # 刷新tab内容

    def startThread_run_Search_qdk(self):#搜索清单库
        self.text = self.window.lineEdit_Search_qdk.text()  # 获取文本内容
        self.window.tableWidget_qdk.clearContents()#可以清除表格所有的内容
        self.window.tableWidget_qdk.setRowCount(0)  # 设置行数
        if self.text!='':
            self.qp_text_list = []
            self.city_name = self.window.comboBox.currentText()  # 获取当前选中的选项的文本
            files_qdk = self.files_qdk+self.city_name
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work_qd(self.text,files_qdk)  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_Search_qdk)
            self.Mywork.stopSing.connect(self.stopThread_Search_qdk)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
            self.window.pushButton_save_qdk.setEnabled(False)
        else:
            return
    def stopThread_Search_qdk(self,all_text_list):
        self.thread.quit()  # 退出
        self.Search_qdk(all_text_list)
    def Search_qdk(self,all_text_list):# 筛选清单库
        for rows,values in enumerate(all_text_list):
            self.window.tableWidget_qdk.setRowCount(rows+1)  # 设置行数
            for colums,value in enumerate(values):
                self.item1 =str(value).replace('None', '')
                if colums ==3:
                    self.qp_text = QPlainTextEdit()
                    self.qp_text.setPlainText(self.item1)
                    self.qp_text_list.append(self.qp_text)
                    self.window.tableWidget_qdk.setCellWidget(rows,colums, self.qp_text)
                if colums !=3:
                    self.window.tableWidget_qdk.setItem(rows,colums,QTableWidgetItem(self.item1))
                self.window.tableWidget_qdk.verticalHeader().resizeSection(rows, 100)  # 调整每一行的大小为100像素
        self.window.tableWidget_qdk.setStyleSheet("gridline-color: rgb(257, 1, 0)")
        self.window.tableWidget_qdk.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive) # 设置列宽，列宽可调
        self.window.tableWidget_qdk.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
        self.window.tableWidget_qdk.horizontalHeader().resizeSection(4, 100)  # 调整第2列的大小为500像素
        self.window.tableWidget_qdk.viewport().update()  # 刷新tab内容

    def startThread_run_Search_qdc(self):# 搜索清单池
        self.text = self.window.lineEdit_Search_qdc.text()  # 获取文本内容
        self.window.tableWidget_qdc.clearContents()#可以清除表格所有的内容
        self.window.tableWidget_qdc.setRowCount(0)  # 设置行数
        if self.text!='':
            self.qdc_text_list = []
            BASE_DIR = os.path.dirname(__file__)  # 清单池
            files_qdc=BASE_DIR+'/'+"清单数据库"+'/'+"清单池"
            self.thread = QThread()  # 实例化一个线程
            self.Mywork = Work_qd(self.text,files_qdc)  # 实例化工作类,并传入参数到工作线程进行加工
            self.Mywork.moveToThread(self.thread)  # 把工作移动到新线程里工作
            self.thread.started.connect(self.Mywork.run_Search_qdc)
            self.Mywork.stopSing.connect(self.stopThread_Search_qdc)  # 停止信号连接到stopThread方法
            self.thread.start()  # 开始线程的运行
            self.window.pushButton_save_qdc.setEnabled(False)
        else:
            return
    def stopThread_Search_qdc(self,all_text_list):
        self.thread.quit()  # 退出
        self.Search_qdc(all_text_list)

    def Search_qdc(self,all_text_list):# 筛选清单池
        for rows,values in enumerate(all_text_list):
            self.window.tableWidget_qdc.setRowCount(rows+1)  # 设置行数
            for colums,value in enumerate(values):
                self.item1 =str(value).replace('None', '')
                if colums ==3:
                    self.qdc_text = QPlainTextEdit()
                    self.qdc_text.setPlainText(self.item1)
                    self.qdc_text_list.append(self.qdc_text)
                    self.window.tableWidget_qdc.setCellWidget(rows,colums, self.qdc_text)
                if colums !=3:
                    self.window.tableWidget_qdc.setItem(rows,colums,QTableWidgetItem(self.item1))
                self.window.tableWidget_qdc.verticalHeader().resizeSection(rows, 100)  # 调整每一行的大小为100像素
        self.window.tableWidget_qdc.setStyleSheet("gridline-color: rgb(257, 1, 0)")
        self.window.tableWidget_qdc.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive) # 设置列宽，列宽可调
        self.window.tableWidget_qdc.horizontalHeader().resizeSection(3, 250)  # 调整第2列的大小为500像素
        self.window.tableWidget_qdc.horizontalHeader().resizeSection(4, 100)  # 调整第2列的大小为500像素
        self.window.tableWidget_qdc.viewport().update()  # 刷新tab内容

    def adaptation_rows(self):#自适应行高
        self.Item_row =win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.rows =self.window.tableWidget_qdk.rowCount()  # 获取单元格行数
            if self.rows != 0:
                for i in range(0, self.rows):
                    self.window.tableWidget_qdk.verticalHeader().setSectionResizeMode(i,QHeaderView.ResizeToContents)

class pdf_tool_windows():#PDF窗口
    def __init__(self):
        super(pdf_tool_windows).__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "pdf_tool_windows.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('PDF编辑工具窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.4)
        self.width = int(self.screenwidth * 0.5)
        self.window.resize(self.width, self.height)
        self.window.treeWidget_1.setColumnWidth(0, 500)  # 0列列宽
        self.window.treeWidget_1.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_1.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_2.setColumnWidth(0, 500)  # 0列列宽
        self.window.treeWidget_2.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_2.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_3.setColumnWidth(0, 500)  # 0列列宽
        self.window.treeWidget_3.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_3.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_4.setColumnWidth(0, 400)  # 0列列宽
        self.window.treeWidget_4.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_4.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_4.setColumnWidth(3, 100)  # 0列列宽
        self.window.treeWidget_5.setColumnWidth(0, 500)  # 0列列宽
        self.window.treeWidget_5.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_5.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_6.setColumnWidth(0, 500)  # 0列列宽
        self.window.treeWidget_6.setColumnWidth(1, 100)  # 0列列宽
        self.window.treeWidget_6.setColumnWidth(2, 100)  # 0列列宽
        self.window.treeWidget_7.setColumnWidth(0, 600)  # 0列列宽
        self.window.treeWidget_7.setColumnWidth(1, 100)  # 0列列宽
        self.window.setWindowFlags(Qt.WindowCloseButtonHint) # 没有最大化
        self.window.pushButton_j.clicked.connect(self.show1)
        self.window.pushButton_f.clicked.connect(self.show2)
        self.window.pushButton_s.clicked.connect(self.show3)
        self.window.pushButton_x.clicked.connect(self.show4)
        self.window.pushButton_t.clicked.connect(self.show5)
        self.window.pushButton_w.clicked.connect(self.show6)
        self.window.pushButton_d.clicked.connect(self.show7)
        self.window.pushButton_file_1.clicked.connect(self.file1)#合并PDF
        self.window.pushButton_save_1.clicked.connect(self.save1)#合并PDF
        self.window.treeWidget_1.itemChanged.connect(self.item_change1)#合并PDF
        self.window.pushButton_file_2.clicked.connect(self.file2)#拆分PDF
        self.window.pushButton_save_2.clicked.connect(self.split_pdf)  # 拆分PDF
        self.window.pushButton_file_3.clicked.connect(self.file3)  # 删除pdf
        self.window.pushButton_save_3.clicked.connect(self.del_pdf)  # 删除pdf
        self.window.pushButton_file_4.clicked.connect(self.file4)  # 旋转pdf页
        self.window.pushButton_save_4.clicked.connect(self.rotating_pdf)  #  旋转pdf页
        self.window.pushButton_file_5.clicked.connect(self.file5)  # 旋转pdf页
        self.window.pushButton_save_5.clicked.connect(self.text_pdf)  #  旋转pdf页
        self.window.pushButton_file_6.clicked.connect(self.file6)  # pdf砖Word
        self.window.pushButton_save_6.clicked.connect(self.pdf_word)  #  pdf砖Word
        self.window.pushButton_file_7.clicked.connect(self.file7)  # pdf砖Word
        self.window.pushButton_save_7.clicked.connect(self.minish_pdf)  #  pdf砖Word

    def show1(self):
        self.window.stackedWidget.setCurrentIndex(0)
    def show2(self):
        self.window.stackedWidget.setCurrentIndex(1)
    def show3(self):
        self.window.stackedWidget.setCurrentIndex(2)
    def show4(self):
        self.window.stackedWidget.setCurrentIndex(3)
    def show5(self):
        self.window.stackedWidget.setCurrentIndex(4)
    def show6(self):
        self.window.stackedWidget.setCurrentIndex(5)
    def show7(self):
        self.window.stackedWidget.setCurrentIndex(6)

    def item_change1(self,item, column):
        if column==1:
            value = item.text(column).isdigit()
            if item.text(column) == '最后页插入':
                return
            elif item.text(column)!='最后页插入' and value==False:
                QMessageBox.information(self.window, '温馨提示', '请输入整数或最后页插入')
                print(value)
        if column==2:
            if item.text(column) == '所有页':
                return
            elif item.text(column)!='所有页':
                it = str(item.text(column)).split('-')
                if len(it)!=2:
                    QMessageBox.information(self.window, '温馨提示', '请按格式输入‘起始页-结束页’或所有页')
                if len(it) == 2:
                    if it[0]>it[1]:
                        QMessageBox.information(self.window, '温馨提示', '起始页应<=结束页')
                        return
                    for value in it:
                        if value.isdigit()==False:
                            QMessageBox.information(self.window, '温馨提示', '请按格式输入‘起始页-结束页’或所有页')


    def file1(self):
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_1.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                self.son = QTreeWidgetItem(self.window.treeWidget_1.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                criteria = '(.*)/(.*)'
                compile = re.compile(criteria, re.S)
                self.results = compile.findall(file)
                self.son.setText(0, str(file))
                self.son.setText(1, '最后页插入')
                self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_1.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_1.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_1.expandAll()
    def save1(self):
        try:
            files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存","未命名", '*.pdf')
            if files_address == "":
                return
            if '/' in files_address:
                if files_address != "":
                    ite = self.window.treeWidget_1.topLevelItem(0)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    merger = PdfWriter()
                    exp = Exception('请输入数字')
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        if string.checkState(0) == Qt.Checked:
                            print(string.text(1))
                            if string.text(1)!='最后页插入':
                                if string.text(1).isdigit()==False:
                                    raise exp
                            if string.text(1) == '0':
                                raise exp
                            it = str(string.text(2)).split('-')
                            if string.text(2)!='所有页':
                                if len(it)!=2:
                                    raise exp
                                for value in it:
                                    if value=='0':
                                        raise exp
                                    if value.isdigit() == False:
                                        print(exp)
                                        raise exp

                            if str(string.text(1))=='最后页插入' and str(string.text(2))=='所有页':
                                merger.append(fileobj=str(string.text(0)))
                            elif str(string.text(1))=='最后页插入' and str(string.text(2))!='所有页' and '-' in str(string.text(2)):
                                merger.append(fileobj=str(string.text(0)), pages=((int(str(string.text(2)).split('-')[0])-1,int(str(string.text(2)).split('-')[1]))))
                            elif str(string.text(1))!='最后页插入' and str(string.text(2))=='所有页':
                                merger.merge(position=int(str(string.text(1)))-1,fileobj=str(string.text(0)))
                            elif str(string.text(1))!='最后页插入' and str(string.text(2))!='所有页' and '-' in str(string.text(2)):
                                merger.merge(position=int(str(string.text(1)))-1,fileobj=str(string.text(0)), pages=((int(str(string.text(2)).split('-')[0])-1,int(str(string.text(2)).split('-')[1]))))
                            print('继续')
                    output = open(files_address, "wb")
                    merger.write(output)
                    merger.close()
                    output.close()
                    QMessageBox.information(self.window, '温馨提示', 'pdf文件合并完成，请检查')
        except:
            QMessageBox.information(self.window, '温馨提示', '格式输入错误，请检查')
            return
    def file2(self):#拆分PDF
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_2.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_2.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                # self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_2.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_2.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_2.expandAll()
    def split_pdf(self):#拆分PDF
        files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", "默认名称", '*.pdf')
        if files_address == "":
            return
        if '/' in files_address:
            if files_address != "":
                try:
                    exp = Exception('起始页不能为0')
                    ite = self.window.treeWidget_2.topLevelItem(0)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    criteria = '(.*)/'
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(files_address)
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        if string.checkState(0) == Qt.Checked:
                            # print(string.text(1))
                            criteria = '.*/(.*).pdf'
                            compile = re.compile(criteria, re.S)
                            name = compile.findall(string.text(0))
                            it = str(string.text(2)).split('|')
                            if '' in it:
                                it.remove('')
                            for sp_it in it:
                                # if str(sp_it).replace('-','').isdigit()==True:
                                if (str(sp_it).split('-')[0])=='0' or (str(sp_it).split('-')[1])=='0':
                                    raise exp

                                merger = PdfWriter()
                                merger.merge(position=0,fileobj=str(string.text(0)), pages=((int(str(sp_it).split('-')[0]) - 1, int(str(sp_it).split('-')[1]))))
                                output = open(os.path.join(self.results[0],name[0]+str(sp_it)+'.pdf'), "wb")
                                print(os.path.join(self.results[0],name[0]+str(sp_it)+'.pdf'))
                                merger.write(output)
                                merger.close()
                                output.close()

                except Exception as e: #  Exception捕获错误的类型，e保存具体错误内容
                    print('出现异常',e)
                    QMessageBox.information(self.window, '温馨提示', '格式错误,正确格式为:起始页-终止页|起始页-终止页。起始页应>=1')
                else:
                    QMessageBox.information(self.window, '温馨提示', 'pdf文件拆分完成，请检查')
    def file3(self):#删除PDF
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_3.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_3.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                # self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_3.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_3.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_3.expandAll()
    def del_pdf(self):#删除PDF
        files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", "默认名称", '*.pdf')
        if files_address == "":
            return
        if '/' in files_address:
            if files_address != "":
                try:
                    exp = Exception('起始页不能为0')
                    ite = self.window.treeWidget_3.topLevelItem(0)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    criteria = '(.*)/'
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(files_address)
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        if string.checkState(0) == Qt.Checked:
                            # print(string.text(1))
                            criteria = '.*/(.*).pdf'
                            compile = re.compile(criteria, re.S)
                            name = compile.findall(string.text(0))
                            it = str(string.text(2)).split('|')
                            if '' in it:
                                it.remove('')
                            pages_list=[]
                            reader = PdfReader(string.text(0))
                            merger = PdfWriter()
                            for page_num in range(0,int(string.text(1))):
                                pages_list.append(page_num+1)
                            if pages_list!=[]:
                                for remove_page in it:
                                    if '-' not in remove_page:
                                        pages=int(remove_page)
                                        pages_list.remove(pages)
                                    if '-' in remove_page:
                                        sp_page=str(remove_page).split('-')
                                        for page in range(int(sp_page[0]),int(sp_page[1])+1):
                                            print(page)
                                            page = int(page)
                                            pages_list.remove(page)
                                for page in pages_list:
                                    merger.add_page(reader.pages[page-1])
                            output = open(os.path.join(self.results[0],name[0]+'.pdf'), "wb")
                            merger.write(output)
                            merger.close()
                            output.close()
                except Exception as e: #  Exception捕获错误的类型，e保存具体错误内容
                    print('出现异常',e)
                    QMessageBox.information(self.window, '温馨提示', '格式错误,正确格式为:起始页-终止页|当前页。起始页应>=1')
                else:
                    QMessageBox.information(self.window, '温馨提示', 'pdf文件页删除完成，请检查')
    def file4(self):#选中PDF页
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_4.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_4.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                self.son.setText(3, '90')
                # self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_4.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_4.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_4.expandAll()
    def rotating_pdf(self):#旋转PDF
        files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", "默认名称", '*.pdf')
        if files_address == "":
            return
        if '/' in files_address:
            if files_address != "":
                try:
                    exp = Exception('起始页不能为0')
                    ite = self.window.treeWidget_4.topLevelItem(0)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    criteria = '(.*)/'
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(files_address)
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        if string.checkState(0) == Qt.Checked:
                            # print(string.text(1))
                            criteria = '.*/(.*).pdf'
                            compile = re.compile(criteria, re.S)
                            name = compile.findall(string.text(0))
                            it = str(string.text(2)).split('|')
                            print(22)
                            if '' in it:
                                it.remove('')
                            pages_list=[]
                            reader = PdfReader(string.text(0))
                            writer = PdfWriter()
                            for app_page in it:
                                if '-' not in app_page:
                                    pages=int(app_page)
                                    pages_list.append(pages)
                                if '-' in app_page:
                                    sp_page=str(app_page).split('-')
                                    for page in range(int(sp_page[0]),int(sp_page[1])+1):
                                        print(page)
                                        page = int(page)
                                        pages_list.append(page)
                            for page_num in range(0, int(string.text(1))):
                                writer.add_page(reader.pages[page_num])
                            if pages_list!=[]:
                                for page in pages_list:
                                    writer.pages[page-1].rotate(int(string.text(3)))
                                output = open(os.path.join(self.results[0],name[0]+'.pdf'), "wb")
                                writer.write(output)
                                writer.close()
                                output.close()
                except Exception as e: #  Exception捕获错误的类型，e保存具体错误内容
                    print('出现异常',e)
                    QMessageBox.information(self.window, '温馨提示', '格式错误,正确格式为:起始页-终止页|当前页。起始页应>=1')
                else:
                    QMessageBox.information(self.window, '温馨提示', 'pdf文件页旋转完成，请检查')
    def file5(self):#选中PDF页
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_5.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_5.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                # self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_5.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_5.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_5.expandAll()
    def text_pdf(self):
        try:
            exp = Exception('起始页不能为0')
            ite = self.window.treeWidget_5.topLevelItem(0)  # 循环获取根节点
            count = ite.childCount()  # 获取当前根节点的子节点数量
            for j in range(0, count):
                string = ite.child(j)  # 子节点的文字信息
                if string.checkState(0) == Qt.Checked:
                    print(string.text(1))
                    it = str(string.text(2)).split('|')
                    if '' in it:
                        it.remove('')
                    pages_list = []
                    print(11)
                    reader = PdfReader(string.text(0))
                    for app_page in it:
                        if '-' not in app_page:
                            pages = int(app_page)
                            pages_list.append(pages)
                        if '-' in app_page:
                            sp_page = str(app_page).split('-')
                            for page in range(int(sp_page[0]), int(sp_page[1]) + 1):
                                print(page)
                                page = int(page)
                                pages_list.append(page)
                    if pages_list != []:
                        for pages in pages_list:
                            page = reader.pages[pages-1]
                            page_text=page.extract_text()
                            self.window.textEdit.append(str(page_text))
                            print(page_text)
        except Exception as e:  # Exception捕获错误的类型，e保存具体错误内容
            print('出现异常', e)
            QMessageBox.information(self.window, '温馨提示', '格式错误,正确格式为:起始页-终止页|当前页。起始页应>=1')
    def file6(self):#选中PDF页
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_6.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_6.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                # self.son.setText(2, '所有页')
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_6.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_6.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_6.expandAll()
    def pdf_word(self):#拆分PDF
        files_address, filetype = QFileDialog.getSaveFileName(self.window, "文件保存", "默认名称", '*.pdf')
        if files_address == "":
            return
        if '/' in files_address:
            if files_address != "":
                try:
                    exp = Exception('起始页不能为0')
                    ite = self.window.treeWidget_6.topLevelItem(0)  # 循环获取根节点
                    count = ite.childCount()  # 获取当前根节点的子节点数量
                    criteria = '(.*)/'
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(files_address)
                    for j in range(0, count):
                        string = ite.child(j)  # 子节点的文字信息
                        if string.checkState(0) == Qt.Checked:
                            # print(string.text(1))
                            criteria = '.*/(.*).pdf'
                            compile = re.compile(criteria, re.S)
                            name = compile.findall(string.text(0))
                            pdf_file = string.text(0)
                            docx_file = os.path.join(self.results[0],name[0]+'.docx')
                            cv = Converter(pdf_file)
                            it = str(string.text(2)).split('|')
                            if '' in it:
                                it.remove('')
                            current_page_list=[]
                            for app_page in it:
                                if '-' not in app_page:
                                    pages= int(app_page)-1
                                    current_page_list.append(pages)
                                if '-' in app_page:
                                    sp_page = str(app_page).split('-')
                                    for page in range(int(sp_page[0]),int(sp_page[1]) + 1):
                                        page = int(page)-1
                                        current_page_list.append(page)
                            if current_page_list!=[]:
                                cv.convert(docx_file, pages=current_page_list)
                                cv.close()
                except Exception as e: #  Exception捕获错误的类型，e保存具体错误内容
                    print('出现异常',e)
                    QMessageBox.information(self.window, '温馨提示', '格式错误,正确格式为:起始页-终止页|当前页。起始页应>=1')
                else:
                    QMessageBox.information(self.window, '温馨提示', 'pdf文件转Word完成，请检查')

    def file7(self):#选中PDF页
        self.filePaths, _ = QFileDialog.getOpenFileNames(self.window, '选择文件', ' ', '文件类型(*.pdf)')
        if self.filePaths != []:
            self.window.treeWidget_7.topLevelItem(0).takeChildren()
            for file in (self.filePaths):
                reader = PdfReader(file)#读取PDF
                numPages = len(reader.pages)#获取总页数
                self.son = QTreeWidgetItem(self.window.treeWidget_7.topLevelItem(0))
                self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                self.son.setText(0, str(file))
                self.son.setText(1, str(numPages))
                self.son.setCheckState(0, Qt.Checked)
                self.window.treeWidget_7.topLevelItem(0).setCheckState(0, Qt.Checked)
                self.window.treeWidget_7.topLevelItem(0).setSizeHint(0, QSize(0, 30))
                self.son.setSizeHint(0, QSize(0, 30))
            self.window.treeWidget_7.expandAll()
    def minish_pdf(self):
        try:
            ite = self.window.treeWidget_7.topLevelItem(0)  # 循环获取根节点
            count = ite.childCount()  # 获取当前根节点的子节点数量
            for j in range(0, count):
                string = ite.child(j)  # 子节点的文字信息
                if string.checkState(0) == Qt.Checked:
                    print(string.text(1))
                    # PDF无损压缩
                    reader = PdfReader(string.text(0))
                    writer = PdfWriter()
                    for page in reader.pages:
                        writer.add_page(page)
                    writer.add_metadata(reader.metadata)
                    with open(string.text(0), "wb") as fp:
                        writer.write(fp)
        except Exception as e:  # Exception捕获错误的类型，e保存具体错误内容
            print('出现异常', e)
            QMessageBox.information(self.window, '温馨提示', '错误，请检查')
        else:
            QMessageBox.information(self.window, '温馨提示', '原pdf文件减小完成，请检查')


class f_r_window():#查找替换窗口
    def __init__(self):
        super(f_r_window, self).__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "find_replace_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('查找/替换/筛选窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.4)
        self.width = int(self.screenwidth * 0.4)
        self.window.resize(self.width, self.height)
        # self.window.setWindowFlags(Qt.WindowCloseButtonHint) # 没有最大化
        self.window.pushButton_find.clicked.connect(self.find_text)  # 查找文字
        self.window.lineEdit_find.setPlaceholderText('输入查找内容')#提示文本
        self.window.lineEdit_replace.setPlaceholderText('输入替换内容')#输入替换
        self.window.lineEdit_filter.setPlaceholderText('输入筛选条件，点击Enter')#输入替换
        self.window.lineEdit_find.returnPressed.connect(self.find_text)#绑定enter键
        self.window.lineEdit_filter.returnPressed.connect(self.check_filter)  # 绑定enter键
        self.window.pushButton_all_repalce.clicked.connect(self.all_replace)
        self.window.listWidget_find.itemSelectionChanged.connect(self.SelectionChange)
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint|Qt.WindowCloseButtonHint)#禁止最大化
        self.window.pushButton_repalce.clicked.connect(self.repalce)
        self.window.toolButton_clear_filter.clicked.connect(self.clear_filter)
        # self.window.comboBox_2.currentIndexChanged.connect(self.change_filter)
        self.window.comboBox_2.activated.connect(self.change_filter)#选中一个下拉选项时发射信号
        # self.window.comboBox_2.highlighted.connect(self.change_filter)  # 选中一个已经选中的下拉选项时发射信号
        self.window.toolButton_clear.clicked.connect(self.clear_up)
        self.trees=[self.window.treeWidget_0,self.window.treeWidget_1,self.window.treeWidget_2,self.window.treeWidget_3,self.window.treeWidget_4,self.window.treeWidget_5,self.window.treeWidget_6,self.window.treeWidget_7,self.window.treeWidget_8,self.window.treeWidget_9,self.window.treeWidget_10]
        self.all_dict = {'项目编码': 0, '项目名称': 1, '项目特征': 2, '项目单位': 3, '匹配清单名称': 4, '匹配定额名称': 5,
                    '专业系统': 6, '材料名称': 7, '材料规格': 8, '材料单位': 9, '工程量': 10}

        self.window.comboBox_2.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 80px;}")

        self.window.comboBox_2.setView(QListView())
        self.window.comboBox_1.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 80px;}")

        self.window.comboBox_1.setView(QListView())
    # def write1_tree(self):
        font = QFont()
        font.setPointSize(10)  # 设置字体大小为10像素
        font.setFamily("宋体")
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.row_list=[]
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                self.row_list.append(str(self.Single_rows))
            for tre in self.trees:
                tre.setColumnWidth(0, 500)  # 0列列宽
                self.root = QTreeWidgetItem(tre)
                self.root.setText(0, '全部')
                self.root.setText(1, '0')
                self.root.setFlags(
                    QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                self.root.setSizeHint(0, QSize(0, 40))
                self.root.setCheckState(0, Qt.Unchecked)
                self.root.setFont(0, font)
                tre.hideColumn(1)  # 隐藏第一列
                if self.row_list!=[]:
                    self.root.setText(1, str('+'.join(self.row_list)))
                    self.root.setSizeHint(1, QSize(0, 40))

    all_re_value=[]
    def find_text(self):#查找
        self.all_re_value = []
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.window.listWidget_find.clear()
            self.tableWidget_allrows = int(win.tableWidget.rowCount())  # 获取总行数
            self.table_column = int(win.tableWidget.columnCount())
            if self.tableWidget_allrows !=0:
                self.find_value = self.window.lineEdit_find.text()  # 获取文本内容
                self.rows = []
                self.column = []
                mytable = win.tableWidget.selectedItems()
                for r in mytable:  # 第三种方法获取值'
                    self.rows.append(r.row())
                    self.column.append(r.column())
                if len(self.rows)>=2:
                    for r in mytable:  # 第三种方法获取值'
                        table_text = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                        if self.find_value in table_text and self.find_value!='':
                            re_text = r.row(),r.column(), table_text
                            self.all_re_value.append(re_text)
                            value = '第{}行 {}列 {}'.format(r.row()+1,r.column()+1, table_text)
                            self.ql_text = QListWidgetItem()  # 创建QListWidgetItem实例
                            self.ql_text.setText(str(value))
                            self.window.listWidget_find.addItem(self.ql_text)  # 添加到列表控件中
                            self.ql_text.setSizeHint(QSize(0, 40))
                else:
                    for self.Single_row in range(0, self.tableWidget_allrows):
                        for self.Single_colum in range(0, self.table_column):
                            table_text = win.tableWidget.item(self.Single_row, self.Single_colum).text()
                            if self.find_value in table_text and self.find_value!='':
                                re_text = self.Single_row, self.Single_colum, table_text
                                self.all_re_value.append(re_text)
                                value='第{}行 {}列 {}'.format(self.Single_row+1,self.Single_colum+1,table_text)
                                self.ql_text = QListWidgetItem()  # 创建QListWidgetItem实例
                                self.ql_text.setText(str(value))
                                self.window.listWidget_find.addItem(self.ql_text)  # 添加到列表控件中
                                self.ql_text.setSizeHint(QSize(0, 40))
    def all_replace(self):#全部替换
        if self.all_re_value!=[]:
            win.tab_change='不提取'
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[win.tableWidget] = self.Click_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            # 替换
            self.re_v = self.window.lineEdit_replace.text()
            for re_value in self.all_re_value:
                re=str(re_value[-1]).replace(self.find_value,self.re_v)
                win.tableWidget.setItem(re_value[0],re_value[1], QTableWidgetItem(re))
            win.tableWidget.viewport().update()  # 刷新tab内容
            # 撤回
            if win.table_do == '执行':
                self.new_undo_dict = {}
                new_text_list = []
                self.new_undo_dict[win.tableWidget] = new_text_list
                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                    t_list = []
                    new_text_list.append(t_list)
                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.Click_list != new_text_list:
                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                    win.undoStack_del.push(command)
            win.tab_change = '提取'

    def SelectionChange(self):#点击list定位table位置
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tableWidget_allrows = int(win.tableWidget.rowCount())  # 获取总行数
            if self.tableWidget_allrows !=0:
                self.replace_text=self.window.listWidget_find.currentItem().text()
                criteria = '第(.*?)行(.*?)列(.*)'
                compile= re.compile(criteria, re.S)
                self.results =compile.findall(self.replace_text)
                location=win.tableWidget.item(int(self.results[0][0])-1,int(self.results[0][1])-1)
                win.tableWidget.scrollToItem(location)# 滚轮定位
    def repalce(self):
        if self.window.listWidget_find.currentIndex().row()!=-1:
            win.tab_change = '不提取'
            # 撤销
            self.old_undo_dict = {}
            self.Click_list = []
            self.old_undo_dict[win.tableWidget] = self.Click_list
            for self.Single_rows in range(0, win.tableWidget.rowCount()):
                t_list = []
                self.Click_list.append(t_list)
                for self.Single_colum in range(0, win.tableWidget.columnCount()):
                    self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                    t_list.append(self.text)
            # 替换
            find_value = self.window.lineEdit_find.text()  # 获取文本内容
            re_v = self.window.lineEdit_replace.text()
            re=str(self.results[0][-1]).replace(find_value,re_v)
            # print(int(self.results[0][0]),int(self.results[0][1]),re)
            win.tableWidget.setItem(int(self.results[0][0])-1,int(self.results[0][1])-1,QTableWidgetItem(re))
            win.tableWidget.viewport().update()  # 刷新tab内容
            # 撤回
            if win.table_do == '执行':
                self.new_undo_dict = {}
                new_text_list = []
                self.new_undo_dict[win.tableWidget] = new_text_list
                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                    t_list = []
                    new_text_list.append(t_list)
                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                if self.Click_list != new_text_list:
                    command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
                    win.undoStack_del.push(command)
            win.tab_change = '提取'

    def change_filter(self):# 筛选
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            font = QFont()
            font.setPointSize(9)  # 设置字体大小为10像素
            font.setFamily("宋体")
            com_text2=self.window.comboBox_2.currentText()
            if com_text2!='':
                self.trees[self.all_dict[com_text2]].topLevelItem(0).takeChildren()
                self.item = self.trees[self.all_dict[com_text2]].topLevelItem(0)
                self.item_1=self.item.text(1)#获取0列内容
                self.values_dict={}
                if self.item_1!='0':
                    self.window.stackedWidget.setCurrentIndex(self.all_dict[com_text2])  # 设置显示页面
                    for self.Single_rows in str(self.item_1).split('+'):
                        for self.Single_colum in range(0, win.tableWidget.columnCount()):
                            if self.Single_colum==self.all_dict[com_text2]:
                                self.text = win.tableWidget.item(int(self.Single_rows), self.Single_colum).text()
                                self.values_dict[str(self.text)]=self.all_dict[com_text2]

                if self.values_dict!={}:
                    for value in self.values_dict.keys():
                        self.son = QTreeWidgetItem(self.trees[self.all_dict[com_text2]].topLevelItem(0))
                        self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
                        self.son.setText(0, value)
                        if com_text2=='项目特征':
                            self.son.setSizeHint(0, QSize(0, 100))
                        else:
                            self.son.setSizeHint(0, QSize(0, 25))
                        self.son.setCheckState(0, Qt.Unchecked)
                        self.son.setFont(0, font)
                    self.trees[self.all_dict[com_text2]].expandAll()


    def check_filter(self):  # 提取材料名称
        com_text1 = self.window.comboBox_1.currentText()
        com_text2 = self.window.comboBox_2.currentText()
        text=self.window.lineEdit_filter.text()
        if com_text2!='' and com_text1!='' and text!='':
            self.item = self.trees[self.all_dict[com_text2]].topLevelItem(0)
            count = self.item.childCount()  # 获取当前根节点的子节点数量
            for j in range(0, count):
                string = self.item.child(j)  # 子节点的文字信息
                check_name = string.text(0)  # 子节点的文字信息
                if com_text1=='等于':
                    if text == check_name:
                        string.setCheckState(0,Qt.Checked)
                if com_text1=='包含':
                    if text in check_name:
                        string.setCheckState(0, Qt.Checked)
                if com_text1=='大于等于':
                    criteria = r"(\d+)"
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(check_name)
                    text_re=compile.findall(text)
                    if self.results!=[]:
                        if int(text_re[0]) <= int(self.results[0]):
                            string.setCheckState(0, Qt.Checked)
                if com_text1=='小于等于':
                    criteria = r"(\d+)"
                    compile = re.compile(criteria, re.S)
                    self.results = compile.findall(check_name)
                    print(self.results)
                    text_re=compile.findall(text)
                    print(text_re)
                    if self.results!=[]:
                        if int(text_re[0]) >= int(self.results[0]):
                            string.setCheckState(0, Qt.Checked)


    def clear_filter(self):  # 提取材料名称
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
            self.filters_dict = {}
            row_list=[]
            com_text2 = self.window.comboBox_2.currentText()
            print(str(com_text2).replace('','空'))
            self.table_rows = win.tableWidget.rowCount()
            if self.table_rows != '0' and com_text2!='':
                ite = self.trees[self.all_dict[com_text2]].topLevelItem(0) # 循环获取根节点
                self.item = self.trees[self.all_dict[com_text2]].topLevelItem(0)
                self.item_1 = self.item.text(1)  # 获取0列内容
                count = ite.childCount()  # 获取当前根节点的子节点数量
                for j in range(0, count):
                    string = ite.child(j)  # 子节点的文字信息
                    if string.checkState(0) == Qt.Checked:
                        self.filters_dict[str(string.text(0))] = self.all_dict[com_text2]
                if self.filters_dict != {} and self.item_1!='0':
                    print(self.filters_dict)
                    for self.Single_rows in str(self.item_1).split('+'):
                        win.tableWidget.hideRow(int(self.Single_rows))
                        for self.Single_colum in range(0, win.tableWidget.columnCount()):
                            for key in self.filters_dict.keys():
                                self.text = win.tableWidget.item(int(self.Single_rows), self.Single_colum).text()
                                if str(key)==self.text:
                                    row_list.append(self.Single_rows)
                                    win.tableWidget.showRow(int(self.Single_rows))
            if row_list!=[]:
                print(row_list)
                for tree in self.trees:
                    if tree!=self.trees[self.all_dict[com_text2]]:
                        tree.topLevelItem(0).setText(1, str('+'.join(row_list)))
                        tree.expandAll()

    def clear_up(self):#清除检查
        com_text2 = self.window.comboBox_2.currentText()
        if com_text2!='':
            self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:
                self.row_list = []
                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                    self.row_list.append(str(self.Single_rows))
                if self.row_list!=[]:
                    for tre in self.trees:
                        tre.topLevelItem(0).takeChildren()
                        tre.topLevelItem(0).setText(1, str('+'.join(self.row_list)))


class file_window():#导入非标工程量窗口
    def __init__(self):
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "impoty_file_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('导入文件窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.6)
        self.width = int(self.screenwidth * 0.6)
        self.window.resize(self.width, self.height)
        self.window.pushButton_import_file.clicked.connect(self.import_file)
        self.window.comboBox_file.currentIndexChanged.connect(self.click_sheet)
        self.window.pushButton_append_import.clicked.connect(self.append_import)
        # self.window.tableWidget_import_file.clearContents()  # 可以清除表格所有的内容
        self.window.pushButton_clear_import.clicked.connect(self.clear_import)
        self.window.pushButton_GQI_F.clicked.connect(self.GQI_F)#分类工程量
        self.window.pushButton_GQI_B.clicked.connect(self.GQI_B)#分类工程量
        self.window.pushButton_insert_row.clicked.connect(self.insert_row)#插入行
        self.window.pushButton_del_row.clicked.connect(self.del_row)#删除行
        self.window.pushButton_adaptation_row.clicked.connect(self.adaptation_row)#自适应行高

        self.window.tableWidget_import_file.setRowCount(1)
        self.window.tableWidget_import_file.setColumnCount(12)
        self.head = ['', '清单编码', '清单名称', '项目特征', '清单单位', '匹配清单', '匹配定额', '专业/系统', '材料名称', '材料规格',
                     '材料单位', '工程量', '备注','超高工程量','竖井内工程量','管内配线','桥架内配线','配线预留(计入管内)','配线预留(计入桥架)','报表工程量名称']
        file_image = os.path.join(BASE_DIR,"image\图标.png")
        self.head0 = QComboBox(self.window)
        self.head0.addItems(self.head)
        self.head0.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        # self.head0.setStyleSheet("QComboBox { min-height: 50px; min-width: 100px;}"
        #                             "QComboBox QAbstractItemView::item { min-height: 50px; min-width: 80px;}")
        self.head0.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 0, self.head0)
        self.head1 = QComboBox(self.window)
        self.head1.addItems(self.head)
        self.head1.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head1.setView(QListView())

        self.window.tableWidget_import_file.setCellWidget(0, 1, self.head1)
        self.head2 = QComboBox(self.window)
        self.head2.addItems(self.head)
        self.head2.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head2.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 2, self.head2)
        self.head3 = QComboBox(self.window)
        self.head3.addItems(self.head)
        self.head3.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head3.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 3, self.head3)
        self.head4 = QComboBox(self.window)
        self.head4.addItems(self.head)
        self.head4.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head4.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 4, self.head4)
        self.head5 = QComboBox(self.window)
        self.head5.addItems(self.head)
        self.head5.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head5.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 5, self.head5)
        self.head6 = QComboBox(self.window)
        self.head6.addItems(self.head)
        self.head6.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head6.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 6, self.head6)
        self.head7 = QComboBox(self.window)
        self.head7.addItems(self.head)
        self.head7.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head7.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 7, self.head7)
        self.head8 = QComboBox(self.window)
        self.head8.addItems(self.head)
        self.head8.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head8.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 8, self.head8)
        self.head9 = QComboBox(self.window)
        self.head9.addItems(self.head)
        self.head9.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head9.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 9, self.head9)
        self.head10 = QComboBox(self.window)
        self.head10.addItems(self.head)
        self.head10.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head10.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 10, self.head10)
        self.head11 = QComboBox(self.window)
        self.head11.addItems(self.head)
        self.head11.setStyleSheet("QComboBox QAbstractItemView::item { min-height: 25px; min-width: 35px;}")

        self.head11.setView(QListView())
        self.window.tableWidget_import_file.setCellWidget(0, 11, self.head11)
        self.window.tableWidget_import_file.horizontalHeader().setDefaultSectionSize(200)
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint)#窗口始终在前面

    def import_file(self):
        self.import_filePath, _ = QFileDialog.getOpenFileName(self.window, '选择文件', ' ', '文件类型(*.xlsx)')
        if self.import_filePath!='':
            self.window.comboBox_file.clear()
            fs_list = []
            self.wb = openpyxl.load_workbook(self.import_filePath, read_only=True, data_only=False, keep_links=False)
            sheets=self.wb.sheetnames
            for s in sheets:
                filepath_sheet=str(self.import_filePath)+'>'+s
                fs_list.append(filepath_sheet)
            self.window.comboBox_file.addItems(fs_list)

    def click_sheet(self):
        if self.window.comboBox_file.currentText()!='':
            self.window.tableWidget_import_file.setRowCount(1)
            self.tableWidget_row = int(self.window.tableWidget_import_file.rowCount())  # 获取总行数
            self.tableWidget_cols = int(self.window.tableWidget_import_file.columnCount())
            self.num = 0
            sheet = str(self.window.comboBox_file.currentText()).split('>')
            self.wb = openpyxl.load_workbook(sheet[0], read_only=True, data_only=True, keep_links=False)
            self.ws = self.wb[sheet[1]]
            self.values_list = list(self.ws.iter_rows(min_row=None, max_row=None, min_col=None, max_col=None, values_only=True))
            max_num = len(self.values_list)
            for wb_ws_row in range(0, max_num):
                data = self.values_list[wb_ws_row]
                self.num += 1
                self.window.tableWidget_import_file.setRowCount(int(self.tableWidget_row) + self.num)  # 设置行数
                for c in range(0,self.tableWidget_cols):
                    self.window.tableWidget_import_file.setItem(int(self.num) + int(self.tableWidget_row) - 1, c, QTableWidgetItem(''))
                for j in range(0, len(data)):
                    self.window.tableWidget_import_file.setItem(int(self.num) + int(self.tableWidget_row) - 1, j,QTableWidgetItem(str(data[j]).replace('None', '')))
    def append_import(self):
        win.tab_change = '不提取'
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.window.comboBox_file.currentText()!='':
                self.head = ['', '清单编码', '清单名称', '项目特征', '清单单位', '匹配清单', '匹配定额', '专业/系统',
                             '材料名称', '材料规格',
                             '材料单位', '工程量', '备注']
                table_rows=int(win.tableWidget.rowCount())
                table_colums=int(win.tableWidget.columnCount())
                for row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                    win.tableWidget.setRowCount(row+table_rows)
                    for column in range(0,table_colums):
                        win.tableWidget.setItem(row-1+table_rows, column, QTableWidgetItem(''))
                head_text0=self.head0.currentText()
                head_text1=self.head1.currentText()
                head_text2=self.head2.currentText()
                head_text3=self.head3.currentText()
                head_text4=self.head4.currentText()
                head_text5=self.head5.currentText()
                head_text6=self.head6.currentText()
                head_text7=self.head7.currentText()
                head_text8=self.head8.currentText()
                head_text9=self.head9.currentText()
                head_text10=self.head10.currentText()
                head_text11=self.head11.currentText()
                for column,value in enumerate(self.head):
                    if head_text0!='' and head_text0==value:
                        # if head_text0!='超高工程量' or head_text0!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text0=self.window.tableWidget_import_file.item(self.row,0).text()
                            print(table_rows+self.row-1,text0)
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text0)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text1!='' and head_text1==value:
                        # if head_text1!='超高工程量' or head_text1!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text1=self.window.tableWidget_import_file.item(self.row,1).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text1)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text2!='' and head_text2==value:
                        # if head_text2!='超高工程量' or head_text2!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text2=self.window.tableWidget_import_file.item(self.row,2).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text2)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text3!='' and head_text3==value:
                        # if head_text3!='超高工程量' or head_text3!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text3=self.window.tableWidget_import_file.item(self.row,3).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text3)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text4!='' and head_text4==value:
                        # if head_text4!='超高工程量' or head_text4!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text4=self.window.tableWidget_import_file.item(self.row,4).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text4)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text5!='' and head_text5==value:
                        # if head_text5!='超高工程量' or head_text5!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text5=self.window.tableWidget_import_file.item(self.row,5).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text5)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text6!='' and head_text6==value:
                        # if head_text6!='超高工程量' or head_text6!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text6=self.window.tableWidget_import_file.item(self.row,6).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text6)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text7!='' and head_text7==value:
                        # if head_text7!='超高工程量' or head_text7!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text7=self.window.tableWidget_import_file.item(self.row,7).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text7)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text8!='' and head_text8==value:
                        # if head_text8!='超高工程量' or head_text8!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text8=self.window.tableWidget_import_file.item(self.row,8).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text8)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text9!='' and head_text9==value:
                        # if head_text9!='超高工程量' or head_text9!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text9=self.window.tableWidget_import_file.item(self.row,9).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text9)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text10!='' and head_text10==value:
                        # if head_text10!='超高工程量' or head_text10!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text10=self.window.tableWidget_import_file.item(self.row,10).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text10)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                    if head_text11!='' and head_text11==value:
                        # if head_text11!='超高工程量' or head_text11!='竖井内工程量':
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text11=self.window.tableWidget_import_file.item(self.row,11).text()
                            win.tableWidget.setItem(table_rows+self.row-1,column-1,QTableWidgetItem(str(text11)))
                            if column-1==7:
                                win.tableWidget.item(table_rows + self.row - 1, column - 1).setCheckState(Qt.Unchecked)
                win.write_sys()
                win.match_name()
            win.tab_change = '提取'
    def clear_import(self):
        win.tab_change = '不提取'
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.window.comboBox_file.currentText()!='':
                self.head = ['', '清单编码', '清单名称', '项目特征', '清单单位', '匹配清单', '匹配定额', '专业/系统',
                             '材料名称', '材料规格',
                             '材料单位', '工程量', '备注']
                win.tableWidget.clearContents()  # 可以清除表格所有的内容
                table_rows=int(win.tableWidget.rowCount())
                table_colums=int(win.tableWidget.columnCount())
                for row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                    print(row)
                    win.tableWidget.setRowCount(row)
                    for column in range(0,table_colums):
                        win.tableWidget.setItem(row-1, column, QTableWidgetItem(''))
                head_text0=self.head0.currentText()
                head_text1=self.head1.currentText()
                head_text2=self.head2.currentText()
                head_text3=self.head3.currentText()
                head_text4=self.head4.currentText()
                head_text5=self.head5.currentText()
                head_text6=self.head6.currentText()
                head_text7=self.head7.currentText()
                head_text8=self.head8.currentText()
                head_text9=self.head9.currentText()
                head_text10=self.head10.currentText()
                head_text11=self.head11.currentText()
                for column,value in enumerate(self.head):
                    if head_text0!='' and head_text0==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text0=self.window.tableWidget_import_file.item(self.row,0).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text0)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text1!='' and head_text1==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text1=self.window.tableWidget_import_file.item(self.row,1).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text1)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text2!='' and head_text2==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text2=self.window.tableWidget_import_file.item(self.row,2).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text2)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text3!='' and head_text3==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text3=self.window.tableWidget_import_file.item(self.row,3).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text3)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text4!='' and head_text4==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text4=self.window.tableWidget_import_file.item(self.row,4).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text4)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text5!='' and head_text5==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text5=self.window.tableWidget_import_file.item(self.row,5).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text5)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text6!='' and head_text6==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text6=self.window.tableWidget_import_file.item(self.row,6).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text6)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text7!='' and head_text7==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text7=self.window.tableWidget_import_file.item(self.row,7).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text7)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text8!='' and head_text8==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text8=self.window.tableWidget_import_file.item(self.row,8).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text8)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text9!='' and head_text9==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text9=self.window.tableWidget_import_file.item(self.row,9).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text9)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text10!='' and head_text10==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text10=self.window.tableWidget_import_file.item(self.row,10).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text10)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                    if head_text11!='' and head_text11==value:
                        for self.row in range(1,int(self.window.tableWidget_import_file.rowCount())):
                            text11=self.window.tableWidget_import_file.item(self.row,11).text()
                            win.tableWidget.setItem(self.row-1,column-1,QTableWidgetItem(str(text11)))
                            if column-1==7:
                                win.tableWidget.item(self.row-1,column-1).setCheckState(Qt.Unchecked)
                win.write_sys()
                win.match_name()
            win.tab_change = '提取'
    def GQI_F(self):
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.window.comboBox_file.currentText()!='':
                self.head = ['', '清单编码', '清单名称', '项目特征', '清单单位', '匹配清单', '匹配定额', '专业/系统',
                             '材料名称', '材料规格',
                             '材料单位', '工程量', '备注', '超高工程量', '竖井内工程量','管内配线','桥架内配线','配线预留(计入管内)','配线预留(计入桥架)','报表工程量名称']
                table_colums=self.window.tableWidget_import_file.columnCount()
                column_list=[]
                head_list=[]
                head_text0=self.head0.currentText()
                head_text1=self.head1.currentText()
                head_text2=self.head2.currentText()
                head_text3=self.head3.currentText()
                head_text4=self.head4.currentText()
                head_text5=self.head5.currentText()
                head_text6=self.head6.currentText()
                head_text7=self.head7.currentText()
                head_text8=self.head8.currentText()
                head_text9=self.head9.currentText()
                head_text10=self.head10.currentText()
                head_text11=self.head11.currentText()
                for column,value in enumerate(self.head):
                    if head_text0!='' and head_text0==value:
                        column_list.append('0&{}'.format(head_text0))
                        head_list.append(head_text0)
                    if head_text1!='' and head_text1==value:
                        column_list.append('1&{}'.format(head_text1))
                        head_list.append(head_text1)
                    if head_text2!='' and head_text2==value:
                        column_list.append('2&{}'.format(head_text2))
                        head_list.append(head_text2)
                    if head_text3!='' and head_text3==value:
                        column_list.append('3&{}'.format(head_text3))
                        head_list.append(head_text3)
                    if head_text4!='' and head_text4==value:
                        column_list.append('4&{}'.format(head_text4))
                        head_list.append(head_text4)
                    if head_text5!='' and head_text5==value:
                        column_list.append('5&{}'.format(head_text5))
                        head_list.append(head_text5)
                    if head_text6!='' and head_text6==value:
                        column_list.append('6&{}'.format(head_text6))
                        head_list.append(head_text6)
                    if head_text7!='' and head_text7==value:
                        column_list.append('7&{}'.format(head_text7))
                        head_list.append(head_text7)
                    if head_text8!='' and head_text8==value:
                        column_list.append('8&{}'.format(head_text8))
                        head_list.append(head_text8)
                    if head_text9!='' and head_text9==value:
                        column_list.append('9&{}'.format(head_text9))
                        head_list.append(head_text9)
                    if head_text10!='' and head_text10==value:
                        column_list.append('10&{}'.format(head_text10))
                        head_list.append(head_text10)
                    if head_text11!='' and head_text11==value:
                        column_list.append('11&{}'.format(head_text11))
                        head_list.append(head_text11)
                if column_list!=[]:
                    if '工程量' in head_list or '超高工程量' in head_list or '竖井内工程量' in head_list or '管内配线' in head_list or '桥架内配线' in head_list:
                        all_list = []
                        column_list.sort(key=lambda x:int(str(x).split('&')[0]),reverse=False)
                        for row in range(1, int(self.window.tableWidget_import_file.rowCount())):
                            t1_list = []#工程量
                            t2_list=[]#超高工程量
                            t3_list=[]#竖井内工程量
                            t4_list = []  # 管内配线
                            t5_list = []  # 桥架配线
                            for column in column_list:
                                text=str(column).split('&')[1]+'$'+self.window.tableWidget_import_file.item(row,int(str(column).split('&')[0])).text()
                                for order, value in enumerate(self.head):
                                    if value==str(text).split('$')[0]:
                                        if '工程量' in head_list:
                                            if order<=10 or order==11:
                                                t1_list.append(text)
                                        if '超高工程量' in head_list:
                                            if value == str(text).split('$')[0]:
                                                if order <= 10 or order == 13:
                                                    text2=text.replace('材料名称$','材料名称$(超高)')
                                                    t2_list.append(text2)
                                        if '竖井内工程量' in head_list:
                                            if value == str(text).split('$')[0]:
                                                if order <= 10 or order == 14:
                                                    text3 = text.replace('材料名称$', '材料名称$(竖井内)')
                                                    t3_list.append(text3)
                                        if '管内配线' in head_list:
                                            if value == str(text).split('$')[0]:
                                                if order <= 10 or order == 15 or order == 17:
                                                    # if order!=17:
                                                    text4 = text.replace('材料名称$', '材料名称$(管内配线)')
                                                    t4_list.append(text4)
                                        if '桥架内配线' in head_list:
                                            if value == str(text).split('$')[0]:
                                                if order <= 10 or order == 16 or order == 18:
                                                    text5 = text.replace('材料名称$', '材料名称$(桥架内配线)')
                                                    t5_list.append(text5)
                            print(t3_list)
                            if t1_list!=[]:
                                # if str(column).split('&')[1] + '$'+'0'
                                if '工程量$0' not in t1_list:
                                    if '工程量$' not in t1_list:
                                        all_list.append(tuple(t1_list))

                            if t2_list != []:
                                if '超高工程量$0' not in t2_list:
                                    if '超高工程量$' not in t2_list:
                                        all_list.append(tuple(t2_list))
                            if t3_list != []:
                                if '竖井内工程量$0' not in t3_list:
                                    if '竖井内工程量$' not in t3_list:
                                        all_list.append(tuple(t3_list))
                            if t4_list != []:
                                if '管内配线$0' not in t4_list:
                                    if '管内配线$' not in t4_list:
                                        if t4_list[-1].split('$')[0]=='配线预留(计入管内)':
                                            t4=t4_list[:-2]+[t4_list[-2].split('$')[-1]+'+'+t4_list[-1].split('$')[-1]]
                                            print(t4)
                                            all_list.append(tuple(t4))
                                        else:
                                            all_list.append(tuple(t4_list))
                            if t5_list != []:
                                if '桥架内配线$0' not in t5_list:
                                    if '桥架内配线$' not in t5_list:
                                        if t5_list[-1].split('$')[0]=='配线预留(计入桥架内)':
                                            t5=t5_list[:-2]+[t5_list[-2].split('$')[-1]+'+'+t5_list[-1].split('$')[-1]]
                                            print(t5)
                                            all_list.append(tuple(t5))
                                        else:
                                            all_list.append(tuple(t5_list))

                        if all_list != []:
                            self.window.tableWidget_import_file.setRowCount(1)
                            for row,values in enumerate(all_list):
                                self.window.tableWidget_import_file.setRowCount(row+2)
                                for table_colum in range(0,table_colums):
                                    self.window.tableWidget_import_file.setItem(row + 1, table_colum,QTableWidgetItem(''))
                                for column, value in enumerate(values):
                                   self.window.tableWidget_import_file.setItem(row+1, column, QTableWidgetItem(str(value).split('$')[-1]))
    def GQI_B(self):
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if self.window.comboBox_file.currentText() != '':
                self.head = ['', '清单编码', '清单名称', '项目特征', '清单单位', '匹配清单', '匹配定额',
                             '专业/系统',
                             '材料名称', '材料规格',
                             '材料单位', '工程量', '备注', '超高工程量', '竖井内工程量', '管内配线', '桥架内配线',
                             '配线预留(计入管内)', '配线预留(计入桥架)', '报表工程量名称']
                table_colums = self.window.tableWidget_import_file.columnCount()
                column_list = []
                head_list = []
                head_text0 = self.head0.currentText()
                head_text1 = self.head1.currentText()
                head_text2 = self.head2.currentText()
                head_text3 = self.head3.currentText()
                head_text4 = self.head4.currentText()
                head_text5 = self.head5.currentText()
                head_text6 = self.head6.currentText()
                head_text7 = self.head7.currentText()
                head_text8 = self.head8.currentText()
                head_text9 = self.head9.currentText()
                head_text10 = self.head10.currentText()
                head_text11 = self.head11.currentText()
                for column, value in enumerate(self.head):
                    if head_text0 != '' and head_text0 == value:
                        column_list.append('0&{}'.format(head_text0))
                        head_list.append(head_text0)
                    if head_text1 != '' and head_text1 == value:
                        column_list.append('1&{}'.format(head_text1))
                        head_list.append(head_text1)
                    if head_text2 != '' and head_text2 == value:
                        column_list.append('2&{}'.format(head_text2))
                        head_list.append(head_text2)
                    if head_text3 != '' and head_text3 == value:
                        column_list.append('3&{}'.format(head_text3))
                        head_list.append(head_text3)
                    if head_text4 != '' and head_text4 == value:
                        column_list.append('4&{}'.format(head_text4))
                        head_list.append(head_text4)
                    if head_text5 != '' and head_text5 == value:
                        column_list.append('5&{}'.format(head_text5))
                        head_list.append(head_text5)
                    if head_text6 != '' and head_text6 == value:
                        column_list.append('6&{}'.format(head_text6))
                        head_list.append(head_text6)
                    if head_text7 != '' and head_text7 == value:
                        column_list.append('7&{}'.format(head_text7))
                        head_list.append(head_text7)
                    if head_text8 != '' and head_text8 == value:
                        column_list.append('8&{}'.format(head_text8))
                        head_list.append(head_text8)
                    if head_text9 != '' and head_text9 == value:
                        column_list.append('9&{}'.format(head_text9))
                        head_list.append(head_text9)
                    if head_text10 != '' and head_text10 == value:
                        column_list.append('10&{}'.format(head_text10))
                        head_list.append(head_text10)
                    if head_text11 != '' and head_text11 == value:
                        column_list.append('11&{}'.format(head_text11))
                        head_list.append(head_text11)

                if column_list != []:
                    if '报表工程量名称' in head_list:
                        all_list = []
                        column_list.sort(key=lambda x:int(str(x).split('&')[0]),reverse=False)
                        # print(column_list)
                        for row in range(1, int(self.window.tableWidget_import_file.rowCount())):
                            t1_list = []#工程量
                            for column in column_list:
                                text=str(column).split('&')[1]+'$'+self.window.tableWidget_import_file.item(row,int(str(column).split('&')[0])).text()
                                # print(text)
                                for order, value in enumerate(self.head):
                                    if value==str(text).split('$')[0]:
                                        if '报表工程量名称' in head_list:
                                            if order<=10 or order==11 or order==19:
                                                # print(order, value)
                                                t1_list.append(text)
                            if t1_list != []:
                                if '工程量$0' not in t1_list:
                                    if '工程量$' not in t1_list:
                                        for num,va in enumerate(t1_list):
                                            if str(va).split('$')[0]=='材料名称':
                                                num1=num
                                            if str(va).split('$')[0]=='报表工程量名称':
                                                num2=num
                                                name=t1_list[num2].split('$')[-1].replace('数量(个)','')
                                                if name!='':
                                                    t1_list[num1]=t1_list[num1]+'('+t1_list[num2].split('$')[-1]+')'
                                                    t1_list.pop(num2)
                                                    all_list.append(t1_list)

                        if all_list != []:
                            self.window.tableWidget_import_file.setRowCount(1)
                            for row,values in enumerate(all_list):
                                self.window.tableWidget_import_file.setRowCount(row+2)
                                for table_colum in range(0,table_colums):
                                    self.window.tableWidget_import_file.setItem(row + 1, table_colum,QTableWidgetItem(''))
                                for column, value in enumerate(values):
                                   self.window.tableWidget_import_file.setItem(row+1, column, QTableWidgetItem(str(value).split('$')[-1]))

    def insert_row(self):#插入行
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.row = self.window.tableWidget_import_file.currentRow()  # 获取单元格行数
            if self.row!=-1:
                self.window.tableWidget_import_file.insertRow(self.row+1)
                for j in range(0,int(self.window.tableWidget_import_file.columnCount())):
                    print(self.row)
                    item= QTableWidgetItem('')
                    self.window.tableWidget_import_file.setItem(self.row+1, j, item)

    def del_row(self):#删除行
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.rows = []
            self.column = []
            mytable =self.window.tableWidget_import_file.selectedItems()
            for r in mytable:  # 第三种方法获取值'
                if r.row() not in self.rows:
                    self.rows.append(r.row())
                self.column.append(r.column())
            self.rows.sort(reverse=False)
            num = 0
            for i in range(0, len(self.rows)):
                if 0 in self.column and 1 in self.column and 2 in self.column and 3 in self.column and 4 in self.column and 5 in self.column and 6 in self.column \
                        and 7 in self.column and 8 in self.column and 9 in self.column and 10 in self.column:
                    self.window.tableWidget_import_file.removeRow(mytable[num].row())
                    num+=self.window.tableWidget_import_file.columnCount()

    def adaptation_row(self):#自适应行高
        self.Item_row =win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.rows =self.window.tableWidget_import_file.rowCount()  # 获取单元格行数
            if self.rows != 0:
                for i in range(1, self.rows):
                    # self.tableWidget.verticalHeader().resizeSection(i, 80)  # 调整每一行的大小为100像素
                    self.window.tableWidget_import_file.verticalHeader().setSectionResizeMode(i,QHeaderView.ResizeToContents)  # 行高根据内容调整，但是行高不可调


class text_Window():
    def __init__(self):
        super().__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "Text_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('文本修改窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.3)
        self.width = int(self.screenwidth * 0.3)
        self.window.resize(self.width, self.height)
        self.window.pushButton_apply.clicked.connect(self.text)
        # self.window.setWindowModality(Qt.ApplicationModal)#阻塞主窗口不能点击
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint)#窗口始终在前面
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.WindowCloseButtonHint)  # 禁止最大化
    def text(self):
        self.row = win.tableWidget.currentRow()
        self.col = win.tableWidget.currentColumn()
        self.texts = self.window.plainTextEdit.toPlainText()
        win.tableWidget.setItem(self.row, self.col, QTableWidgetItem(self.texts))

class check_Window():# 检查窗口
    def __init__(self):
        super().__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "checks_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('检查窗口')# 设置主窗口的标题
        self.window.desktop = QApplication.desktop()
        self.screenRect = self.window.desktop.screenGeometry()
        self.screenheight = self.screenRect.height()
        self.screenwidth = self.screenRect.width()
        self.height = int(self.screenheight * 0.5)
        self.width = int(self.screenwidth * 0.5)
        self.window.resize(self.width, self.height)
        self.window.pushButton_check.clicked.connect(self.checks)
        self.window.treeWidget_check.setColumnWidth(0, 1500)  # 0列列宽
        self.window.treeWidget_check.itemDoubleClicked.connect(self.location)  # 双击
        # self.window.setWindowModality(Qt.ApplicationModal)#阻塞主窗口不能点击
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint)#窗口始终在前面
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.WindowCloseButtonHint)  # 禁止最大化
        self.tree_font = QFont()
        self.tree_font.setPointSize(9)  # 设置字体大小为9像素
        self.tree_font.setFamily("黑体")
        self.root0 = self.window.treeWidget_check.topLevelItem(0)
        self.root0.setSizeHint(0, QSize(0, 40))
        self.root0.setFont(0, self.tree_font)
        self.root1 = self.window.treeWidget_check.topLevelItem(1)
        self.root1.setSizeHint(0, QSize(0, 40))
        self.root1.setFont(0, self.tree_font)
        self.root2 = self.window.treeWidget_check.topLevelItem(2)
        self.root2.setSizeHint(0, QSize(0, 40))
        self.root2.setFont(0, self.tree_font)
    def checks(self):
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            if win.tableWidget.rowCount() != 0:
                all_row_list=[]
                rows_list=[]
                for self.Single_rows in range(0, int(win.tableWidget.rowCount())):
                    all_row_list.append(self.Single_rows)
                    self.checken_list = []
                    for self.Single_column in range(0, int(win.tableWidget.columnCount())):
                        self.text = win.tableWidget.item(self.Single_rows, self.Single_column).text()
                        self.checken_list.append(self.text)
                    # print(self.checken_list)
                    if '0' not in self.checken_list[0] and len(self.checken_list[0])!=12 and '《定额》'!=self.checken_list[5]:
                        value = '第{}行 {} {}'.format(str(self.Single_rows + 1), self.checken_list[7], self.checken_list[8])
                        self.son = QTreeWidgetItem(self.window.treeWidget_check.topLevelItem(0))
                        self.son.setText(0, value)
                        self.son.setSizeHint(0, QSize(0, 40))
                        self.son.setFont(0, self.tree_font)
                        self.son.setFlags(
                            QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                    if self.checken_list[3]!=self.checken_list[9] and '《定额》'!=self.checken_list[5]:
                        value = '第{}行 {} {}'.format(str(self.Single_rows + 1), self.checken_list[7], self.checken_list[8])
                        self.son = QTreeWidgetItem(self.window.treeWidget_check.topLevelItem(1))
                        self.son.setText(0, value)
                        self.son.setSizeHint(0, QSize(0, 40))
                        self.son.setFont(0, self.tree_font)
                        self.son.setFlags(
                            QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)


                    text = win.tableWidget.item(self.Single_rows, 5).text()
                    if text== '《定额》':
                        rows_list.append(self.Single_rows)
                    if text!= '《定额》':
                        win.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Checked)
                        if win.tableWidget.item(self.Single_rows, 7).checkState() == Qt.Checked:
                            if self.Single_rows+1<=win.tableWidget.rowCount()-1:
                                text5 = win.tableWidget.item(self.Single_rows+1, 5).text()
                                if str(text5) == '《定额》':
                                    rows_list.append(self.Single_rows)

                if rows_list!=[]:
                    for row in rows_list:
                        all_row_list.remove(row)
                for self.Single_rows in all_row_list:
                    text7 = win.tableWidget.item(self.Single_rows, 7).text()
                    text8 = win.tableWidget.item(self.Single_rows, 8).text()
                    value = '第{}行 {} {}'.format(str(self.Single_rows+1), text7, text8)
                    self.son = QTreeWidgetItem(self.window.treeWidget_check.topLevelItem(2))
                    self.son.setText(0, value)
                    self.son.setSizeHint(0, QSize(0, 40))
                    self.son.setFont(0, self.tree_font)
                    self.son.setFlags(
                        QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
                    win.tableWidget.item(self.Single_rows, 7).setCheckState(Qt.Unchecked)
                # self.window.treeWidget_check.expandAll()

    def location(self):
        self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
        if int(self.Item_row) != -1:
            self.tableWidget_allrows = int(win.tableWidget.rowCount())  # 获取总行数
            if self.tableWidget_allrows !=0:
                current=self.window.treeWidget_check.currentItem()
                self.replace_text=current.text(0)
                if '行' in self.replace_text:
                    criteria = '第(.*?)行'
                    compile= re.compile(criteria, re.S)
                    self.results =compile.findall(self.replace_text)
                    location=win.tableWidget.item(int(self.results[0])-1,7)
                    win.tableWidget.scrollToItem(location)# 滚轮定位
                    
class Specifications_window():# 解析规格
    def __init__(self):
        super().__init__()
        # BASE_DIR = os.path.dirname(os.path.realpath(sys.argv[0]))
        BASE_DIR = os.path.dirname(__file__)
        file_path = os.path.join(BASE_DIR, "Specifications_window.ui")
        self.window = uic.loadUi(file_path)
        self.window.setWindowTitle('解析窗')# 设置主窗口的标题
        # self.window.desktop = QApplication.desktop()
        # self.screenRect = self.window.desktop.screenGeometry()
        # self.screenheight = self.screenRect.height()
        # self.screenwidth = self.screenRect.width()
        # self.height = int(self.screenheight * 0.4)
        # self.width = int(self.screenwidth * 0.2)
        # self.window.resize(500, 300)
        self.window.pushButton_1.clicked.connect(self.specifications)#材料解析提取
        self.window.pushButton_circumference.clicked.connect(self.circumference)  # 周长解析
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint)#窗口始终在前面
        self.window.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.WindowCloseButtonHint)  # 禁止最大化

    def specifications(self):# 解析规格
        try:
            self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:
                win.tab_change = '不提取'
                # 撤销
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[win.tableWidget] = self.Click_list
                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)
                row = win.tableWidget.currentRow()  # 获取总行数
                if row !=-1:
                    text1 = self.window.comboBox_1.currentText()  # 获取当前选中的选项的文本
                    text2 = self.window.comboBox_2.currentText()  # 获取当前选中的选项的文本
                    place1=self.window.comboBox_place1.currentText()
                    place_column={'匹配定额名称':5,'材料名称':7,'材料规格':8}
                    print(place1)
                    if text1!='' and text2!='':
                        criteria = '{}(.*?){}'.format(text1,text2)
                        compile = re.compile(criteria, re.S)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict={}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys()))==1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                if r.column()==8 and str(place1)!='材料规格':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(), 7).setCheckState(Qt.Unchecked)
                                if r.column()==7 and str(place1)!='材料名称':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                if r.column()==2:
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容criteria = '{}(.*)'.format(text1)
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(),7).setCheckState(Qt.Unchecked)
                    elif text1!='' and text2=='':
                        criteria = '{}(.*)'.format(text1)
                        compile = re.compile(criteria)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                if r.column() == 8 and str(place1)!='材料规格':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(), 7).setCheckState(Qt.Unchecked)
                                if r.column()==7 and str(place1)!='材料名称':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                if r.column() == 2:
                                    all_names = win.tableWidget.item(r.row(),
                                                                     r.column()).text()  # 获取单元格内容criteria = '{}(.*)'.format(text1)
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(),7).setCheckState(Qt.Unchecked)

                    elif text1=='' and text2!='':
                        criteria = '(.*){}'.format(text2)
                        compile = re.compile(criteria)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                if r.column() == 8 and str(place1)!='材料规格':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(), 7).setCheckState(Qt.Unchecked)
                                if r.column()==7 and str(place1)!='材料名称':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                if r.column() == 2:
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        all_names = win.tableWidget.item(r.row(),
                                                                         r.column()).text()  # 获取单元格内容criteria = '{}(.*)'.format(text1)
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(),7).setCheckState(Qt.Unchecked)
                    elif text1=='' and text2=='':
                        criteria = '(.*)'
                        compile = re.compile(criteria, re.S)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                if r.column() == 8 and str(place1)!='材料规格':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(), 7).setCheckState(Qt.Unchecked)
                                if r.column()==7 and str(place1)!='材料名称':
                                    all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result!=[]:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                if r.column() == 2:
                                    all_names = win.tableWidget.item(r.row(),
                                                                     r.column()).text()  # 获取单元格内容criteria = '{}(.*)'.format(text1)
                                    if win.tableWidget.item(r.row(), 5).text()!='《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            win.tableWidget.setItem(r.row(), place_column[place1], QTableWidgetItem(str(result[0]).replace('(','').replace(')','')))
                                            win.tableWidget.item(r.row(),7).setCheckState(Qt.Unchecked)
                    win.tableWidget.viewport().update()  # 刷新tab内容
                    # 撤回
                    if win.table_do == '执行':
                        self.new_undo_dict = {}
                        new_text_list = []
                        self.new_undo_dict[win.tableWidget] = new_text_list
                        for self.Single_rows in range(0, win.tableWidget.rowCount()):
                            t_list = []
                            new_text_list.append(t_list)
                            for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)
                        if self.Click_list != new_text_list:
                            command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict,
                                                         self.new_undo_dict)
                            win.undoStack_del.push(command)
                    win.tab_change = '提取'
        except:
            pass
    def circumference(self):# 周长解析
        try:
            self.Item_row = win.window.treeWidget_Items.currentIndex().row()  # 获取行
            if int(self.Item_row) != -1:
                win.tab_change = '不提取'
                # 撤销
                self.old_undo_dict = {}
                self.Click_list = []
                self.old_undo_dict[win.tableWidget] = self.Click_list
                for self.Single_rows in range(0, win.tableWidget.rowCount()):
                    t_list = []
                    self.Click_list.append(t_list)
                    for self.Single_colum in range(0, win.tableWidget.columnCount()):
                        self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                        t_list.append(self.text)

                row = win.tableWidget.currentRow()  # 获取总行数
                if row != -1:
                    text1 = self.window.comboBox_3.currentText()  # 获取当前选中的选项的文本
                    text2 = self.window.comboBox_4.currentText()  # 获取当前选中的选项的文本
                    text3 = self.window.comboBox_5.currentText()  # 获取当前选中的选项的文本
                    place2=self.window.comboBox_place2.currentText()
                    place_column={'匹配定额名称':5,'材料规格':8}
                    print(text3)
                    radio_text = self.window.buttonGroup.checkedButton().text()
                    print(radio_text)
                    # print(text1,text2)
                    if text1 != '' and text2 != '' and text3!='':
                        criteria = '{}(.*?){}'.format(text1, text2)
                        compile = re.compile(criteria, re.S)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                if r.column() == 8 and str(place2)!='材料规格':
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text=='周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(2 * cf,0))))
                                            if radio_text=='半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14/2
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                if r.column() == 2:
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2],QTableWidgetItem(str(round(2 * cf, 0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14/2
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2],QTableWidgetItem(str(round(cf, 0))))
                    if text1 != '' and text2 == '' and text3!='':
                        criteria = '{}(.*)'.format(text1)
                        compile = re.compile(criteria)
                        mytable = win.tableWidget.selectedItems()

                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                if r.column() == 8 and str(place2)!='材料规格':
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)

                                        if result != []:
                                            print(result[0])
                                            circumference=str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    cf=int(circumference[0])*3.14
                                                    win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference)>=2:
                                                    cf=(int(circumference[0])+int(circumference[1]))
                                                    win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(2*cf,0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    cf=int(circumference[0])*3.14/2
                                                    win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference)>=2:
                                                    cf=(int(circumference[0])+int(circumference[1]))
                                                    win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                if r.column() == 2:
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        print(result)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    cf=int(circumference[0])*3.14
                                                    win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    cf = (int(circumference[0]) + int(circumference[1]))
                                                    win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(2*cf,0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    cf = int(circumference[0]) * 3.14/2
                                                    win.tableWidget.setItem(r.row(), place_column[place2],
                                                                            QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    cf = (int(circumference[0]) + int(circumference[1]))
                                                    win.tableWidget.setItem(r.row(), place_column[place2],
                                                                            QTableWidgetItem(str(round(cf, 0))))
                    if text1 == '' and text2 != '' and text3 != '':
                        criteria = '(.*){}'.format(text2)
                        compile = re.compile(criteria)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                if r.column() == 8 and str(place2)!='材料规格':
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(2 * cf,0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14/2
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                if r.column() == 2:
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2],QTableWidgetItem(str(round(2 * cf, 0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14/2
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2],QTableWidgetItem(str(round(cf, 0))))

                    if text1 == '' and text2 == '' and text3 != '':
                        criteria = '(.*)'
                        compile = re.compile(criteria, re.S)
                        mytable = win.tableWidget.selectedItems()
                        self.column_dict = {}
                        for r in mytable:  # 第三种方法获取值
                            if r.column() not in self.column_dict:
                                self.column_dict[r.column()] = ''
                        if len(list(self.column_dict.keys())) == 1 and 5 not in list(self.column_dict.keys()):
                            for r in mytable:  # 第三种方法获取值
                                all_names = win.tableWidget.item(r.row(), r.column()).text()  # 获取单元格内容
                                if r.column() == 8 and str(place2)!='材料规格':
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(2 * cf,0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == [] :
                                                        cf=int(circumference[0])*3.14/2
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se==[]:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2], QTableWidgetItem(str(round(cf,0))))
                                if r.column() == 2:
                                    if win.tableWidget.item(r.row(), 5).text() != '《定额》':
                                        result = compile.findall(all_names)
                                        if result != []:
                                            circumference = str(result[0]).split(text3)
                                            if radio_text == '周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2],QTableWidgetItem(str(round(2 * cf, 0))))
                                            if radio_text == '半周长解析':
                                                if len(circumference) == 1:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    if float_re == []:
                                                        cf = int(circumference[0]) * 3.14/2
                                                        win.tableWidget.setItem(r.row(),place_column[place2], QTableWidgetItem(str(round(cf, 0))))
                                                if len(circumference) >= 2:
                                                    float_num = re.compile(r'\D')
                                                    float_re = float_num.findall(circumference[0])
                                                    float_se = float_num.findall(circumference[1])
                                                    if float_re == [] and float_se == []:
                                                        cf = (int(circumference[0]) + int(circumference[1]))
                                                        win.tableWidget.setItem(r.row(), place_column[place2],QTableWidgetItem(str(round(cf, 0))))
                    win.tableWidget.viewport().update()  # 刷新tab内容
                    # 撤回
                    if win.table_do == '执行':
                        self.new_undo_dict = {}
                        new_text_list = []
                        self.new_undo_dict[win.tableWidget] = new_text_list
                        for self.Single_rows in range(0, win.tableWidget.rowCount()):
                            t_list = []
                            new_text_list.append(t_list)
                            for self.Single_colum in range(0, win.tableWidget.columnCount()):
                                self.text = win.tableWidget.item(self.Single_rows, self.Single_colum).text()
                                t_list.append(self.text)
                        if self.Click_list != new_text_list:
                            command = items_tableCommand('单位工程', '单位工程', self.old_undo_dict,
                                                         self.new_undo_dict)
                            win.undoStack_del.push(command)
                    win.tab_change = '提取'
        except:
            pass
if __name__ == '__main__':
    app = QApplication(sys.argv)
    win = Windows()
    BASE_DIR = os.path.dirname(__file__)
    file_path = BASE_DIR + '/' + "image" + '/' + "QLZS.jpg"
    # file_path = os.path.join(BASE_DIR, 'image', "QLZS.jpg")
    app.setWindowIcon(QIcon(file_path))
    className = ("清量计价助手V2.0")
    ext = (".ZJB2.0")
    extDes = ("清量计价工程文件")
    appPath = (sys.argv[0])
    baseUrl = ("HKEY_CURRENT_USER\\Software\\Classes")
    settingClasses = QSettings(baseUrl, QSettings.NativeFormat)
    settingClasses.setValue("/" + className + "/Shell/Open/Command/.", "\"" + appPath + "\" \"%1\"")
    settingClasses.setValue("/" + className + "/.", extDes)
    settingClasses.setValue("/" + className + "/DefaultIcon/.", appPath + ",0")
    settingClasses.setValue("/" + ext + "/OpenWithProgIds/" + className, "")
    settingClasses.sync()
    win.window.show()
    sys.exit(app.exec_())

# pyuic5 -o Manifest_window.py Manifest_window.ui
# pyuic5 -o untitled1.py untitled1.ui


  # def Click_sys(self):  # 获取系统
  #       font = QFont()
  #       font.setPointSize(10)  # 设置字体大小为10像素
  #       font.setFamily("宋体")
  #       n = self.window.treeWidget_system.topLevelItemCount()  # 获取根节点数量
  #       self.window.treeWidget_system.clear()
  #       self.root = QTreeWidgetItem(self.window.treeWidget_system)
  #       self.root.setText(0, '整个工程')
  #       self.root.setFlags(
  #           QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsTristate)
  #       self.root.setText(1, '0')
  #       self.root.setCheckState(0, Qt.Unchecked)
  #       self.window.treeWidget_system.expandAll()
  #       for i in range(0, n):
  #           ite = self.window.treeWidget_system.topLevelItem(i)  # 循环获取根节点
  #           count = ite.childCount()  # 获取当前根节点的子节点数量
  #           for j in range(0, count):
  #               strin = ite.child(j).text(0)  # 子节点的文字信息
  #               if strin != None:
  #                   self.son = QTreeWidgetItem(self.root)
  #                   self.son.setFlags(
  #                       QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
  #                   self.son.setText(0, strin)
  #                   self.son.setSizeHint(0, QSize(0, 40))
  #                   self.son.setCheckState(0, Qt.Unchecked)
  #                   self.son.setFont(0, font)
  #                   print(strin)

# def system_Clicked(self):# 点击系统获取对应的材料
#     font = QFont()
#     font.setPointSize(10)  # 设置字体大小为10像素
#     font.setFamily("宋体")
#     self.tab_change = '不提取'
#     system_dict = {}
#     self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
#     if int(self.Item_row) != -1:  # 先给单位工程打上对钩会返回行-1，值为NONE，所以要不等-1.
#         self.item = self.window.treeWidget_Items.currentItem()
#         self.item_0 = self.item.text(0)  # 获取0列内容
#         self.item_1 = self.item.text(1)  # 获取1列内容
#         self.sys_row = self.window.treeWidget_system.currentIndex().row()  # 获取行
#         if int(self.sys_row) != -1:
#             self.sys = self.window.treeWidget_system.currentItem()
#             self.sys_0=self.sys.text(0)#获取0列内容
#             self.window.treeWidget_system.headerItem().setText(0, self.sys_0)  # 表头写入内容
#             BASE_DIR = os.path.dirname(__file__)
#             files_address = BASE_DIR + '/' + "json_save"
#             for file_name in os.listdir(files_address):
#                 if self.item_0 + '$' + self.item_1 + '.json' ==file_name:
#                     self.table_rows = self.tableWidget.rowCount()
#                     if self.table_rows != '0':
#                         for self.all_row in range(0, self.table_rows):
#                             self.all_values = self.tableWidget.item(self.all_row, 6).text()
#                             self.all_name = self.tableWidget.item(self.all_row, 7).text()
#                             if self.all_values not in system_dict:
#                                 system_dict[self.all_values] = [self.all_name]
#                             elif self.all_values in system_dict and self.all_name not in system_dict[self.all_values]:
#                                 system_dict[self.all_values].append(self.all_name)
#                     # file_path = os.path.join(files_address,self.item_0 + '$' + self.item_1 + '.json')
#                     file_path = os.path.join(files_address,file_name)
#                     # if os.path.exists(file_path):
#                     with open(file_path, 'r') as f:
#                         self.data_json = json.load(f)
#                         self.python_data = json.loads(self.data_json)  # json转python
#                         for key, values in self.python_data[0].items():
#                             num = 0
#                             for value in values:
#                                 if self.sys_0=='整个工程':
#                                     num+=1
#                                     # self.tableWidget.clearContents()  # 可以清除表格所有的内容
#                                     self.tableWidget.setRowCount(num)  # 设置行数
#                                     for colum in range(0,len(value)):
#                                         self.tableWidget.setItem(num-1, colum,QTableWidgetItem(value[colum]))
#                                         if colum==7:
#                                             self.tableWidget.item(num-1, 7).setCheckState(Qt.Unchecked)
#
#                                 if self.sys_0==value[6]:
#                                     num+=1
#                                     self.tableWidget.setRowCount(num)  # 设置行数
#                                     for colum in range(0,len(value)):
#                                         self.tableWidget.setItem(num-1, colum,QTableWidgetItem(value[colum]))
#                                         if colum==7:
#                                             self.tableWidget.item(num-1, 7).setCheckState(Qt.Unchecked)
# if system_dict != {}:
# self.window.treeWidget_system.clear()
# self.root = QTreeWidgetItem(self.window.treeWidget_system)
# self.root.setText(0, '整个工程')
# self.root.setText(1, '0')
# self.root.setCheckState(0, Qt.Unchecked)
# self.window.treeWidget_system.expandAll()
# for sys, name in system_dict.items():
#     self.son = QTreeWidgetItem(self.root)
#     self.son.setFlags(
#         QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
#     self.son.setText(0, sys)
#     self.son.setSizeHint(0, QSize(0, 40))
#     self.son.setCheckState(0, Qt.Unchecked)
#     self.son.setFont(0, font)
# self.window.treeWidget_name.clear()
# self.r_name = QTreeWidgetItem(self.window.treeWidget_name)
# self.r_name.setText(0, '整个工程')
# self.r_name.setText(1, '0')
# self.r_name.setCheckState(0, Qt.Unchecked)
# self.window.treeWidget_name.expandAll()
# for sys, name in system_dict.items():
#     for n in name:
#         self.grandson = QTreeWidgetItem(self.r_name)
#         self.grandson.setText(0, n)
#         self.grandson.setSizeHint(0, QSize(0, 40))
#         self.grandson.setCheckState(0, Qt.Unchecked)
#         self.grandson.setFont(0, font)


# 创建 菜单栏QMenuBar 对象 并返回
# menuBar = QMenuBar(self.window)
# menuBar.resize(500,50)
# menuBar.move(100, 0)
# 一级菜单
# fileMenu = menuBar.addMenu("文件")
# editMenu = menuBar.addMenu("编辑")
# helpMenu = menuBar.addMenu("帮助")
# # 一级Action
# actionHomePage = menuBar.addAction('主页')
# actionHomePage.triggered.connect(self.Excel_values)
# # 二级菜单
# edit_1 = editMenu.addMenu("插入图表")
# editMenu.addSeparator()  # 分隔符
# edit_2 = editMenu.addMenu("插入图片")
#
# # 二级菜单的 action项
# action1 = edit_1.addAction("action1")
# edit_1.addSeparator()  # 分隔符
# action2 = edit_1.addAction("action2")
# sun_menu = QMenu(self.menu)
# self.menu.addMenu(sun_menu)  # 先在主菜单栏中添加一个子菜单
# sun_menu.setTitle('子菜单标题')  # 设置子菜单标题
# self.window.toolButton_hide_column.setContextMenuPolicy(Qt.CustomContextMenu)


# selectRect = self.tableWidget.selectedRanges()
# for r in selectRect:  # 获取范围边界
#     self.top = r.topRow()
#     # print(self.top)
#     self.left = r.leftColumn()
#     self.bottom = r.bottomRow()
#     self.right = r.rightColumn()
# self.column_n = 0
# self.number = 0
# self.row_n = 0
# self.column_n = self.right - self.left + 1
# self.row_n = self.bottom - self.top + 1
# self.number = self.row_n * self.column_n
# self.c = []
# for i in range(self.number):
#     self.c.append(' \t')  # 注意，是空格+\t
#     if (i % self.column_n) == (self.column_n - 1):
#         self.c.append('\n')
#     else:
#         pass
#     # 这里生成了一个列表，大小是：行X（列+1），换行符占了一列。
#     # 默认情况下，列表中全部是空格，
# self.c.pop()  # 删去最后多余的换行符

'''
self.page = QWidget()
self.verticalLayout_10 = QtWidgets.QVBoxLayout(self.page)
self.table = QTableWidget(self.page)
self.table.resize(1000, 500)
self.table.move(500,200)
self.table.setRowCount(1)
self.table.setColumnCount(10)
self.table.horizontalHeader()
self.table.setHorizontalHeaderLabels(['项目编码', '项目名称', '项目特征', '单位', '匹配清单', '匹配定额','专业/系统','材料名称','材料规格','材料单位','工程量'])
self.window.stackedWidget.addWidget(self.page)
'''


# def new_Unit(self):
#     pass
    # self.item = self.window.treeWidget_Items.currentItem()
    # self.Unit = QTreeWidgetItem(self.item)
    # self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
    # self.Unit.setText(0, '单位工程')
    # self.Unit.setText(1, str(self.Item_row))
    # self.Unit.setFlags(QtCore.Qt.ItemIsSelectable 单击选中| QtCore.Qt.ItemIsEditable(文本是否可编辑) | QtCore.Qt.ItemIsDragEnabled (可拖拉)| QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable用户检查 | QtCore.Qt.ItemIsEnabled文本灰显不可编辑 | QtCore.Qt.ItemIsTristate半选中状态)
    # self.window.treeWidget_Items.expandAll()
    #
    # self.page = QWidget()
    # self.verticalLayout_10 = QtWidgets.QVBoxLayout(self.page)
    # self.window.stackedWidget.addWidget(self.page)
    # self.table = QTableWidget(self.page)
    # self.table.resize(1000, 500)
    # self.table.move(500,200)
    # self.table.setRowCount(10)
    # self.table.setColumnCount(10)
    # self.table.horizontalHeader()
    # self.table.setHorizontalHeaderLabels(['项目编码', '项目名称', '项目特征', '单位', '匹配清单', '匹配定额','专业/系统','材料名称','材料规格','材料单位','工程量'])
    # self.window.stackedWidget.addWidget(self.page)


# def Click_Select(self):  # 单击获取表格内任意值，并且或撤销值
#     pass
    # self.t7_list=[]
    # self.Item_row = self.window.treeWidget_Items.currentIndex().row()  # 获取行
    # if int(self.Item_row) != -1:
    #     if self.tableWidget.currentRow()!=-1 :
    #         t7 = self.tableWidget.item(self.tableWidget.currentRow(), 7).text()
    #         self.old_undo_dict={}
    #         self.Click_list = []
    #         self.old_undo_dict[self.tableWidget] = self.Click_list
    #         for self.Single_rows in range(0, self.tableWidget.rowCount()):
    #             t_list = []
    #             self.Click_list.append(t_list)
    #             for self.Single_colum in range(0, self.tableWidget.columnCount()):
    #                 self.text = self.tableWidget.item(self.Single_rows, self.Single_colum).text()
    #                 t_list.append(self.text)
    #         self.tab_change='变化'


# table_do = '执行'
# new_undo_dict = {}

#
# def undo_redo(self):  # 文本内容变化后转让撤销类中
#     pass
    # try:
    #     if self.tab_change == '变化':
    #         if self.tableWidget.currentRow() != -1 and self.tableWidget.currentRow() != 0:#复选框打上对√如果返回行号为0，无法获取总行和总列
    #             self.new_undo_dict = {}
    #             new_text_list = []
    #             # print( self.tableWidget.rowCount(),self.tableWidget.columnCount())
    #             self.new_undo_dict[self.tableWidget] = new_text_list
    #             for self.Single_rows in range(0, self.tableWidget.rowCount()):
    #                 t_list=[]
    #                 new_text_list.append(t_list)
    #                 for self.Single_colum in range(0, self.tableWidget.columnCount()):
    #                     self.text =self.tableWidget.item(self.Single_rows, self.Single_colum).text()
    #                     t_list.append(self.text)
    #             if self.old_undo_dict!=self.new_undo_dict and self.old_undo_dict!={} and self.new_undo_dict!={}:
    #                 command = items_tableCommand('单位工程','单位工程',self.old_undo_dict, self.new_undo_dict)
    #                 self.undoStack_del.push(command)
    #                 self.tab_change='还原'
    #
    # except:
    #     pass



# def open_double_file(self,jsonPath):
    #     with open(jsonPath, 'r') as f:
    #         self.data_json = json.load(f)
    #         self.python_data = json.loads(self.data_json)  # json转python
    #         for value in self.python_data:
    #             for key, value in value[0].items():
    #                 print(str(key).split('$')[1], value)
    #                 if str(key).split('$')[0]=='建设项目':
    #                     self.tableWidget = self.tablelist[0]
    #                     self.window.treeWidget_Items.topLevelItem(0).setText(2, str(key).split('$')[1])
    #                     for row,va in enumerate(value):
    #                         self.tableWidget.setRowCount(row + 1)  # 设置行数
    #                         for column ,text in enumerate(va):
    #                             self.tableWidget.setItem(row,column, QTableWidgetItem(''))
    #                             self.item1 = QTableWidgetItem(str(text))
    #                             if column == 7:
    #                                 self.item1.setCheckState(Qt.Unchecked)
    #                                 self.tableWidget.setItem(row,column, self.item1)
    #                 else:
    #                     self.son = QTreeWidgetItem(self.window.treeWidget_Items.topLevelItem(0))
    #                     # self.son.setFlags(QtCore.Qt.ItemIsSelectable | QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsDragEnabled | QtCore.Qt.ItemIsDropEnabled | QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
    #                     self.son.setText(0, str(key).split('$')[0])
    #                     self.son.setText(1, str(key).split('$')[1])
    #                     self.son.setCheckState(0, Qt.Unchecked)
    #                     self.window.treeWidget_Items.topLevelItem(0).setCheckState(0, Qt.Unchecked)
    #                     self.window.treeWidget_Items.topLevelItem(0).setSizeHint(0, QSize(0, 30))
    #                     self.son.setSizeHint(0, QSize(0, 30))
    #                     font = QFont()
    #                     font.setPointSize(10)  # 设置字体大小为10像素
    #                     font.setFamily("宋体")
    #                     self.son.setFont(0, font)
    #                     self.tableWidget = self.tablelist[int(str(key).split('$')[1])]
    #                     for row,va in enumerate(value):
    #                         self.tableWidget.setRowCount(row + 1)  # 设置行数
    #                         for column ,text in enumerate(va):
    #                             self.tableWidget.setItem(row,column, QTableWidgetItem(''))
    #                             self.item1 = QTableWidgetItem(str(text))
    #                             if column == 7:
    #                                 self.item1.setCheckState(Qt.Unchecked)
    #                                 self.tableWidget.setItem(row,column, self.item1)


# self.label_nowTime = QLabel('初始内容',self.window)
# self.label_nowTime.resize(2000,30)
# self.label_nowTime.setText('')
# self.label_nowTime.move(800,10)
# self.label_nowTime.setFont(font)
# self.movie = QMovie(r"C:\Users\zhang\Desktop\无标题.gif")
# self.label_nowTime.setMovie(self.movie)
# self.movie.start()
# self.window.pushButton_Export_Excel.setEnabled(False)